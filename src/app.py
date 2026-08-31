import streamlit as st
import os
import base64
import time
import httpx
import json
import glob
import markdown
import subprocess
import fitz  # PyMuPDF
import re
from datetime import datetime
from docx import Document
from openai import OpenAI
from playwright.sync_api import sync_playwright
from apscheduler.schedulers.background import BackgroundScheduler
import sys
import wechat_publisher
import prompts
import image_engine
import wechat_render
import llm
from cli_paths import NPX_CMD, DREAMINA_CMD
import zsxq_client
import scheduler
from config_store import (load_config, save_config, get_exec_env, load_indicators, save_indicator, delete_indicator)
import prompts
import image_engine
import wechat_render
import llm
import importlib
importlib.reload(wechat_publisher)


# --- Config & Constants ---
IMAGE_OUTPUT_DIR = os.path.join("outputs", "images")
SCRIPT_OUTPUT_DIR = os.path.join("outputs", "scripts")
INDICATOR_DOCS_DIR = os.path.join("outputs", "indicator_docs")
WECHAT_OUTPUT_DIR = os.path.join("outputs", "wechat")
WECHAT_IMAGES_DIR = os.path.join("outputs", "wechat", "images")
os.makedirs(IMAGE_OUTPUT_DIR, exist_ok=True)
os.makedirs(SCRIPT_OUTPUT_DIR, exist_ok=True)
os.makedirs(INDICATOR_DOCS_DIR, exist_ok=True)
os.makedirs(WECHAT_OUTPUT_DIR, exist_ok=True)
os.makedirs(WECHAT_IMAGES_DIR, exist_ok=True)










if "config" not in st.session_state:
    st.session_state.config = load_config()

config = st.session_state.config





def render_wechat_preview(md_text):
    import re
    # split text by markdown images
    parts = re.split(r'(!\[.*?\]\(.*?\))', md_text)
    for part in parts:
        if part.startswith("!["):
            m = re.match(r'!\[(.*?)\]\((.*?)\)', part)
            if m:
                img_path = m.group(2)
                if img_path.startswith("http"):
                    st.image(img_path)
                elif os.path.exists(img_path):
                    st.image(img_path)
                else:
                    st.markdown(part)
        else:
            if part.strip():
                st.markdown(part, unsafe_allow_html=True)





# Default Config Setup
defaults = {
    "user_groups": {},
    "groups": {"默认群组": ""},
    "platform": "自定义/OpenAI",
    "mode": "常规总结",
    "selected_group": "默认群组",
    "auto_run": False,
    "run_time": "08:00",
    "channel_configs": {
        "自定义/OpenAI": {"api_key": "", "base_url": "https://api.openai.com/v1", "selected_model": "gpt-4o", "available_models": ["gpt-4o", "gpt-3.5-turbo"]},
        "火山方舟 (Volcengine)": {"api_key": "", "base_url": "https://ark.cn-beijing.volces.com/api/v3", "selected_model": "", "available_models": []},
        "魔塔 (ModelScope)": {"api_key": "", "base_url": "https://api.modelscope.cn/v1", "selected_model": "", "available_models": []},
        "DeepSeek": {"api_key": "", "base_url": "https://api.deepseek.com", "selected_model": "deepseek-chat", "available_models": ["deepseek-chat", "deepseek-reasoner"], "enable_thinking": True, "reasoning_effort": "high"}
    },
    "custom_prompts": {"合规化处理": "请对以上内容进行合规化处理：1. 隐藏具体的人名和联系方式；2. 增加‘以上内容仅供参考，不构成投资建议’的免责声明；3. 语气调整为客观中立的中台视角。"}
}
for key, val in defaults.items():
    if key not in config:
        config[key] = val

# Ensure all default channels exist in channel_configs
if "channel_configs" not in config:
    config["channel_configs"] = defaults["channel_configs"].copy()
else:
    for plat, plat_cfg in defaults["channel_configs"].items():
        if plat not in config["channel_configs"]:
            config["channel_configs"][plat] = plat_cfg.copy()

# Migrate old config to channel_configs specifically, and clean up polluted platforms
if "api_key" in config and config["api_key"]:
    legacy_key = config.get("api_key", "")
    legacy_url = config.get("base_url", "")
    legacy_model = config.get("selected_model", "")
    legacy_models = config.get("available_models", [])
    
    # Check if the legacy URL belongs to ModelScope
    is_modelscope = "modelscope.cn" in legacy_url
    target_plat = "魔塔 (ModelScope)" if is_modelscope else "自定义/OpenAI"
    
    # Migrate Specifically to target_plat if it doesn't already have an api_key
    target_cfg = config["channel_configs"].setdefault(target_plat, {})
    if not target_cfg.get("api_key"):
        target_cfg["api_key"] = legacy_key
        target_cfg["base_url"] = legacy_url if legacy_url else defaults["channel_configs"][target_plat]["base_url"]
        target_cfg["selected_model"] = legacy_model if legacy_model else defaults["channel_configs"][target_plat]["selected_model"]
        target_cfg["available_models"] = legacy_models if legacy_models else defaults["channel_configs"][target_plat]["available_models"]
        
    # Reset other channels if they were polluted by legacy_key
    for plat, default_cfg in defaults["channel_configs"].items():
        if plat in config["channel_configs"] and plat != target_plat:
            chan_cfg = config["channel_configs"][plat]
            if chan_cfg.get("api_key") == legacy_key:
                config["channel_configs"][plat] = default_cfg.copy()
                
    # Remove all legacy root keys
    for k in ["api_key", "base_url", "selected_model", "available_models"]:
        config.pop(k, None)
    save_config(config)

if "available_models" not in st.session_state:
    st.session_state.available_models = config["channel_configs"][config.get("platform", "自定义/OpenAI")].get("available_models", [])

if "scheduler" not in st.session_state:
    st.session_state.scheduler = scheduler.get_scheduler()

if "virtual_history" not in st.session_state:
    st.session_state.virtual_history = []






MANUAL_TASK_STATE_FILE = os.path.join("outputs", "manual_task_state.json")
MANUAL_LOG_FILE = os.path.join("outputs", "manual_execution.log")












scheduler.update_scheduler()

@st.dialog("✨ 动态 Cron 表达式配置器")
def cron_configurator_dialog(page_selection, curr_cron):
    st.info("💡 提示：留空表示不限制（即 *）。侧边栏支持直接手写诸如 `*/5` 这种高阶语法。")
    
    parts = curr_cron.split()
    if len(parts) >= 5:
        c_min, c_hour, c_dom, c_mon, c_dow = parts[0], parts[1], parts[2], parts[3], parts[4]
    else:
        c_min, c_hour, c_dom, c_mon, c_dow = "0", "8", "*", "*", "*"

    def safe_parse(val, opts):
        return [v for v in val.split(',') if v in opts] if val != '*' else []
    
    min_opts = [str(i) for i in range(60)]
    hour_opts = [str(i) for i in range(24)]
    dom_opts = [str(i) for i in range(1, 32)]
    mon_opts = [str(i) for i in range(1, 13)]
    dow_opts = [str(i) for i in range(7)]

    col1, col2 = st.columns(2)
    with col1:
        sel_min = st.multiselect("分钟 (多选)", min_opts, default=safe_parse(c_min, min_opts), placeholder="留空为 *")
        sel_dom = st.multiselect("日期 (多选)", dom_opts, default=safe_parse(c_dom, dom_opts), placeholder="留空为 *")
        sel_dow = st.multiselect("星期 (多选, 0为周日)", dow_opts, default=safe_parse(c_dow, dow_opts), placeholder="留空为 *")
    with col2:
        sel_hour = st.multiselect("小时 (多选)", hour_opts, default=safe_parse(c_hour, hour_opts), placeholder="留空为 *")
        sel_mon = st.multiselect("月份 (多选)", mon_opts, default=safe_parse(c_mon, mon_opts), placeholder="留空为 *")
        
    n_min = ",".join(sel_min) if sel_min else "*"
    n_hour = ",".join(sel_hour) if sel_hour else "*"
    n_dom = ",".join(sel_dom) if sel_dom else "*"
    n_mon = ",".join(sel_mon) if sel_mon else "*"
    n_dow = ",".join(sel_dow) if sel_dow else "*"
    
    new_cron = f"{n_min} {n_hour} {n_dom} {n_mon} {n_dow}"
    
    st.markdown(f"**当前生成的表达式:** `{new_cron}`")
    
    try:
        from apscheduler.triggers.cron import CronTrigger
        from datetime import datetime
        trigger = CronTrigger.from_crontab(new_cron)
        now = datetime.now()
        runs = []
        from datetime import timedelta
        for _ in range(5):
            now = trigger.get_next_fire_time(None, now)
            runs.append(now.strftime("%Y-%m-%d %H:%M:%S"))
            now = now + timedelta(seconds=1)
        st.success("🔮 **未来 5 次执行时间预演:**\n\n" + "\n".join([f"- {r}" for r in runs]))
        is_valid = True
    except Exception as e:
        st.error(f"❌ 表达式无效或无法解析，请检查输入格式。({str(e)})")
        is_valid = False

    if st.button("✅ 确定使用此配置", use_container_width=True, disabled=not is_valid):
        st.session_state[f"cron_input_{page_selection}"] = new_cron
        st.rerun()


# --- Helpers ---









# --- Streamlit UI ---
st.set_page_config(page_title="AI 多功能控制台", layout="wide")

if 'show_success_toast' in st.session_state:
    st.toast(st.session_state['show_success_toast'], icon="🎉")
    st.balloons()
    del st.session_state['show_success_toast']

with st.sidebar:
    st.header("🧭 导航")
    page_selection = st.radio("选择功能模块", ["AI 深度分析", "视频脚本制作器", "指标文档制作"], label_visibility="collapsed")
    st.markdown("---")
    
    st.header("⚙️ 全局配置")
    
    @st.cache_data(ttl=300, show_spinner=False)
    def get_zsxq_auth_status():
        res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8', env=get_exec_env())
        return res.stdout

    @st.cache_data(ttl=300, show_spinner=False)
    def get_dreamina_credit_status():
        res = subprocess.run([DREAMINA_CMD, "user_credit"], capture_output=True, text=True, encoding='utf-8')
        return res.returncode, res.stdout
        
    def clear_auth_cache():
        get_zsxq_auth_status.clear()
        get_dreamina_credit_status.clear()
        
    if hasattr(st, "dialog") and "show_log_dialog" not in globals():
        @st.dialog("📋 任务执行日志", width="large")
        def show_log_dialog(log_content, page_selection="AI 深度分析"):
            try:
                cfg = load_config()
                da_sched = cfg.get("schedulers", {}).get(page_selection, {})
                ui_state = da_sched.get("ui_state", {})
                if ui_state:
                    st.markdown(f"**📢 当前定时任务配置 ({page_selection})：**")
                    if page_selection == "AI 深度分析":
                        c1, c2, c3, c4 = st.columns(4)
                        with c1:
                            st.caption(f"👥 **星球群组**：{ui_state.get('selected_group', '默认')}")
                            st.caption(f"⏰ **调度周期**：{da_sched.get('cron_expr', '未配置')}")
                            st.caption(f"🎯 **获取范围**：{ui_state.get('scope_ui', '最新总结 (话题+文件)')}")
                        with c2:
                            st.caption(f"📝 **分析模式**：{ui_state.get('a_mode', '常规')}")
                            st.caption(f"🔢 **拉取篇数**：{ui_state.get('l_limit', 3)} 篇")
                        with c3:
                            st.caption(f"📱 **微信推文**：{'开启' if ui_state.get('use_wechat') else '关闭'}")
                            st.caption(f"🧭 **推文定位**：{ui_state.get('wechat_article_orientation', '产业宏观与行业趋势')}")
                            st.caption(f"🎨 **排版模式**：{ui_state.get('wechat_mode', 'AI模式')}")
                            if ui_state.get('use_wechat'):
                                st.caption(f"🚀 **发布模式**：{ui_state.get('wechat_publish_mode', '仅生成本地文件 (不上传)')}")
                        with c4:
                            st.caption(f"🍁 **排版主题**：{ui_state.get('wechat_theme', 'spring-fresh')}")
                            st.caption(f"📐 **字号/背景**：{ui_state.get('wechat_font_size', 'medium')} / {ui_state.get('wechat_background_type', 'none')}")
                            if ui_state.get('use_wechat'):
                                st.caption(f"🎯 **发布账号**：{ui_state.get('wechat_publish_account', '未指定')}")
                    else:
                        st.json(ui_state)
                    st.divider()
            except Exception as e:
                st.caption(f"无法加载配置快照: {str(e)}")
            st.text_area("", value=log_content, height=500, disabled=True)
            
    if hasattr(st, "dialog") and "show_wechat_publish_dialog" not in globals():
        @st.dialog("📲 微信公众号一键发布与预览", width="large")
        def show_wechat_publish_dialog():
            import urllib.parse
            
            # 自动从关联 of .draft.json 中加载草稿状态，防止页面刷新丢失
            filepath = st.session_state.get('latest_wechat_file_path')
            draft_info = wechat_publisher.load_draft_info(filepath) if filepath else {}
            
            is_published = (draft_info.get("status") == "published")
            is_scheduled = (draft_info.get("status") == "scheduled")
            
            if is_published:
                st.error(f"⚠️ **防重群发拦截**：该文章已于 `{draft_info.get('publish_time', '此前')}` 正式群发到公众号，不可重复操作！")
                if st.button("🔓 解除群发锁定并允许重新发布", use_container_width=True):
                    wechat_publisher.save_draft_info(
                        filepath,
                        draft_info.get("media_id"),
                        draft_info.get("url"),
                        status="draft",
                        publish_time=""
                    )
                    st.success("已成功解除群发锁定！")
                    st.rerun()
            elif is_scheduled:
                st.info(f"📅 **定时发布中**：该文章已排期于 `{draft_info.get('scheduled_time', '未来')}` 自动定时群发。")
                if st.button("🔓 取消定时发布并解除锁定", use_container_width=True):
                    try:
                        job_id = f"schedule_publish_{draft_info.get('media_id')}"
                        st.session_state.scheduler.remove_job(job_id)
                    except Exception as e:
                        pass
                    wechat_publisher.save_draft_info(
                        filepath,
                        draft_info.get("media_id"),
                        draft_info.get("url"),
                        status="draft",
                        scheduled_time=""
                    )
                    st.success("已成功取消定时并解除锁定！")
                    st.rerun()
            
            if "wechat_draft_media_id" not in st.session_state or st.session_state.wechat_draft_media_id is None:
                st.session_state.wechat_draft_media_id = draft_info.get("media_id")
            if "wechat_draft_url" not in st.session_state or st.session_state.wechat_draft_url is None:
                st.session_state.wechat_draft_url = draft_info.get("url")
            if "wechat_publish_result" not in st.session_state:
                st.session_state.wechat_publish_result = None
            if "wechat_preview_status" not in st.session_state:
                st.session_state.wechat_preview_status = None
                
            st.markdown("### 1. 公众号账号选择与管理")
            accounts = wechat_publisher.load_accounts()
            
            account_names = [a["name"] for a in accounts]
            selected_acc_name = st.selectbox(
                "选择当前操作的公众号", 
                account_names + ["➕ 新增公众号账号配置..."], 
                index=0 if account_names else 0
            )
            
            active_account = None
            if selected_acc_name == "➕ 新增公众号账号配置...":
                with st.form("add_wechat_acc_form"):
                    st.write("🔑 **添加公众号开发者凭证 (AppID / AppSecret)**")
                    new_name = st.text_input("公众号名称 (例如：慧峰金融)", placeholder="请输入公众号名称")
                    new_appid = st.text_input("开发者 ID (AppID)", placeholder="请输入 AppID")
                    new_secret = st.text_input("开发者密码 (AppSecret)", type="password", placeholder="请输入 AppSecret")
                    
                    submitted = st.form_submit_button("💾 保存配置", use_container_width=True)
                    if submitted:
                        if new_name and new_appid and new_secret:
                            accounts = [a for a in accounts if a["name"] != new_name.strip()]
                            accounts.append({
                                "name": new_name.strip(),
                                "appid": new_appid.strip(),
                                "secret": new_secret.strip()
                            })
                            wechat_publisher.save_accounts(accounts)
                            st.success(f"公众号「{new_name}」配置保存成功！")
                            st.rerun()
                        else:
                            st.error("❌ 所有字段均为必填项！")
            elif selected_acc_name:
                active_account = next((a for a in accounts if a["name"] == selected_acc_name), None)
                if active_account:
                    col_det, col_del = st.columns([5, 1.5])
                    with col_det:
                        st.caption(f"🛡️ AppID: `{active_account['appid'][:6]}******` | 密钥已妥善加密保存")
                    with col_del:
                        if st.button("🗑️ 删除该公众号", key="del_wechat_acc_btn", use_container_width=True):
                            accounts = [a for a in accounts if a["name"] != selected_acc_name]
                            wechat_publisher.save_accounts(accounts)
                            st.success("账号已删除")
                            st.rerun()
                            
            if not active_account:
                st.info("💡 请先添加或选择一个有效的微信公众号。")
                return
                
            st.divider()
            
            st.markdown("### 2. IP 白名单安全配置")
            with st.spinner("正在获取服务器外网出口 IP..."):
                server_ip = wechat_publisher.get_server_ip()
            st.info(
                f"🖥️ **当前服务器公网 IP：`{server_ip}`**\n\n"
                "⚠️ **重要提示**：请登录「微信公众平台 -> 开发接口管理 -> IP白名单」，把上述公网 IP 加进去。否则接口调用会返回 `ip not in whitelist` 报错。"
            )
            
            st.divider()
            
            st.markdown("### 3. 推文封面图及基本元数据")
            
            latest_wechat = st.session_state.get('latest_wechat', '')
            parsed_title = "未命名推文"
            for line in latest_wechat.split("\n"):
                line = line.strip()
                if line.startswith("#"):
                    t = re.sub(r'^#+\s*', '', line).strip()
                    if t:
                        parsed_title = t
                        break
            if len(parsed_title) > 32:
                parsed_title = parsed_title[:32]
                
            meta_title = st.text_input("推文标题 (微信要求 ≤ 32字)", value=parsed_title)
            meta_author = st.text_input("作者名称 (建议 ≤ 8字)", value=config.get("wechat_publish_author", ""))
            
            # 强健地过滤各种 markdown 图片格式（兼容空格及换行）
            clean_text = re.sub(r'!\[[^\]]*\]\s*\([^)]*\)', '', latest_wechat)
            # 强健地过滤 markdown 链接格式（兼容空格）
            clean_text = re.sub(r'\[[^\]]*\]\s*\([^)]*\)', '', clean_text)
            # 过滤任何可能存在的 HTML 标签
            clean_text = re.sub(r'<[^>]*>', '', clean_text)
            # 安全防线：彻底清除可能泄露的本地图片路径特征，确保摘要纯净
            clean_text = re.sub(r'outputs[/\\]wechat[/\\]images[/\\][^\s]+', '', clean_text)
            clean_text = re.sub(r'\(\s*outputs[/\\]wechat[/\\]images[/\\][^\s]*\)', '', clean_text)
            # 清理特殊排版字符及多余空格
            clean_text = re.sub(r'[#\*_`\-\>\+\n\r\t]', ' ', clean_text)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            default_digest = clean_text[:80] + "..." if len(clean_text) > 80 else clean_text
            meta_digest = st.text_area("推文摘要 (微信要求 ≤ 128字)", value=default_digest, height=70)
            
            st.markdown("**封面图预览与设置**")
            local_images = []
            # 强健识别包含空格在内的所有配图路径，支持 WebP
            img_pattern = re.compile(r'!\[.*?\]\s*\(\s*([^)\s]+\.(?:jpg|png|jpeg|webp))\s*\)', re.IGNORECASE)
            found_imgs = img_pattern.findall(latest_wechat)
            for img in found_imgs:
                img_path = img.replace("\\", "/").strip()
                # 兼容相对路径、绝对路径、前导斜杠以及不同系统的路径前缀
                paths_to_check = [img_path]
                if img_path.startswith("/"):
                    paths_to_check.append(img_path[1:])
                else:
                    paths_to_check.append("/" + img_path)
                
                resolved_path = None
                for p in paths_to_check:
                    if os.path.exists(p) and not os.path.isdir(p):
                        resolved_path = p
                        break
                
                if resolved_path and resolved_path not in local_images:
                    local_images.append(resolved_path)
                    
            cover_mode_sel = st.radio("选择封面图来源", ["从推文已生成的配图中选择", "手动上传自定义封面图片"], horizontal=True)
            
            selected_cover_path = None
            if cover_mode_sel == "从推文已生成的配图中选择":
                if local_images:
                    selected_cover_path = st.selectbox("选择配图", local_images, format_func=lambda x: os.path.basename(x))
                    if selected_cover_path:
                        st.image(selected_cover_path, width=240, caption="选定的封面图")
                else:
                    st.warning("⚠️ 推文中未检测到本地配图，请选择手动上传封面。")
            else:
                uploaded_cover = st.file_uploader("上传 JPG/PNG 格式封面 (建议宽高比 2.35:1)", type=["jpg", "png", "jpeg"])
                if uploaded_cover:
                    os.makedirs("tests", exist_ok=True)
                    temp_cover_path = os.path.join("tests", f"temp_upload_cover_{int(time.time())}.jpg")
                    with open(temp_cover_path, "wb") as f_cover:
                        f_cover.write(uploaded_cover.getbuffer())
                    selected_cover_path = temp_cover_path
                    st.image(uploaded_cover, width=240, caption="上传的自定义封面")
                    
            st.divider()
            
            st.markdown("### 4. 预览与正式发布控制台")
            
            c_action1, c_action2, c_action3 = st.columns(3)
            
            with c_action1:
                if st.button("1️⃣ 生成草稿箱文章", use_container_width=True, type="primary"):
                    if not selected_cover_path:
                        st.error("❌ 请先设置封面图！")
                        return
                    with st.spinner("🚀 正在上传所有配图、上传封面图、并在微信后台生成草稿..."):
                        try:
                            token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                            
                            # 优先重用已经预览生成好的 HTML 缓存，保证内容/加粗重点字/排版与本地预览 100% 绝对一致
                            if st.session_state.get('latest_wechat_html_raw'):
                                raw_html = st.session_state['latest_wechat_html_raw']
                            else:
                                raw_html = wechat_render.convert_to_wechat_html(
                                    st.session_state['latest_wechat'],
                                    st.session_state.get('latest_wechat_rendered_theme', 'spring-fresh'),
                                    st.session_state.get('latest_wechat_rendered_mode', 'AI 模式 (免费)'),
                                    api_key=config.get("md2wechat_api_key", ""),
                                    font_size=config.get("wechat_font_size", "medium"),
                                    bg_type=config.get("wechat_background_type", "none"),
                                    chan_config=current_chan_config,
                                    custom_prompt=config.get("wechat_custom_prompt", ""),
                                    for_wechat_api=True
                                )
                            
                            theme_now = st.session_state.get('latest_wechat_rendered_theme', 'spring-fresh')
                            raw_html = wechat_render.post_process_wechat_html(raw_html, theme_now)
                            final_html = wechat_publisher.replace_local_images_with_wechat_urls(raw_html, token)
                            cover_media_id = wechat_publisher.upload_cover_image(selected_cover_path, token)
                            draft_media_id = wechat_publisher.create_draft(
                                token, 
                                meta_title, 
                                meta_author, 
                                meta_digest, 
                                final_html, 
                                cover_media_id
                            )
                            draft_url = wechat_publisher.get_draft_url(token, draft_media_id)
                            
                            st.session_state.wechat_draft_media_id = draft_media_id
                            st.session_state.wechat_draft_url = draft_url
                            st.session_state.wechat_publish_result = None
                            st.session_state.wechat_preview_status = None
                            
                            # 将草稿状态持久化到文件，防止刷新丢失
                            filepath = st.session_state.get('latest_wechat_file_path')
                            if filepath:
                                wechat_publisher.save_draft_info(
                                    filepath, 
                                    draft_media_id, 
                                    draft_url, 
                                    status="draft", 
                                    appid=active_account["appid"], 
                                    secret=active_account["secret"]
                                )
                            st.success("🎉 微信草稿文章创建成功！")
                            st.rerun()
                        except Exception as e:
                            err_str = str(e)
                            m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                            if m_ip:
                                ip_addr = m_ip.group(1)
                                st.error(f"❌ 操作失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                st.warning(f"💡 **解决方法**：请登录「微信公众平台 -> 开发 -> 开发设置 -> IP白名单」，把 IP **`{ip_addr}`** 加进去，然后重新尝试。")
                            else:
                                st.error(f"❌ 操作失败: {err_str}")
                            
            if st.session_state.wechat_draft_url:
                st.markdown("#### 📱 微信扫码预览草稿")
                col_qr, col_info = st.columns([1.2, 2])
                with col_qr:
                    qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={urllib.parse.quote(st.session_state.wechat_draft_url)}"
                    st.image(qr_api_url, width=180, caption="扫码即可在手机微信预览草稿")
                with col_info:
                    st.markdown("**草稿箱 MediaID：**")
                    st.code(st.session_state.wechat_draft_media_id)
                    st.markdown("**临时预览链接 (在微信内打开有效)：**")
                    st.markdown(f"[👉 点击直接跳转预览]({st.session_state.wechat_draft_url})")
                    st.caption("注：微信临时预览链接有短期时效限制，过期后需重新点击「生成草稿箱文章」获取新链接。")
                
                with c_action2:
                    st.write("**手机预览推送**")
                    preview_wx_id = st.text_input("接收人的微信号", placeholder="如: my_wechat_id")
                    if st.button("2️⃣ 发送预览推送", use_container_width=True, disabled=not preview_wx_id):
                        with st.spinner("正在向指定微信号发送预览..."):
                            try:
                                token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                                wechat_publisher.send_preview(token, st.session_state.wechat_draft_media_id, preview_wx_id)
                                st.session_state.wechat_preview_status = f"✅ 预览成功发送至微信号: {preview_wx_id}，请检查手机微信通知！"
                                st.rerun()
                            except Exception as e:
                                err_str = str(e)
                                m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                                if m_ip:
                                    ip_addr = m_ip.group(1)
                                    st.error(f"❌ 发送失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                    st.warning(f"💡 **解决方法**：请登录「微信公众平台」，把 IP **`{ip_addr}`** 加到 IP 白名单中，然后重新尝试。")
                                else:
                                    st.error(f"❌ 发送失败: {err_str}")
                                
                if st.session_state.wechat_preview_status:
                    st.info(st.session_state.wechat_preview_status)
                    
                with c_action3:
                    st.write("**正式发布**")
                    st.warning("⚠️ 正式发布为不可逆的群发操作，将对所有关注者可见！")
                    
                    publish_action_type = st.radio(
                        "发布类型选择",
                        ["群发 (推送粉丝，主页可见，占额度)", 
                         "发布 (仅生成永久链接，不推送，主页历史看不到)"],
                        index=0,
                        disabled=is_published,
                        key="wechat_publish_action_type"
                    )
                    chosen_action = "mass_send" if "群发" in publish_action_type else "publish"
                    
                    enable_schedule = st.checkbox("定时发布（不勾选则立即执行）", key="wechat_enable_schedule", disabled=is_published)
                    
                    sched_dt = None
                    if enable_schedule:
                        col_d, col_t = st.columns(2)
                        with col_d:
                            import datetime as dt_module
                            sched_date = st.date_input("发布日期", min_value=dt_module.date.today(), key="wechat_sched_date")
                        with col_t:
                            sched_time = st.time_input("发布时间", key="wechat_sched_time")
                        
                        sched_dt = dt_module.datetime.combine(sched_date, sched_time)
                        
                    confirm_publish = st.checkbox("我已确认扫码预览无误，同意正式发布/群发", disabled=is_published)
                    
                    btn_label = "⏰ 安排定时发布" if enable_schedule else ("3️⃣ 正式群发到公众号" if chosen_action == "mass_send" else "3️⃣ 正式发布到公众号 (仅生成链接)")
                    
                    if st.button(btn_label, use_container_width=True, type="primary", disabled=(not confirm_publish) or is_published):
                        if enable_schedule and sched_dt:
                            import datetime as dt_module
                            now_dt = dt_module.datetime.now()
                            if sched_dt <= now_dt:
                                st.error("❌ 定时发布时间必须是未来的时间！")
                            else:
                                with st.spinner("正在安排定时发布..."):
                                    try:
                                        # 1. 保存状态为 scheduled 并记入 scheduled_time 与 publish_mode
                                        wechat_publisher.save_draft_info(
                                            filepath,
                                            st.session_state.wechat_draft_media_id,
                                            st.session_state.wechat_draft_url,
                                            status="scheduled",
                                            scheduled_time=sched_dt.strftime("%Y-%m-%d %H:%M:%S"),
                                            appid=active_account["appid"],
                                            secret=active_account["secret"],
                                            publish_mode=chosen_action
                                        )
                                        # 2. 注册定时任务到 BackgroundScheduler
                                        job_id = f"schedule_publish_{st.session_state.wechat_draft_media_id}"
                                        
                                        # 移除可能已存在的同 ID 任务
                                        try:
                                            st.session_state.scheduler.remove_job(job_id)
                                        except Exception:
                                            pass
                                            
                                        # 计算绝对路径
                                        draft_json_path = os.path.abspath(filepath).replace(".md", ".draft.json")
                                        st.session_state.scheduler.add_job(
                                            run_scheduled_wechat_publish,
                                            'date',
                                            run_date=sched_dt,
                                            args=[draft_json_path],
                                            id=job_id
                                        )
                                        st.session_state.wechat_publish_result = f"📅 定时任务已成功设置！安排在 `{sched_dt.strftime('%Y-%m-%d %H:%M:%S')}` 自动执行（模式：{'群发' if chosen_action == 'mass_send' else '仅发布'}）。请保持系统后台运行。"
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"❌ 设置定时发布失败: {str(e)}")
                        else:
                            # 立即群发 / 发布
                            with st.spinner("正在发送中，请稍候..."):
                                try:
                                    token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                                    if chosen_action == "mass_send":
                                        pub_id = wechat_publisher.mass_send_draft(token, st.session_state.wechat_draft_media_id)
                                        st.session_state.wechat_publish_result = f"🚀 文章已正式群发！消息 ID: `{pub_id}`。您可在公众平台后台查看群发状态。"
                                    else:
                                        pub_id = wechat_publisher.publish_draft(token, st.session_state.wechat_draft_media_id)
                                        st.session_state.wechat_publish_result = f"🚀 文章已正式发布！发布 ID: `{pub_id}`。您已通过「发布」发表，可使用链接分享。"
                                    
                                    # 更新为已发布状态
                                    if filepath:
                                        wechat_publisher.save_draft_info(
                                            filepath, 
                                            st.session_state.wechat_draft_media_id, 
                                            st.session_state.wechat_draft_url, 
                                            status="published",
                                            publish_mode=chosen_action
                                        )
                                    st.rerun()
                                except Exception as e:
                                    err_str = str(e)
                                    m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                                    if m_ip:
                                        ip_addr = m_ip.group(1)
                                        st.error(f"❌ 发布失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                        st.warning(f"💡 **解决方法**：请登录「微信公众平台」，把 IP **`{ip_addr}`** 加到 IP 白名单中，然后重新尝试。")
                                    else:
                                        st.error(f"❌ 发布失败: {err_str}")
                                
                if st.session_state.wechat_publish_result:
                    st.success(st.session_state.wechat_publish_result)

    # Pre-load current channel config so it is available globally
    current_chan_config = config["channel_configs"].get(config.get("platform", "自定义/OpenAI"), {})

    # Expander 1: 🔑 授权与星球管理
    with st.expander("🔑 授权与星球管理", expanded=False):
        if st.button("🔄 刷新授权状态", use_container_width=True):
            clear_auth_cache()
            st.rerun()
            
        st.markdown("**知识星球授权**")
        zsxq_stdout = get_zsxq_auth_status()
        logged_in = False
        user_id = ""
        user_name = ""
        try:
            m_auth = re.search(r'\{.*\}', zsxq_stdout, re.DOTALL)
            if m_auth:
                auth_data = json.loads(m_auth.group(0))
                if auth_data.get("ok") and auth_data.get("data", {}).get("loggedIn"):
                    logged_in = True
                    user_id = auth_data["data"].get("userId", "")
                    user_name = auth_data["data"].get("userName", "")
        except Exception: pass
        
        net_proxy = st.text_input("🌐 网络代理 (HTTP/HTTPS Proxy)", value=config.get("network_proxy", ""), placeholder="例如: http://127.0.0.1:7890", help="当服务器访问知识星球 TLS 超时或无法直连时，请填入可用的 HTTP/HTTPS 代理。")
        if net_proxy != config.get("network_proxy", ""):
            config["network_proxy"] = net_proxy
            save_config(config)
            clear_auth_cache()
            st.rerun()

        if logged_in:
            st.success(f"✅ {user_name} (已授权)")
            if st.button("退出登录"): 
                subprocess.run([NPX_CMD, "zsxq-cli", "auth", "logout"], env=get_exec_env())
                clear_auth_cache()
                st.rerun()
        else:
            st.warning("⚠️ 未授权")
            if st.button("🔗 获取授权链接"):
                login_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "login", "--json", "--no-browser", "--no-wait"], capture_output=True, text=True, encoding='utf-8', env=get_exec_env())
                try:
                    m_login = re.search(r'\{.*\}', login_res.stdout, re.DOTALL)
                    if m_login:
                        d_login = json.loads(m_login.group(0))["data"]
                        st.session_state.zsxq_device_code = d_login["device_code"]
                        st.markdown(f"**[👉 点击授权]({d_login['verification_uri_complete']})**")
                        st.info(f"确认码：`{d_login['user_code']}`")
                except Exception: st.error("无法获取链接")
            if "zsxq_device_code" in st.session_state and st.button("我已完成授权"):
                with st.spinner("验证中..."):
                    verify_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "login", "--device-code", st.session_state.zsxq_device_code, "--json"], capture_output=True, text=True, encoding='utf-8', env=get_exec_env())
                    try:
                        m_verify = re.search(r'\{.*\}', verify_res.stdout, re.DOTALL)
                        if m_verify and json.loads(m_verify.group(0)).get("ok"):
                            st.success("授权成功！")
                            del st.session_state.zsxq_device_code
                            clear_auth_cache()
                            st.rerun()
                        else:
                            st.error("未检测到成功授权，请确认您已在手机端扫码并输入了确认码！")
                    except Exception:
                        st.error("验证失败：无法解析登录状态。")
        
        st.markdown("---")
        st.markdown("**🔍 故障诊断工具**")
        if st.button("🌐 诊断服务器网络与授权", use_container_width=True):
            with st.spinner("正在进行多维度网络连通性与授权诊断..."):
                st.markdown("#### 1. DNS 解析测试")
                import socket
                try:
                    ip = socket.gethostbyname("mcp.zsxq.com")
                    st.success(f"✅ DNS 解析成功: `mcp.zsxq.com` -> `{ip}`")
                except Exception as e:
                    st.error(f"❌ DNS 解析失败: {str(e)}")
                
                st.markdown("#### 2. Curl 详细连接握手诊断 (TLS/WAF)")
                curl_cmd = "curl.exe" if sys.platform == "win32" else "curl"
                try:
                    res = subprocess.run([curl_cmd, "-iv", "-X", "POST", "https://mcp.zsxq.com/topic/mcp"], capture_output=True, text=True, timeout=15, env=get_exec_env())
                    st.text("执行命令: curl -iv -X POST https://mcp.zsxq.com/topic/mcp")
                    st.code(f"STDOUT:\n{res.stdout}\n\nSTDERR:\n{res.stderr}", language="bash")
                except Exception as e:
                    st.error(f"❌ Curl 执行异常: {str(e)}")
                    if "timed out" in str(e).lower() or "timeout" in str(e).lower():
                        st.warning("""💡 **TLS 握手超时 (`TLS handshake timeout`) 诊断解决指引：**
您的服务器在访问 `mcp.zsxq.com` 时遭遇 TLS 握手超时。这通常代表服务器 IP 被知识星球安全网关 (WAF) 防火墙阻断。
- **推荐方案**：请在上方【🌐 网络代理】配置项中填入有效的 HTTP/HTTPS 代理（如 `http://127.0.0.1:7890`）。
- **备选方案**：在 Debian 服务端拉起应用前，于终端中运行 `export HTTPS_PROXY=http://your-proxy:port`。""")
                    
                st.markdown("#### 3. CLI 授权状态诊断")
                try:
                    res_cli = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8', timeout=15, env=get_exec_env())
                    st.code(res_cli.stdout or res_cli.stderr, language="json")
                except Exception as e:
                    st.error(f"❌ CLI 诊断异常: {str(e)}")

        st.markdown("---")
        # Dynamic groups based on userId
        if logged_in and user_id:
            if "user_groups" not in config: config["user_groups"] = {}
            if user_id not in config["user_groups"]:
                config["user_groups"][user_id] = config.get("groups", {"默认群组": ""})
                save_config(config)
            groups_dict = config["user_groups"][user_id]
        else:
            groups_dict = {"默认群组": ""}
            
        group_keys = list(groups_dict.keys())
        if not group_keys:
            groups_dict = {"默认群组": ""}
            group_keys = ["默认群组"]
            
        if config.get("selected_group") not in group_keys:
            config["selected_group"] = group_keys[0]

        sel_g_name = st.selectbox("选择星球/群组", group_keys, index=group_keys.index(config["selected_group"]))
        if sel_g_name != config["selected_group"]:
            config["selected_group"] = sel_g_name
            save_config(config)
            
        curr_group_id = groups_dict[sel_g_name]
        
        if logged_in and user_id:
            st.markdown("---")
            st.markdown("**➕ 星群管理**")
            st.markdown("**新增或修改群组**")
            ng_name = st.text_input("星球名称", placeholder="例如：阿铭linux")
            ng_id = st.text_input("Group ID", placeholder="星球数字ID")
            if st.button("💾 保存/更新群组"):
                if ng_name and ng_id:
                    config["user_groups"][user_id][ng_name] = ng_id
                    config["selected_group"] = ng_name
                    save_config(config)
                    st.success("更新成功！")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("名称和ID不能为空！")
                    
            st.markdown("---")
            st.markdown("**删除群组**")
            del_g_name = st.selectbox("选择要删除的群组", group_keys)
            if st.button("🗑️ 删除选中群组"):
                if del_g_name in config["user_groups"][user_id]:
                    del config["user_groups"][user_id][del_g_name]
                    if config["selected_group"] == del_g_name:
                        config["selected_group"] = list(config["user_groups"][user_id].keys())[0] if config["user_groups"][user_id] else ""
                    save_config(config)
                    st.success("删除成功！")
                    time.sleep(0.5)
                    st.rerun()

    # Expander 2: 🎨 绘图引擎配置
    with st.expander("🎨 绘图引擎配置", expanded=False):
        img_gen_ops = ["即梦 (Dreamina)", "Google Gemini (Imagen 3)"]
        prev_img_gen = config.get("image_generator", "即梦 (Dreamina)")
        selected_img_gen = st.selectbox("图片生成引擎", img_gen_ops, index=img_gen_ops.index(prev_img_gen) if prev_img_gen in img_gen_ops else 0)
        
        if selected_img_gen != prev_img_gen:
            config["image_generator"] = selected_img_gen
            save_config(config)
            st.rerun()

        if selected_img_gen == "Google Gemini (Imagen 3)":
            google_api_key_val = config.get("google_api_key", "")
            new_google_key = st.text_input("Google API Key", value=google_api_key_val, type="password")
            
            model_ops = ["imagen-4.0-generate-001", "imagen-4.0-ultra-generate-001", "imagen-4.0-fast-generate-001", "gemini-3.1-flash-image"]
            prev_model = config.get("gemini_image_model", "imagen-4.0-generate-001")
            selected_model = st.selectbox("Gemini 绘图模型", model_ops, index=model_ops.index(prev_model) if prev_model in model_ops else 0)
            
            aspect_ratio_ops = ["1:1", "4:3", "16:9", "3:4", "9:16"]
            prev_aspect_ratio = config.get("image_aspect_ratio", "1:1")
            selected_ratio = st.selectbox("图片比例", aspect_ratio_ops, index=aspect_ratio_ops.index(prev_aspect_ratio) if prev_aspect_ratio in aspect_ratio_ops else 0)
            
            if new_google_key != google_api_key_val or selected_model != prev_model or selected_ratio != prev_aspect_ratio:
                config["google_api_key"] = new_google_key
                config["gemini_image_model"] = selected_model
                config["image_aspect_ratio"] = selected_ratio
                save_config(config)
                
            if new_google_key:
                st.success("✅ Google Gemini 已配置")
            else:
                st.warning("⚠️ Google Gemini 未配置 API Key")
        else:
            d_returncode, d_stdout = get_dreamina_credit_status()
            dreamina_logged_in = (d_returncode == 0)
            dreamina_credit_info = ""
            
            if dreamina_logged_in:
                try:
                    credit_data = json.loads(d_stdout)
                    dreamina_credit_info = f"剩余积分: {credit_data.get('total_credit', '未知')}"
                except Exception:
                    dreamina_credit_info = d_stdout.strip()
                st.success(f"✅ 即梦已授权 ({dreamina_credit_info})")
                if st.button("退出即梦登录"): 
                    subprocess.run([DREAMINA_CMD, "logout"])
                    clear_auth_cache()
                    st.rerun()
            else:
                st.warning("⚠️ 即梦未授权")
                if st.button("🔗 获取即梦授权链接"):
                    login_res = subprocess.run([DREAMINA_CMD, "login", "--headless"], capture_output=True, text=True, encoding='utf-8')
                    m_uri = re.search(r'verification_uri:\s*(https://\S+)', login_res.stdout)
                    m_code = re.search(r'user_code:\s*(\w+)', login_res.stdout)
                    m_device = re.search(r'device_code:\s*(\w+)', login_res.stdout)
                    if m_uri and m_code and m_device:
                        st.session_state.dreamina_device_code = m_device.group(1)
                        st.markdown(f"**[👉 点击授权]({m_uri.group(1)})**")
                        st.info(f"确认码：`{m_code.group(1)}`")
                    else:
                        st.error("无法获取链接")
                if "dreamina_device_code" in st.session_state and st.button("我已完成即梦授权"):
                    with st.spinner("验证即梦授权中（最多等待30秒）..."):
                        verify_res = subprocess.run(
                            [DREAMINA_CMD, "login", "checklogin",
                             f"--device_code={st.session_state.dreamina_device_code}",
                             "--poll=30"],
                            capture_output=True, text=True, encoding='utf-8'
                        )
                        combined_output = verify_res.stdout + verify_res.stderr
                        if verify_res.returncode == 0 or "LOGIN_SUCCESS" in combined_output or "登录成功" in combined_output or "Successfully" in combined_output:
                            st.success("授权成功！")
                            del st.session_state.dreamina_device_code
                            clear_auth_cache()
                            st.rerun()
                        else:
                            st.error("未检测到成功授权，请确认您已在网页端登录并确认！")

    # Expander 3: 🤖 AI 渠道配置
    with st.expander("🤖 AI 渠道配置", expanded=False):
        plat_ops = ["自定义/OpenAI", "火山方舟 (Volcengine)", "魔塔 (ModelScope)", "DeepSeek"]
        prev_platform = config.get("platform", "自定义/OpenAI")
        
        selected_platform = st.selectbox("AI 渠道", plat_ops, index=plat_ops.index(prev_platform) if prev_platform in plat_ops else 0)
        if selected_platform != prev_platform:
            config["platform"] = selected_platform
            save_config(config)
            st.rerun()
            
        current_chan_config = config["channel_configs"].get(selected_platform, {})
        
        new_api_key = st.text_input("API Key", value=current_chan_config.get("api_key", ""), type="password")
        new_base_url = st.text_input("Base URL", value=current_chan_config.get("base_url", ""))
        
        if new_api_key != current_chan_config.get("api_key") or new_base_url != current_chan_config.get("base_url"):
            config["channel_configs"][selected_platform]["api_key"] = new_api_key
            config["channel_configs"][selected_platform]["base_url"] = new_base_url
            save_config(config)
        
        if st.button("🔄 获取模型列表") and current_chan_config.get("api_key") and current_chan_config.get("base_url"):
            with st.spinner("获取中..."):
                try:
                    h = {"Authorization": f"Bearer {current_chan_config['api_key']}"}
                    r = httpx.get(f"{current_chan_config['base_url'].rstrip('/')}/models", headers=h, timeout=10)
                    if r.status_code == 200:
                        models = [m["id"] for m in r.json().get("data", [])]
                        config["channel_configs"][selected_platform]["available_models"] = models
                        save_config(config)
                        st.success(f"获取成功！共获取到 {len(models)} 个可用模型。")
                    else:
                        st.error(f"获取失败，HTTP 状态码: {r.status_code}")
                except Exception as e:
                    st.error(f"获取异常: {str(e)}")

        all_mods = list(dict.fromkeys(current_chan_config.get("available_models", []) + config.get("manual_models", [])))
        if not all_mods: all_mods = ["gpt-4o"]
        
        new_model = st.selectbox("当前模型", all_mods, index=all_mods.index(current_chan_config.get("selected_model")) if current_chan_config.get("selected_model") in all_mods else 0)
        if new_model != current_chan_config.get("selected_model"):
            config["channel_configs"][selected_platform]["selected_model"] = new_model
            save_config(config)

        if selected_platform == "DeepSeek":
            enable_thinking = st.toggle("开启思考模式 (Thinking Mode)", value=current_chan_config.get("enable_thinking", True))
            reasoning_effort = st.selectbox("思考强度 (Reasoning Effort)", ["high", "max"], index=0 if current_chan_config.get("reasoning_effort", "high") == "high" else 1)
            if enable_thinking != current_chan_config.get("enable_thinking", True) or reasoning_effort != current_chan_config.get("reasoning_effort", "high"):
                config["channel_configs"][selected_platform]["enable_thinking"] = enable_thinking
                config["channel_configs"][selected_platform]["reasoning_effort"] = reasoning_effort
                save_config(config)

    # Expander 4: ⏰ 定时任务 (Cron)
    # Expander 4: ⏰ 定时任务 (Cron)
    with st.expander("⏰ 定时任务 (Cron)", expanded=False):
        if "schedulers" not in config:
            config["schedulers"] = {}
            
        # Migrate old config if present
        if "auto_run" in config and "cron_expr" in config and not config["schedulers"]:
            config["schedulers"]["AI 深度分析"] = {
                "auto_run": config.get("auto_run", False),
                "cron_expr": config.get("cron_expr", "0 8 * * *"),
                "ui_state": {}
            }
            
        sched_config = config["schedulers"].setdefault(page_selection, {"auto_run": False, "cron_expr": "0 8 * * *", "ui_state": {}})
        
        curr_cron = sched_config.get("cron_expr", "0 8 * * *")
        new_cron = st.text_input("✏️ 手动输入 Cron 表达式", value=curr_cron, key=f"cron_input_{page_selection}")
        
        if st.button("✨ 打开 Cron 可视化配置器", use_container_width=True):
            cron_configurator_dialog(page_selection, new_cron)
            
        try:
            from apscheduler.triggers.cron import CronTrigger
            from datetime import datetime
            trigger = CronTrigger.from_crontab(new_cron)
            now = datetime.now()
            runs = []
            from datetime import timedelta
            for _ in range(5):
                now = trigger.get_next_fire_time(None, now)
                runs.append(now.strftime("%Y-%m-%d %H:%M:%S"))
                now = now + timedelta(seconds=1)
            st.info("🔮 **主界面实时预演 (未来 5 次):**\n\n" + "\n".join([f"- {r}" for r in runs]))
        except Exception:
            st.error("❌ 表达式无效，无法计算执行时间。")
            
        new_auto_run = st.toggle("开启定时自动运行", value=sched_config.get("auto_run", False), key=f"auto_run_{page_selection}")
        
        if new_auto_run:
            st.success("🟢 定时任务已激活并运行中")
        else:
            st.warning("⚪ 定时任务未开启")
            
        col_save, col_test, col_log = st.columns(3)
        with col_save:
            if st.button("💾 保存当前配置", use_container_width=True):
                # Capture UI state snapshot
                ui_state = {}
                if page_selection == "AI 深度分析":
                    ui_state = {
                        "scope_ui": st.session_state.get("da_scope_ui", "最新总结 (话题+文件)"),
                        "l_limit": st.session_state.get("da_l_limit", 3),
                        "a_mode": st.session_state.get("da_a_mode", "常规总结"),
                        "use_p_ui": st.session_state.get("da_use_p_ui", False),
                        "cp_text": st.session_state.get("da_cp_text", ""),
                        "use_wechat": st.session_state.get("da_use_wechat", False),
                        "wechat_article_orientation": st.session_state.get("da_wechat_article_orientation", config.get("wechat_article_orientation", "产业宏观与行业趋势")),
                        "also_generate_report": st.session_state.get("da_also_generate_report", False),
                        "wechat_mode": st.session_state.get("da_wechat_mode", "AI 模式 (免费)"),
                        "wechat_theme": st.session_state.get("da_wechat_theme", "spring-fresh"),
                        "wechat_custom_prompt": st.session_state.get("da_wechat_custom_prompt", ""),
                        "wechat_font_size": st.session_state.get("da_wechat_font_size", "medium"),
                        "wechat_background_type": st.session_state.get("da_wechat_background_type", "none"),
                        "wechat_prompt": st.session_state.get("da_wechat_prompt", ""),
                        "wechat_publish_mode": st.session_state.get("da_wechat_publish_mode", "仅生成本地文件 (不上传)"),
                        "wechat_publish_account": st.session_state.get("da_wechat_publish_account", ""),
                        "wechat_publish_author": st.session_state.get("da_wechat_publish_author", ""),
                        "selected_group": config.get("selected_group", "默认群组")
                    }
                elif page_selection == "视频脚本制作器":
                    v_hist_serial = []
                    for vf in st.session_state.get("virtual_history", []):
                        v_hist_serial.append({"name": vf["name"], "text": vf["text"]})
                    ui_state = {
                        "script_mode": st.session_state.get("vs_script_mode", "仿写现有格式生成新脚本"),
                        "export_format": st.session_state.get("vs_export_format", ".docx"),
                        "prompt_input": st.session_state.get("vs_prompt_input", ""),
                        "virtual_history": v_hist_serial
                    }
                elif page_selection == "指标文档制作":
                    ui_state = {
                        "selected_indicator": st.session_state.get("ind_selected", "")
                    }
                    
                config["schedulers"][page_selection] = {
                    "auto_run": new_auto_run,
                    "cron_expr": new_cron,
                    "ui_state": ui_state
                }
                save_config(config)
                scheduler.update_scheduler()
                st.success("配置与定时任务已保存！")
                time.sleep(1)
                st.rerun()
                
        with col_test:
            if st.button("⚡ 测试运行任务", use_container_width=True):
                import threading
                def test_runner():
                    if page_selection == "AI 深度分析":
                        scheduler.run_scheduled_deep_analysis(ignore_auto_run=True)
                    elif page_selection == "视频脚本制作器":
                        scheduler.run_scheduled_video_script(ignore_auto_run=True)
                    elif page_selection == "指标文档制作":
                        scheduler.run_scheduled_indicator_docs(ignore_auto_run=True)
                
                threading.Thread(target=test_runner, daemon=True).start()
                st.success("⚡ 测试任务已在后台启动！请稍后点击『查看执行日志』关注运行进展。")
                
        with col_log:
            if st.button("📋 查看执行日志", use_container_width=True):
                log_path = os.path.join("outputs", "cron_execution.log")
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8") as f:
                        log_content = f.read()
                else:
                    log_content = "暂无执行日志。"
                if len(log_content) > 30000:
                    log_content = "...(已截断历史日志)...\n" + log_content[-30000:]
                if hasattr(st, "dialog"):
                    show_log_dialog(log_content, page_selection)
                else:
                    st.session_state["show_log_fallback"] = log_content
                    st.rerun()

    if st.button("💾 手动保存所有配置", use_container_width=True): save_config(config); st.success("已保存")

# Main Area
if "show_log_fallback" in st.session_state:
    with st.container(border=True):
        st.subheader("📋 定时任务执行日志")
        st.text_area("", value=st.session_state["show_log_fallback"], height=450, disabled=True)
        if st.button("关闭日志"):
            del st.session_state["show_log_fallback"]
            st.rerun()

if page_selection == "AI 深度分析":
    st.title("🤖 AI 深度分析控制台")
    col1, col2 = st.columns([2, 1.3])
    with col1:
        st.subheader("📊 分析模式与数据源")
        scope_ui = st.radio("获取范围", ["最新总结 (话题+文件)", "文件总结 (仅限附件)"], horizontal=True, key="da_scope_ui")
        s_key = "all" if "最新" in scope_ui else "files"
        l_limit = st.number_input("获取近多少条消息", min_value=1, max_value=100, value=3, key="da_l_limit")
        a_mode = st.radio("分析模式", ["常规总结", "个股分析", "行业分析"], horizontal=True, key="da_a_mode")
        
        st.markdown("---")
        st.subheader("✨ 个性化输出处理")
        use_p_ui = st.checkbox("总结后进行个性化二次加工", key="da_use_p_ui")
        cp_text = ""
        if use_p_ui:
            cp_ops = ["自定义输入"] + list(config["custom_prompts"].keys())
            sp_name = st.selectbox("选择预设", cp_ops)
            
            if sp_name == "自定义输入":
                cp_text = st.text_area("指令", key="da_cp_text")
                np_name = st.text_input("预设名称")
                if st.button("💾 保存预设", use_container_width=True) and cp_text and np_name: 
                    config["custom_prompts"][np_name] = cp_text
                    save_config(config)
                    st.rerun()
            else:
                cp_text = st.text_area("编辑指令", value=config["custom_prompts"][sp_name], height=120, key="da_cp_text")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("🔄 更新当前预设", use_container_width=True):
                        config["custom_prompts"][sp_name] = cp_text
                        save_config(config)
                        st.success("更新成功！")
                        time.sleep(1)
                        st.rerun()
                with col_btn2:
                    if st.button("🗑️ 删除当前预设", use_container_width=True):
                        st.session_state['delete_confirm'] = sp_name
                
                if st.session_state.get('delete_confirm') == sp_name:
                    with st.container(border=True):
                        st.warning(f"⚠️ 确定要删除预设【{sp_name}】吗？")
                        del_col1, del_col2 = st.columns(2)
                        if del_col1.button("✅ 确认删除", type="primary", use_container_width=True):
                            del config["custom_prompts"][sp_name]
                            save_config(config)
                            del st.session_state['delete_confirm']
                            st.rerun()
                        if del_col2.button("❌ 取消", use_container_width=True):
                            del st.session_state['delete_confirm']
                            st.rerun()

        st.markdown("---")
        st.subheader("📱 微信公众号推文生成")
        use_wechat = st.checkbox("生成微信公众号推文 (跳过常规图表渲染加速)", value=config.get("use_wechat", False), key="da_use_wechat")
        if use_wechat != config.get("use_wechat", False):
            config["use_wechat"] = use_wechat
            save_config(config)
            st.rerun()

        also_generate_report = False
        if use_wechat:
            wechat_orientations = ["产业宏观与行业趋势", "上市公司财报深度解读"]
            prev_orientation = config.get("wechat_article_orientation", "产业宏观与行业趋势")
            if prev_orientation not in wechat_orientations:
                prev_orientation = wechat_orientations[0]
            selected_orientation = st.radio(
                "🧭 推文内容定位 (文风与合规模式)",
                wechat_orientations,
                index=wechat_orientations.index(prev_orientation),
                horizontal=True,
                key="da_wechat_article_orientation",
                help="• 产业宏观与行业趋势：严格金融合规，严禁任何未来具体数字预测与个股点名，专注产业链供需与技术格局。\n• 上市公司财报深度解读：专注上市公司财报拆解、利润引擎归因、财务质量与经营风险，严禁买卖指令。"
            )
            if selected_orientation != prev_orientation:
                config["wechat_article_orientation"] = selected_orientation
                save_config(config)

            also_generate_report = st.checkbox("同时生成分析报告", value=config.get("also_generate_report", False), key="da_also_generate_report")
            if also_generate_report != config.get("also_generate_report", False):
                config["also_generate_report"] = also_generate_report
                save_config(config)
                st.rerun()
                
            # 定时自动发布模式
            wechat_pub_modes = [
                "仅生成本地文件 (不上传)", 
                "自动保存至微信草稿箱", 
                "自动保存草稿并‘发布’ (不推送，不占群发额度，主页历史看不到)", 
                "自动保存草稿并‘群发’ (推送给所有粉丝，占用群发额度，主页可见，超限自动降级)"
            ]
            prev_pub_mode = config.get("wechat_publish_mode", "仅生成本地文件 (不上传)")
            # 兼容老配置选项
            if prev_pub_mode == "自动保存草稿并正式发布":
                prev_pub_mode = "自动保存草稿并‘发布’ (不推送，不占群发额度，主页历史看不到)"
                
            selected_pub_mode = st.selectbox(
                "定时任务微信发布模式",
                wechat_pub_modes,
                index=wechat_pub_modes.index(prev_pub_mode) if prev_pub_mode in wechat_pub_modes else 0,
                key="da_wechat_publish_mode"
            )
            
            # 微信发布目标账号
            accounts = wechat_publisher.load_accounts()
            account_names = [a["name"] for a in accounts]
            prev_pub_acc = config.get("wechat_publish_account", account_names[0] if account_names else "")
            
            selected_pub_acc = st.selectbox(
                "微信发布目标账号",
                account_names if account_names else ["无可用账号"],
                index=account_names.index(prev_pub_acc) if (account_names and prev_pub_acc in account_names) else 0,
                key="da_wechat_publish_account"
            )
            
            # 微信发布默认作者
            prev_pub_author = config.get("wechat_publish_author", "")
            selected_pub_author = st.text_input(
                "定时任务/发布默认作者 (不填则不显示)",
                value=prev_pub_author,
                key="da_wechat_publish_author"
            )
            
            if (selected_pub_mode != prev_pub_mode or 
                selected_pub_acc != prev_pub_acc or 
                selected_pub_author != prev_pub_author):
                config["wechat_publish_mode"] = selected_pub_mode
                config["wechat_publish_account"] = selected_pub_acc
                config["wechat_publish_author"] = selected_pub_author
                save_config(config)

        wechat_prompt = ""
        if use_wechat:
            col_mode, col_theme = st.columns(2)
            with col_mode:
                wechat_mode_ops = ["AI 模式 (免费)", "API 模式 (专业)"]
                prev_wechat_mode = config.get("wechat_mode", "AI 模式 (免费)")
                selected_wechat_mode = st.selectbox(
                    "排版模式", 
                    wechat_mode_ops, 
                    index=wechat_mode_ops.index(prev_wechat_mode) if prev_wechat_mode in wechat_mode_ops else 0,
                    key="da_wechat_mode"
                )
            
            is_api_mode = (selected_wechat_mode == "API 模式 (专业)")
            if is_api_mode:
                theme_ops = [
                    "default", "bytedance", "apple", "sspai-red", "wechat-native", 
                    "nyt-classic", "sunset-amber", "mint-fresh", "lavender-dream",
                    "elegant-gold", "elegant-green", "elegant-blue", "elegant-red",
                    "focus-gold", "focus-green", "focus-blue", "focus-red",
                    "minimal-gold", "minimal-green", "minimal-blue", "minimal-red",
                    "bold-gold", "bold-green", "bold-blue", "bold-red",
                    "chinese", "cyber", "sports"
                ]
            else:
                theme_ops = ["spring-fresh", "autumn-warm", "ocean-calm", "custom"]
                
            with col_theme:
                prev_wechat_theme = config.get("wechat_theme", "spring-fresh" if not is_api_mode else "default")
                if prev_wechat_theme not in theme_ops:
                    prev_wechat_theme = theme_ops[0]
                selected_wechat_theme = st.selectbox(
                    "排版主题",
                    theme_ops,
                    index=theme_ops.index(prev_wechat_theme),
                    key="da_wechat_theme"
                )

            md2wechat_api_key = config.get("md2wechat_api_key", "")
            if is_api_mode:
                md2wechat_api_key = st.text_input("md2wechat API Key", value=md2wechat_api_key, type="password", placeholder="填入 md2wechat.cn 专属 API Key")
            
            wechat_custom_prompt = config.get("wechat_custom_prompt", "")
            if not is_api_mode and selected_wechat_theme == "custom":
                wechat_custom_prompt = st.text_area("自定义 AI 排版 Prompt", value=wechat_custom_prompt, placeholder="输入您自定义的 CSS/HTML 设计提示词...", key="da_wechat_custom_prompt")

            with st.expander("⚙️ 微信高级排版选项"):
                prev_font_size = config.get("wechat_font_size", "medium")
                font_size_ops = ["medium", "small", "large"]
                selected_font_size = st.selectbox("正文字号", font_size_ops, index=font_size_ops.index(prev_font_size), key="da_wechat_font_size")
                
                prev_bg_type = config.get("wechat_background_type", "none")
                bg_type_ops = ["none", "default", "grid"]
                selected_bg_type = st.selectbox("背景类型", bg_type_ops, index=bg_type_ops.index(prev_bg_type), key="da_wechat_background_type")

            if (selected_wechat_mode != prev_wechat_mode or 
                selected_wechat_theme != prev_wechat_theme or 
                md2wechat_api_key != config.get("md2wechat_api_key", "") or
                wechat_custom_prompt != config.get("wechat_custom_prompt", "") or
                selected_font_size != prev_font_size or
                selected_bg_type != prev_bg_type):
                
                config["wechat_mode"] = selected_wechat_mode
                config["wechat_theme"] = selected_wechat_theme
                config["md2wechat_api_key"] = md2wechat_api_key
                config["wechat_custom_prompt"] = wechat_custom_prompt
                config["wechat_font_size"] = selected_font_size
                config["wechat_background_type"] = selected_bg_type
                save_config(config)
            
            wechat_prompt = st.text_area("推文个性化要求 (可选)", placeholder="例如：语气更加诙谐幽默，重点强调端侧存储，多用短句...", height=68, key="da_wechat_prompt")

        task_state = scheduler.get_manual_task_state()
        is_running = task_state.get("status") == "running"

        if is_running:
            st.warning(f"⏳ 后台深度分析任务正在解耦运行中... (启动时间: {task_state.get('started_at', '')}，刷新 F5 或切换导航任务不会中断)")
            if st.button("🛑 停止当前任务", use_container_width=True, type="primary"):
                scheduler.request_stop_manual_task()
                scheduler.write_manual_log("🛑 用户在界面点击了停止当前任务按钮...")
                st.toast("已发送停止信号！", icon="🛑")
                time.sleep(1)
                st.rerun()
        else:
            if st.button("🚀 立即开始深度分析", use_container_width=True, type="primary"):
                if not current_chan_config.get("api_key"):
                    st.error("请先在左侧全局配置中填写 API Key！")
                    st.stop()
                    
                # 启动后台工作线程
                import threading
                worker_thread = threading.Thread(
                    target=scheduler.run_manual_deep_analysis_worker,
                    args=(
                        curr_group_id, l_limit, s_key, current_chan_config, a_mode, 
                        use_p_ui, cp_text, use_wechat, also_generate_report, 
                        wechat_prompt, config
                    ),
                    daemon=True
                )
                worker_thread.start()
                st.toast("🚀 后台深度分析任务已启动！", icon="🚀")
                time.sleep(0.5)
                st.rerun()

        # 无论是否在运行，均常驻展示 outputs/manual_execution.log 的最新运行日志
        st.write("---")
        with st.expander("📋 深度分析任务执行日志 (已持久化保存)", expanded=is_running):
            if os.path.exists(MANUAL_LOG_FILE):
                with open(MANUAL_LOG_FILE, "r", encoding="utf-8") as f:
                    log_text = f.read()
                if log_text.strip():
                    st.text_area("日志输出窗口", value=log_text, height=320, key="manual_log_display_window")
                else:
                    st.info("暂无日志输出")
            else:
                st.info("暂无日志记录")

        # 若任务在后台运行中，自动隔 2 秒轮询刷新 UI 呈现最新日志
        if is_running:
            time.sleep(2)
            st.rerun()

    def auto_load_latest_outputs_to_session(force=False):
        """从 outputs/ 磁盘目录自动装载最新的推文 Markdown、HTML 和分析图表到 st.session_state，解决刷新后预览丢失问题"""
        # 1. 恢复最新微信推文
        if force or 'latest_wechat' not in st.session_state or not st.session_state.get('latest_wechat'):
            hist_mds = sorted(glob.glob(os.path.join(WECHAT_OUTPUT_DIR, "*.md")), reverse=True)
            if hist_mds:
                latest_md_file = hist_mds[0]
                try:
                    with open(latest_md_file, "r", encoding="utf-8") as f:
                        st.session_state['latest_wechat'] = f.read()
                    st.session_state['latest_wechat_file_path'] = latest_md_file
                    
                    latest_html_file = latest_md_file.replace(".md", ".html")
                    if os.path.exists(latest_html_file):
                        with open(latest_html_file, "r", encoding="utf-8") as hf:
                            raw_html = hf.read()
                            theme_now = config.get("wechat_theme", "spring-fresh")
                            mode_now = config.get("wechat_mode", "AI 模式 (免费)")
                            raw_html = wechat_render.post_process_wechat_html(raw_html, theme_now)
                            st.session_state['latest_wechat_html_raw'] = raw_html
                            st.session_state['latest_wechat_html'] = wechat_render.make_preview_html(raw_html)
                            st.session_state['latest_wechat_rendered_theme'] = theme_now
                            st.session_state['latest_wechat_rendered_mode'] = mode_now
                            st.session_state['preview_wechat_theme_select'] = theme_now
                            st.session_state['preview_wechat_mode_select'] = mode_now
                except Exception as e:
                    pass

        # 2. 恢复最新分析报表图片
        if force or 'latest_img' not in st.session_state or not st.session_state.get('latest_img'):
            hist_imgs = sorted(glob.glob(os.path.join(IMAGE_OUTPUT_DIR, "*.png")), reverse=True)
            if hist_imgs:
                st.session_state['latest_img'] = hist_imgs[0]

    # 自动恢复磁盘最新成果，防止页面刷新或切换后右侧预览丢失
    if not is_running and task_state.get("status") == "success":
        auto_load_latest_outputs_to_session(force=True)
    else:
        auto_load_latest_outputs_to_session(force=False)

    with col2:
        tab1, tab2 = st.tabs(["📊 分析报告预览", "📱 公众号推文预览"])
        
        with tab1:
            if 'latest_img' in st.session_state and os.path.exists(st.session_state['latest_img']):
                st.image(st.session_state['latest_img'], use_container_width=True)
            
            st.write("---")
            st.subheader("📅 历史报告归档")
            hist_imgs = sorted(glob.glob(os.path.join(IMAGE_OUTPUT_DIR, "*.png")), reverse=True)
            if not hist_imgs: st.info("暂无历史记录")
            else:
                d_groups = {}
                for h_im in hist_imgs:
                    try:
                        ds = os.path.basename(h_im).split('_')[1][:8]
                        dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                    except Exception: dfmt = "其他"
                    if dfmt not in d_groups: d_groups[dfmt] = []
                    d_groups[dfmt].append(h_im)
                
                for d_key, d_ims in d_groups.items():
                    with st.expander(f"📅 {d_key} ({len(d_ims)}份)"):
                        for single_im in d_ims:
                            st.image(single_im, caption=os.path.basename(single_im))
                            
        with tab2:
            if st.session_state.get('latest_wechat'):
                # 初始化 selectbox 状态以防止切换两次的冲突问题
                if 'preview_wechat_mode_select' not in st.session_state:
                    st.session_state['preview_wechat_mode_select'] = st.session_state.get('latest_wechat_rendered_mode', config.get("wechat_mode", "AI 模式 (免费)"))
                if 'preview_wechat_theme_select' not in st.session_state:
                    st.session_state['preview_wechat_theme_select'] = st.session_state.get('latest_wechat_rendered_theme', config.get("wechat_theme", "spring-fresh"))
                
                # 🎨 实时排版主题/模式切换面板
                st.markdown("### 🎨 实时排版主题切换")
                col_preview_mode, col_preview_theme = st.columns(2)
                with col_preview_mode:
                    preview_mode_ops = ["AI 模式 (免费)", "API 模式 (专业)"]
                    val_mode = st.session_state['preview_wechat_mode_select']
                    if val_mode not in preview_mode_ops:
                        val_mode = preview_mode_ops[0]
                    preview_mode = st.selectbox(
                        "预览排版模式", 
                        preview_mode_ops, 
                        key="preview_wechat_mode_select",
                        index=preview_mode_ops.index(val_mode)
                    )
                
                is_preview_api = "API" in preview_mode
                if is_preview_api:
                    preview_theme_ops = [
                        "default", "bytedance", "apple", "sspai-red", "wechat-native", 
                        "nyt-classic", "sunset-amber", "mint-fresh", "lavender-dream",
                        "elegant-gold", "elegant-green", "elegant-blue", "elegant-red",
                        "focus-gold", "focus-green", "focus-blue", "focus-red",
                        "minimal-gold", "minimal-green", "minimal-blue", "minimal-red",
                        "bold-gold", "bold-green", "bold-blue", "bold-red",
                        "chinese", "cyber", "sports"
                    ]
                else:
                    preview_theme_ops = ["spring-fresh", "autumn-warm", "ocean-calm", "custom"]
                    
                with col_preview_theme:
                    val_theme = st.session_state['preview_wechat_theme_select']
                    if val_theme not in preview_theme_ops:
                        val_theme = preview_theme_ops[0]
                    preview_theme = st.selectbox(
                        "预览排版主题",
                        preview_theme_ops,
                        key="preview_wechat_theme_select",
                        index=preview_theme_ops.index(val_theme)
                    )
                
                # 响应式渲染检查
                if st.session_state.get('just_loaded_history'):
                    st.session_state['latest_wechat_rendered_theme'] = preview_theme
                    st.session_state['latest_wechat_rendered_mode'] = preview_mode
                    st.session_state['just_loaded_history'] = False
                
                need_re_render = False
                if not st.session_state.get('latest_wechat_html'):
                    need_re_render = True
                elif st.session_state.get('latest_wechat_rendered_theme') != preview_theme:
                    need_re_render = True
                elif st.session_state.get('latest_wechat_rendered_mode') != preview_mode:
                    need_re_render = True
                    
                if need_re_render:
                    with st.spinner("🔄 正在为当前推文切换主题并重新排版..."):
                        font_size = config.get("wechat_font_size", "medium")
                        bg_type = config.get("wechat_background_type", "none")
                        api_key = config.get("md2wechat_api_key", "")
                        custom_prompt = config.get("wechat_custom_prompt", "")
                        
                        html_res = wechat_render.convert_to_wechat_html(
                            st.session_state['latest_wechat'], 
                            preview_theme, 
                            preview_mode, 
                            api_key=api_key, 
                            font_size=font_size, 
                            bg_type=bg_type, 
                            chan_config=current_chan_config, 
                            custom_prompt=custom_prompt
                        )
                        st.session_state['latest_wechat_html_raw'] = html_res
                        st.session_state['latest_wechat_html'] = wechat_render.make_preview_html(html_res)
                        st.session_state['latest_wechat_rendered_theme'] = preview_theme
                        st.session_state['latest_wechat_rendered_mode'] = preview_mode
                        
                        if 'wechat_long_image' in st.session_state:
                            del st.session_state['wechat_long_image']
                            
                        file_path = st.session_state.get('latest_wechat_file_path')
                        if file_path:
                            html_p = file_path.replace(".md", ".html")
                            try:
                                with open(html_p, "w", encoding="utf-8") as hf:
                                    hf.write(html_res)
                            except Exception:
                                pass
                        
                        st.rerun()
                
                st.write("---")
                col_btn_docx, col_btn_img, col_btn_wechat = st.columns(3)
                with col_btn_docx:

                    st.download_button(
                        label="📄 导出推文为 Docx 文档",
                        data=wechat_render.markdown_to_wechat_docx_bytes(st.session_state['latest_wechat']),
                        file_name=f"微信公众号推文_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary"
                    )
                with col_btn_img:
                    if st.button("📸 导出推文为长图", use_container_width=True):
                        if st.session_state.get('latest_wechat_html'):
                            with st.spinner("正在启动 Playwright 渲染生成高保真长图..."):
                                try:
                                    img_path = image_engine.generate_wechat_long_image(st.session_state['latest_wechat_html'])
                                    st.session_state['wechat_long_image'] = img_path
                                    st.success("长图生成成功！请在下方查看并下载。")
                                except Exception as e:
                                    st.error(f"长图生成失败: {str(e)}")
                        else:
                            st.warning("暂无已美化的推文 HTML 样式，请重新运行分析。")
                with col_btn_wechat:
                    if st.button("📲 公众号一键发布/预览", use_container_width=True):
                        if st.session_state.get('latest_wechat_html'):
                            show_wechat_publish_dialog()
                        else:
                            st.warning("暂无已美化的推文 HTML 样式，请重新运行分析。")
                
                if st.session_state.get('wechat_long_image') and os.path.exists(st.session_state['wechat_long_image']):
                    st.write("---")
                    st.subheader("🖼️ 已生成的推文长图")
                    st.image(st.session_state['wechat_long_image'], use_container_width=True)
                    with open(st.session_state['wechat_long_image'], "rb") as img_file:
                        st.download_button(
                            label="📥 下载超长排版图",
                            data=img_file.read(),
                            file_name=os.path.basename(st.session_state['wechat_long_image']),
                            mime="image/png",
                            use_container_width=True
                        )
                
                st.write("---")
                st.subheader("📱 排版预览 (微信内置视口)")
                
                if st.session_state.get('latest_wechat_html'):
                    import streamlit.components.v1 as components
                    components.html(st.session_state['latest_wechat_html'], height=800, scrolling=True)
                else:
                    st.markdown(
                        """
                        <style>
                        .wechat-container {
                            max-width: 480px;
                            margin: 16px auto;
                            background-color: white;
                            border-radius: 12px;
                            padding: 24px 16px;
                            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
                            color: #333;
                            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                        }
                        .wechat-container img {
                            width: 100%;
                            border-radius: 8px;
                            margin: 16px 0;
                        }
                        .wechat-container h1, .wechat-container h2, .wechat-container h3 {
                            color: #1a1a1a;
                        }
                        .wechat-container blockquote {
                            border-left: 4px solid #07c160;
                            background: #f7f7f7;
                            margin: 16px 0;
                            padding: 12px;
                            color: #666;
                            font-size: 0.9em;
                        }
                        </style>
                        """, unsafe_allow_html=True
                    )
                    st.markdown('<div class="wechat-container">', unsafe_allow_html=True)
                    render_wechat_preview(st.session_state['latest_wechat'])
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("暂无生成的公众号推文。请在左侧勾选「生成微信公众号推文 (跳过常规图表渲染加速)」并开始深度分析。")
            
            st.write("---")
            st.subheader("📅 历史推文归档")
            hist_mds = sorted(glob.glob(os.path.join(WECHAT_OUTPUT_DIR, "*.md")), reverse=True)
            if not hist_mds: st.info("暂无历史记录")
            else:
                d_groups = {}
                for h_md in hist_mds:
                    try:
                        ds = os.path.basename(h_md).split('_')[1][:8]
                        dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                    except Exception: dfmt = "其他"
                    if dfmt not in d_groups: d_groups[dfmt] = []
                    d_groups[dfmt].append(h_md)
                
                for d_key, d_mds in d_groups.items():
                    with st.expander(f"📅 {d_key} ({len(d_mds)}篇)"):
                        for single_md in d_mds:
                            col_a, col_b = st.columns([3, 1])
                            try:
                                with open(single_md, "r", encoding="utf-8") as f:
                                    content = f.read()
                                    title = content.split('\n')[0].replace('#', '').strip()[:20]
                                    if not title: title = "无标题"
                            except Exception: title = "读取失败"
                            
                            col_a.write(f"📄 {title}...")
                            if col_b.button("查看", key=f"view_{os.path.basename(single_md)}"):
                                st.session_state['latest_wechat'] = content
                                st.session_state['latest_wechat_file_path'] = single_md
                                if 'wechat_long_image' in st.session_state:
                                    del st.session_state['wechat_long_image']
                                # 加载历史文章时重置内存中的微信草稿状态，弹窗会自动重新读取对应文件的 .draft.json
                                st.session_state.wechat_draft_media_id = None
                                st.session_state.wechat_draft_url = None
                                st.session_state.wechat_publish_result = None
                                st.session_state.wechat_preview_status = None
                                    
                                html_p = single_md.replace(".md", ".html")
                                if os.path.exists(html_p):
                                    with open(html_p, "r", encoding="utf-8") as hf:
                                        raw_html = hf.read()
                                        theme_now = st.session_state.get('preview_wechat_theme_select', config.get("wechat_theme", "spring-fresh"))
                                        mode_now = st.session_state.get('preview_wechat_mode_select', config.get("wechat_mode", "AI 模式 (免费)"))
                                        raw_html = wechat_render.post_process_wechat_html(raw_html, theme_now)
                                        st.session_state['latest_wechat_html_raw'] = raw_html
                                        st.session_state['latest_wechat_html'] = wechat_render.make_preview_html(raw_html)
                                        st.session_state['latest_wechat_rendered_theme'] = theme_now
                                        st.session_state['latest_wechat_rendered_mode'] = mode_now
                                    st.session_state['just_loaded_history'] = True
                                else:
                                    st.session_state['latest_wechat_html'] = None
                                    st.session_state['just_loaded_history'] = False
                                    
                                st.rerun()

elif page_selection == "视频脚本制作器":
    st.title("🎥 视频脚本制作器")
    
    col1, col2 = st.columns([2, 1.3])
    with col1:
        st.subheader("📁 历史脚本库 (供AI学习风格/续写)")
        uploaded_files = st.file_uploader("支持多文件上传 (.docx, .md, .xlsx, .csv)", type=['docx', 'md', 'xlsx', 'csv'], accept_multiple_files=True)
        
        if st.session_state.virtual_history:
            st.markdown("**📌 已追加的虚拟历史脚本:**")
            for vf in st.session_state.virtual_history:
                st.write(f"📄 `{vf['name']}`")
            if st.button("🗑️ 清空虚拟历史"):
                st.session_state.virtual_history = []
                st.rerun()
        
        st.markdown("---")
        st.subheader("⚙️ 生成设置")
        script_mode = st.radio("生成模式", ["仿写现有格式生成新脚本", "根据历史序列向后发散续写（例如基于0,1,2,3,4续写5,6）"], horizontal=False, key="vs_script_mode")
        export_format = st.selectbox("导出格式", [".docx", ".md"], key="vs_export_format")
        
        st.markdown("---")
        st.subheader("📝 素材与提示词")
        prompt_input = st.text_area("输入新的核心观点、素材内容或具体要求...", height=150, placeholder="例如：今天我们来讲一下人工智能在医疗领域的最新应用，重点突出AI辅助诊断的高效性...", key="vs_prompt_input")
        
        if st.button("🚀 开始生成脚本", use_container_width=True, type="primary"):
            if not current_chan_config.get("api_key"):
                st.error("请先在左侧全局配置中填写 API Key！")
                st.stop()
                
            with st.status("正在启动脚本制作流程...", expanded=True) as status:
                st.write("🔍 分析历史脚本与素材...")
                
                history_texts = []
                for f in uploaded_files:
                    txt = llm.parse_uploaded_file(f)
                    if txt: history_texts.append(f"【历史脚本：{f.name}】\n{txt}")
                for vf in st.session_state.virtual_history:
                    history_texts.append(f"【历史脚本（追加）：{vf['name']}】\n{vf['text']}")
                history_context = "\n\n".join(history_texts)
                
                system_prompt = "你是一个专业的金融/交易类视频脚本编导。请严格学习并模仿用户提供的历史脚本的文案风格、语气（如口语化、设问式）和排版格式。\n\n【重要排版指令】：请必须使用 Markdown 语法进行排版输出。为了作为提词器使用时的重音提示，请务必对文案中的核心观点、金句或转折词使用**加粗**（如 `**重点内容**`）或引用块（如 `> 核心金句`）进行高亮。"
                
                user_content = ""
                if history_context:
                    user_content += f"以下是你需要学习参考的历史脚本序列：\n\n{history_context}\n\n====================\n\n"
                
                if "仿写" in script_mode:
                    user_content += f"请根据以上提供的历史脚本风格进行学习。你的核心任务是**直接生成一篇全新的完整视频脚本（最终提词器念稿版本）**，绝对不要只输出分析总结，也不要向我提问索要素材。\n"
                    if prompt_input.strip():
                        user_content += f"\n[新素材与要求]：\n{prompt_input}\n\n请结合上述新素材生成脚本内容。"
                    else:
                        user_content += "\n由于用户没有提供新素材，请自行发挥你的专业金融编导水平，拟定一个符合当前市场热点的主题，直接撰写出这篇完整的视频脚本。"
                else:
                    user_content += f"请深度分析前面提供的历史脚本序列的故事线、知识递进逻辑和表达手法，自动推理出下一期的主题，并**直接【顺延生成】最新一期的完整视频脚本内容（最终提词器念稿版本）**。要求保持一贯的口语化、设问式风格。绝对不要向我提问索要素材，也绝对不要仅仅输出风格特征分析表！\n"
                    if prompt_input.strip():
                        user_content += f"\n在续写时，请必须结合以下新素材或要求：\n[新素材与要求]：\n{prompt_input}"
                    else:
                        user_content += "\n请注意：由于我没有提供新素材，请直接根据历史上下文的逻辑推演，自动拟定下一期主题并生成完整的视频文案！"
                
                st.write(f"🧠 AI 正在构思与生成 (模型: {current_chan_config.get('selected_model')})...")
                
                try:
                    client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                    response = llm.call_chat_completion(
                        client,
                        current_chan_config.get("selected_model"),
                        [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content}
                        ],
                        chan_config=current_chan_config
                    )
                    if not response.choices:
                        raise Exception(f"接口返回空数据，可能是该模型暂不支持或网络限流。详情: {response}")
                    script_content = response.choices[0].message.content
                    st.session_state['generated_script_preview'] = script_content
                    st.session_state['export_format_choice'] = export_format
                    
                    # 保存到历史归档
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    save_path = os.path.join(SCRIPT_OUTPUT_DIR, f"script_{timestamp}.md")
                    with open(save_path, "w", encoding="utf-8") as f:
                        f.write(script_content)
                        
                    status.update(label="脚本生成完成！", state="complete", expanded=False)
                    st.success("生成成功！请在右侧预览并下载。")
                except Exception as e:
                    status.update(label="任务异常终止", state="error")
                    st.error(f"脚本生成失败: {str(e)}")

    with col2:
        st.subheader("📺 生成结果预览")
        
        preview_text = st.session_state.get('generated_script_preview', '')
        if preview_text:
            styled_html = f"""
            <style>
                .script-preview-container {{
                    font-size: 16px;
                    line-height: 1.8;
                    color: #333;
                }}
                .script-preview-container strong {{
                    color: #d97706; /* 深橙色/主题色 */
                    font-weight: 900;
                    background-color: #fef3c7;
                    padding: 0 4px;
                    border-radius: 3px;
                }}
                .script-preview-container blockquote {{
                    border-left: 5px solid #07c160; /* 主题绿 */
                    padding: 12px 15px;
                    margin: 15px 0;
                    color: #4b5563;
                    background-color: #f9fafb;
                    border-radius: 0 8px 8px 0;
                    font-style: italic;
                }}
            </style>
            """
            st.markdown(styled_html, unsafe_allow_html=True)
            
            with st.container(border=True):
                st.markdown(f'<div class="script-preview-container">{markdown.markdown(preview_text, extensions=["extra", "nl2br"])}</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("💾 导出与下载")
            
            # Real download logic for generated scripts
            ext = st.session_state.get('export_format_choice', '.md')
            file_name = f"视频脚本_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
            
            if ext == ".docx":
                from io import BytesIO
                from docx import Document
                from docx.shared import RGBColor
                from docx.enum.text import WD_COLOR_INDEX
                import re
                
                doc = Document()
                for line in preview_text.split('\n'):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                        continue
                    
                    if line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                        continue
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                        continue
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                        continue
                    
                    is_quote = line.startswith('> ')
                    is_bullet = line.startswith('- ') or line.startswith('* ')
                    
                    if is_bullet:
                        p = doc.add_paragraph(style='List Bullet')
                        text = line[2:]
                    else:
                        p = doc.add_paragraph()
                        text = line[2:] if is_quote else line
                        
                    # 解析并映射加粗与颜色
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.color.rgb = RGBColor(217, 119, 6) # 对应 CSS 的深橙色 #d97706
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW # 黄色背景高亮
                        elif part:
                            run = p.add_run(part)
                            if is_quote:
                                run.font.color.rgb = RGBColor(7, 193, 96) # 对应 CSS 的主题绿 #07c160
                                run.italic = True
                                
                bio = BytesIO()
                doc.save(bio)
                file_data = bio.getvalue()
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                file_data = preview_text.encode('utf-8')
                mime_type = "text/plain"
            
            st.download_button(
                label=f"⬇️ 下载 {ext} 文件",
                data=file_data,
                file_name=file_name,
                mime=mime_type,
                use_container_width=True
            )
            
            st.markdown("---")
            if st.button("➕ 将此篇加入参考库，继续生成下一期", use_container_width=True):
                next_index = len(st.session_state.virtual_history) + 1
                v_name = f"新生成_追加_{next_index}.md"
                st.session_state.virtual_history.append({"name": v_name, "text": preview_text})
                st.success(f"已成功加入左侧参考库：{v_name}")
                time.sleep(1)
                st.rerun()
        else:
            st.info("暂无生成内容，请先在左侧输入素材并点击开始生成。")
            
        st.write("---")
        st.subheader("📅 历史生成归档")
        hist_scripts = sorted(glob.glob(os.path.join(SCRIPT_OUTPUT_DIR, "*.md")), reverse=True)
        if not hist_scripts:
            st.info("暂无历史记录")
        else:
            d_groups = {}
            for h_sc in hist_scripts:
                try:
                    ds = os.path.basename(h_sc).split('_')[1]
                    dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                except Exception:
                    dfmt = "其他"
                if dfmt not in d_groups: d_groups[dfmt] = []
                d_groups[dfmt].append(h_sc)
            
            for d_key, d_scs in d_groups.items():
                with st.expander(f"📅 {d_key} ({len(d_scs)}份)"):
                    for single_sc in d_scs:
                        sc_name = os.path.basename(single_sc)
                        if st.button(f"📄 {sc_name}", key=f"hist_{single_sc}"):
                            with open(single_sc, "r", encoding="utf-8") as f:
                                st.session_state['generated_script_preview'] = f.read()
                            st.session_state['export_format_choice'] = '.md'
                            st.session_state['show_success_toast'] = "视频脚本生成完成！请在右侧预览"
                            st.rerun()

elif page_selection == "指标文档制作":
    st.title("📈 指标文档制作")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("⚙️ 1. 指标库管理")
        indicators = load_indicators()
        ind_options = ["✨ [新建空白指标]"] + list(indicators.keys())
        
        # Maintain selection state
        if "ind_selected" not in st.session_state: st.session_state["ind_selected"] = ind_options[0]
        sel_ind = st.selectbox("选择或新建指标", ind_options, index=ind_options.index(st.session_state["ind_selected"]) if st.session_state["ind_selected"] in ind_options else 0)
        
        if sel_ind != st.session_state["ind_selected"]:
            st.session_state["ind_selected"] = sel_ind
            st.rerun()
        
        if sel_ind == "✨ [新建空白指标]":
            def_name = ""
            def_code = ""
        else:
            def_name = sel_ind
            def_code = indicators[sel_ind]["code"]
            
        new_ind_name = st.text_input("指标名称", value=def_name, placeholder="例如：震荡顶底模型")
        new_ind_code = st.text_area("指标源码", value=def_code, height=200, placeholder="在此粘贴 Pine Script 或 Python 源码...")
        
        c_btn1, c_btn2 = st.columns(2)
        with c_btn1:
            if st.button("💾 保存/更新", use_container_width=True):
                if new_ind_name.strip() and new_ind_code.strip():
                    save_indicator(new_ind_name.strip(), new_ind_code.strip())
                    st.session_state["ind_selected"] = new_ind_name.strip()
                    st.success(f"指标 '{new_ind_name}' 已保存！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("名称和源码不能为空！")
        with c_btn2:
            if st.button("🗑️ 删除", use_container_width=True):
                if sel_ind != "✨ [新建空白指标]":
                    delete_indicator(sel_ind)
                    st.session_state["ind_selected"] = "✨ [新建空白指标]"
                    st.success(f"指标 '{sel_ind}' 已删除！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.warning("新建状态不可删除")
                    
        st.markdown("---")
        st.subheader("⚡ 2. 生成分析文档")
        
        if not indicators:
            st.info("暂无已保存的指标，请先保存指标后再生成。")
        else:
            gen_ind = st.selectbox("选择目标指标", list(indicators.keys()), index=list(indicators.keys()).index(st.session_state["ind_selected"]) if st.session_state["ind_selected"] in indicators else 0)
            
            if st.button("🚀 生成标准合规版", type="primary", use_container_width=True):
                if not current_chan_config.get("api_key"): st.error("请先在左侧全局配置中填写 API Key！")
                else:
                    with st.status("正在生成标准合规分析文档...", expanded=True) as status:
                        st.write(f"🧠 AI 正在分析 {gen_ind} 的源码...")
                        user_content = f"【指标名称】：{gen_ind}\n【指标源码】：\n{indicators[gen_ind]['code']}"
                        try:
                            client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                            response = llm.call_chat_completion(
                                client,
                                current_chan_config.get("selected_model"),
                                [
                                    {"role": "system", "content": prompts.PROMPT_INDICATOR_STANDARD},
                                    {"role": "user", "content": user_content}
                                ],
                                chan_config=current_chan_config
                            )
                            if not response.choices: raise Exception("接口返回空数据")
                            doc_content = response.choices[0].message.content
                            
                            ts = datetime.now().strftime("%Y%m%d%H%M%S")
                            md_filename = f"《{gen_ind}》---合规化_{ts}.md"
                            docx_filename = f"《{gen_ind}》---合规化_{ts}.docx"
                            md_path = os.path.join(INDICATOR_DOCS_DIR, md_filename)
                            docx_path = os.path.join(INDICATOR_DOCS_DIR, docx_filename)
                            
                            with open(md_path, "w", encoding="utf-8") as f: f.write(doc_content)
                            wechat_render.markdown_to_docx_file(doc_content, docx_path, indicator_name=gen_ind)
                            
                            st.session_state["ind_preview_title"] = f"{gen_ind} (标准合规版)"
                            st.session_state["ind_preview_content"] = doc_content
                            st.session_state["ind_preview_docx"] = docx_path
                            
                            status.update(label="标准文档生成完成！", state="complete", expanded=False)
                            st.rerun()
                        except Exception as e:
                            status.update(label="生成失败", state="error")
                            st.error(f"报错信息: {str(e)}")
                            
            if st.button("✨ 基于标准版一键转化【社群互动教学】版", use_container_width=True):
                if not current_chan_config.get("api_key"): st.error("请先填写 API Key！")
                elif "ind_preview_content" not in st.session_state or "合规化" not in st.session_state.get("ind_preview_title", ""):
                    st.warning("请先生成或从右侧历史预览一份【标准合规版】文档，再进行转化！")
                else:
                    with st.status("正在进行社群风格转换...", expanded=True) as status2:
                        try:
                            client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                            response2 = llm.call_chat_completion(
                                client,
                                current_chan_config.get("selected_model"),
                                [
                                    {"role": "system", "content": prompts.PROMPT_INDICATOR_COMMUNITY},
                                    {"role": "user", "content": f"请将以下标准技术文档转化为社群教学版本：\n\n{st.session_state['ind_preview_content']}"}
                                ],
                                chan_config=current_chan_config
                            )
                            if not response2.choices: raise Exception("接口返回空数据")
                            comm_doc = response2.choices[0].message.content
                            
                            ts = datetime.now().strftime("%Y%m%d%H%M%S")
                            md_filename = f"《{gen_ind}》--社群指标互动手册_{ts}.md"
                            docx_filename = f"《{gen_ind}》--社群指标互动手册_{ts}.docx"
                            md_path = os.path.join(INDICATOR_DOCS_DIR, md_filename)
                            docx_path = os.path.join(INDICATOR_DOCS_DIR, docx_filename)
                            
                            with open(md_path, "w", encoding="utf-8") as f: f.write(comm_doc)
                            wechat_render.markdown_to_docx_file(comm_doc, docx_path, indicator_name=gen_ind)
                            
                            st.session_state["ind_preview_title"] = f"{gen_ind} (社群互动教学版)"
                            st.session_state["ind_preview_content"] = comm_doc
                            st.session_state["ind_preview_docx"] = docx_path
                            
                            status2.update(label="社群文档转换完成！", state="complete", expanded=False)
                            st.rerun()
                        except Exception as e:
                            status2.update(label="转换失败", state="error")
                            st.error(f"报错信息: {str(e)}")

    with col2:
        st.subheader("👁️ 3. 预览与归档")
        
        preview_title = st.session_state.get("ind_preview_title", "预览区")
        preview_content = st.session_state.get("ind_preview_content", "")
        
        with st.container(border=True):
            st.markdown(f"#### {preview_title}")
            if preview_content:
                st.markdown(preview_content)
                docx_path = st.session_state.get("ind_preview_docx")
                if docx_path and os.path.exists(docx_path):
                    with open(docx_path, "rb") as f: docx_data = f.read()
                    st.download_button(
                        label="⬇️ 下载该 Docx 文档",
                        data=docx_data,
                        file_name=os.path.basename(docx_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary"
                    )
            else:
                st.info("右侧暂无预览，请在左侧操作生成或从历史归档中选择。")
                
        st.markdown("---")
        st.write("### 📅 历史报告归档")
        hist_docs = sorted(glob.glob(os.path.join(INDICATOR_DOCS_DIR, "*.md")), reverse=True)
        if not hist_docs:
            st.info("暂无生成的文档归档。")
        else:
            # 按日期分组
            grouped_files = {}
            for h_md in hist_docs:
                basename = os.path.basename(h_md)
                match = re.search(r'_(\d{8})\d{6}\.md$', basename)
                if match:
                    date_str = datetime.strptime(match.group(1), "%Y%m%d").strftime("%Y-%m-%d")
                else:
                    mtime = os.path.getmtime(h_md)
                    date_str = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
                
                if date_str not in grouped_files:
                    grouped_files[date_str] = []
                grouped_files[date_str].append(h_md)
            
            # 日期下拉菜单
            date_list = sorted(grouped_files.keys(), reverse=True)
            selected_date = st.selectbox("选择日期", date_list, index=0, key="hist_date_select")
            
            # 仅展示所选日期的文件
            if selected_date in grouped_files:
                for h_md in grouped_files[selected_date]:
                    basename = os.path.basename(h_md)
                    with st.expander(f"📄 {basename}"):
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("👁️ 在上方预览", key=f"prev_{h_md}", use_container_width=True):
                                with open(h_md, "r", encoding="utf-8") as f: content = f.read()
                                docx_p = h_md.replace(".md", ".docx")
                                st.session_state["ind_preview_title"] = basename.replace(".md", "")
                                st.session_state["ind_preview_content"] = content
                                st.session_state["ind_preview_docx"] = docx_p if os.path.exists(docx_p) else ""
                                st.rerun()
                        with c2:
                            docx_p = h_md.replace(".md", ".docx")
                            if os.path.exists(docx_p):
                                with open(docx_p, "rb") as f: db = f.read()
                                st.download_button(
                                    label="⬇️ 下载 Docx",
                                    data=db,
                                    file_name=os.path.basename(docx_p),
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_{docx_p}",
                                    use_container_width=True
                                )

