import llm
import os
import re
import subprocess
import json
import base64


def markdown_to_docx_file(md_text, filepath, indicator_name="本指标"):
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    # 标题颜色分级体系
    HEADING_COLORS = {
        1: RGBColor(0, 63, 114),    # 深靛蓝 — 一级标题
        2: RGBColor(0, 82, 148),    # 靛蓝 — 二级标题
        3: RGBColor(34, 107, 156),  # 钢蓝 — 三级标题
    }
    
    doc = Document()
    lines = md_text.split('\n')
    
    in_table = False
    table = None
    in_code_block = False
    prev_was_blank = False
    
    def _render_inline(p, text, is_quote=False, is_th=False):
        # 清洗可能存在的全角冒号及嵌套加粗
        text = text.replace("：green[", ":green[").replace("：red[", ":red[").replace("：orange[", ":orange[")
        # 脱掉颜色标签外层的加粗标记 (例如 **:red[升]** -> :red[升])
        text = re.sub(r'\*\*\s*(:(?:green|red|orange)\[.*?\])\s*\*\*', r'\1', text)
        
        pattern = r'(:(?:green|red|orange)\[.*?\]|\*\*.*?\*\*)'
        parts = re.split(pattern, text)
        for part in parts:
            if not part: continue
            run = p.add_run()
            if is_quote:
                run.bold = True
                run.font.color.rgb = RGBColor(64, 64, 64)
            
            if is_th: run.bold = True
                
            if part.startswith(':green[') and part.endswith(']'):
                run.text = part[7:-1]
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.bold = True
            elif part.startswith(':red[') and part.endswith(']'):
                run.text = part[5:-1]
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
            elif part.startswith(':orange[') and part.endswith(']'):
                run.text = part[8:-1]
                run.font.color.rgb = RGBColor(255, 165, 0)
                run.bold = True
            elif part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            else:
                run.text = part
                
    def _set_heading_color(h, level):
        color = HEADING_COLORS.get(level, RGBColor(0, 82, 148))
        for run in h.runs:
            run.font.color.rgb = color
            run.bold = True
                
    for line in lines:
        stripped = line.strip()
        
        # 处理代码块标记
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            prev_was_blank = False
            continue
        
        # 代码块内容：等宽缩进段落
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
            pf = p.paragraph_format
            pf.left_indent = Pt(24)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            prev_was_blank = False
            continue
        
        # 过滤水平分隔线 --- / *** / ___
        if re.match(r'^[-*_]{3,}$', stripped):
            prev_was_blank = False
            continue
        
        # 空行处理：直接跳过，避免生成无意义的大缝隙空行
        if not stripped:
            if in_table:
                in_table = False
            continue
        
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                row_cells = table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells): _render_inline(row_cells[i].paragraphs[0], cell_text, is_th=True)
            else:
                if all(re.match(r'^[-:\s]+$', c) for c in cells): continue 
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells): _render_inline(row_cells[i].paragraphs[0], cell_text)
            continue
        else:
            if in_table: in_table = False
        
        if stripped.startswith('#### '):
            h = doc.add_heading(level=4)
            _render_inline(h, stripped[5:])
            _set_heading_color(h, 3)
        elif stripped.startswith('### '): 
            h = doc.add_heading(level=3)
            _render_inline(h, stripped[4:])
            _set_heading_color(h, 3)
        elif stripped.startswith('## '): 
            h = doc.add_heading(level=2)
            _render_inline(h, stripped[3:])
            _set_heading_color(h, 2)
        elif stripped.startswith('# '): 
            h = doc.add_heading(level=1)
            _render_inline(h, stripped[2:])
            _set_heading_color(h, 1)
        else:
            is_quote = stripped.startswith('>') or stripped.startswith('`>`') or stripped.startswith('`> `')
            is_bullet = stripped.startswith('- ') or stripped.startswith('* ')
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                text = stripped[2:]
                if text.startswith('`>`'):
                    text = text[3:].strip()
                elif text.startswith('`> `'):
                    text = text[4:].strip()
                elif text.startswith('>'):
                    text = text[1:].strip()
            else:
                p = doc.add_paragraph()
                if is_quote:
                    text = stripped
                    if text.startswith('`>'):
                        text = re.sub(r'^`>\s*`?', '', text).strip()
                    else:
                        text = text[1:].strip()
                else:
                    text = stripped
            if text:
                _render_inline(p, text, is_quote)

    # 动态追加美化版的合规免责声明（与指标名称绑定）
    doc.add_paragraph()
    dtbl = doc.add_table(rows=1, cols=1)
    dcell = dtbl.cell(0, 0)
    
    tcPr = dcell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF0F0')
    tcPr.append(shd)
    
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), 'C00000')
    tcBorders.append(left)
    for border_name in ['top', 'right', 'bottom']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    dp1 = dcell.paragraphs[0]
    r1 = dp1.add_run("重要合规化声明：\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(192, 0, 0)
    
    dp2 = dcell.add_paragraph()
    r2 = dp2.add_run(f"以上关于《{indicator_name}》的所有信号提示和区间标注，均是基于历史收盘价等统计数据的技术测算展示，不保证任何未来走势预测的准确性。\n\n")
    r2.font.color.rgb = RGBColor(192, 0, 0)
    
    r3 = dp2.add_run(f"本手册中涉及的所有技术观察标签和计算逻辑仅供技术分析学习与参考，不构成任何形式的投资建议或操作指令。投资者据此操作风险自担，市场有风险，投资需谨慎。请务必结合自身风险承受能力，严格执行止盈止损纪律，独立做出判断。")
    r3.font.color.rgb = RGBColor(192, 0, 0)

    doc.save(filepath)

def markdown_to_wechat_docx_bytes(md_text):
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from io import BytesIO
    import re
    import httpx

    doc = Document()
    lines = md_text.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        # 匹配图片 ![alt](url)
        img_match = re.search(r'!\[.*?\]\((.*?)\)', stripped)
        if img_match:
            img_path = img_match.group(1)
            try:
                if img_path.startswith("http"):
                    resp = httpx.get(img_path, timeout=15)
                    if resp.status_code == 200:
                        doc.add_picture(BytesIO(resp.content), width=Inches(6.0))
                else:
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6.0))
            except Exception:
                pass
            continue
            
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            h = doc.add_heading(level=level)
            h.add_run(stripped[level:].strip())
        elif stripped.startswith('>') or stripped.startswith('`>`') or stripped.startswith('`> `'):
            p = doc.add_paragraph()
            text = stripped
            if text.startswith('`>'):
                text = re.sub(r'^`>\s*`?', '', text).strip()
            else:
                text = text[1:].strip()
            r = p.add_run(text)
            r.font.color.rgb = RGBColor(100, 100, 100)
            r.bold = True
        else:
            p = doc.add_paragraph()
            # 简单处理加粗
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(12)
        
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def make_preview_html(html_raw):
    """为 Streamlit 预览将 HTML 中的本地图片路径转换为 base64 编码，解决 sandbox iframe 限制"""
    import re
    import os
    import base64
    if not html_raw:
        return html_raw
    def replace_src(match):
        prefix = match.group(1)
        src = match.group(2)
        if not src.startswith("http") and not src.startswith("data:") and os.path.exists(src):
            try:
                with open(src, "rb") as f_f:
                    b64 = base64.b64encode(f_f.read()).decode()
                ext = os.path.splitext(src)[1].lower().replace(".", "")
                if ext == "jpg": ext = "jpeg"
                return f'{prefix}="data:image/{ext};base64,{b64}"'
            except Exception:
                pass
        return match.group(0)
    return re.sub(r'(src)=["\']([^"\']+)["\']', replace_src, html_raw, flags=re.IGNORECASE)

def post_process_wechat_html(html_raw, theme, title_color=None):
    import re
    if not html_raw:
        return html_raw
        
    accent_color = "#d97758"  # 默认暖橙 (autumn-warm)
    title_color_default = "#4a413d"   # 默认大标题字色 (黑褐)
    if theme == "spring-fresh":
        accent_color = "#6b9b7a"  # 嫩绿
        title_color_default = "#3d4a3d"   # 深绿灰
    elif theme == "ocean-calm":
        accent_color = "#4a7c9b"  # 蔚蓝
        title_color_default = "#3a4150"   # 深蓝灰
        
    t_color = title_color if title_color else title_color_default

    # 重点强调 (strong/b) 字色强制高亮为主题强调色，并转化为 span 标签以防止微信编辑器强制脱落加粗样式
    def highlight_strong(match):
        tag = match.group(1)
        attrs = match.group(2) or ""
        content = match.group(3)
        if 'style=' in attrs or 'style =' in attrs:
            def repl_style(m):
                style_content = m.group(2)
                # 滤除原先可能带有的任何 color 声明和 font-weight 声明
                style_content = re.sub(r'color\s*:\s*[^;]+;?', '', style_content).strip()
                style_content = re.sub(r'font-weight\s*:\s*[^;]+;?', '', style_content).strip()
                # 拼装新的主题色与强力加粗
                return f'style="{style_content}; color: {accent_color}; font-weight: bold !important;"'
            new_attrs = re.sub(r'style\s*=\s*(["\'])(.*?)(\1)', repl_style, attrs)
        else:
            new_attrs = attrs + f' style="color: {accent_color}; font-weight: bold !important;"'
        return f'<span{new_attrs}>{content}</span>'

    html_raw = re.sub(r'<(strong|b)(\s+[^>]*?)?>(.*?)</\1>', highlight_strong, html_raw, flags=re.IGNORECASE | re.DOTALL)

    # 自动将所有 h1-h6 标题内部的文本用 <span> 包裹，以在微信后台强制保留颜色与粗体样式，并去除默认下划线
    def wrap_headings_in_span(match):
        tag = match.group(1).lower()
        attrs = match.group(2) or ""
        content = match.group(3)
        
        # 默认标题的大小和行高
        default_sizes = {
            'h1': '24px',
            'h2': '20px',
            'h3': '18px',
            'h4': '16px',
            'h5': '15px',
            'h6': '14px'
        }
        size_val = default_sizes.get(tag, '16px')
        
        # 1. 确保 h 标签本身具有 font-weight: bold !important; 和 text-decoration: none !important;
        # 以及默认的 font-size 和 line-height (如果原样式中没有设定)
        if 'style=' in attrs or 'style =' in attrs:
            def repl_h_style(m):
                style_content = m.group(2)
                # 统一清理现有的 text-decoration 和 font-weight 声明，以防冲突
                style_content = re.sub(r'text-decoration\s*:\s*[^;]+;?', '', style_content).strip()
                style_content = re.sub(r'font-weight\s*:\s*[^;]+;?', '', style_content).strip()
                
                # 检查是否包含 font-size 和 line-height
                extra_styles = []
                if 'font-size' not in style_content:
                    extra_styles.append(f"font-size: {size_val}")
                if 'line-height' not in style_content:
                    extra_styles.append("line-height: 1.5")
                    
                style_content_str = style_content
                if not style_content_str.endswith(';') and style_content_str:
                    style_content_str += ';'
                    
                extra_str = "; ".join(extra_styles) + ";" if extra_styles else ""
                
                return f'style="{style_content_str} {extra_str} text-decoration: none !important; font-weight: bold !important;"'
            new_attrs = re.sub(r'style\s*=\s*(["\'])(.*?)(\1)', repl_h_style, attrs)
        else:
            new_attrs = attrs + f' style="font-size: {size_val}; line-height: 1.5; text-decoration: none !important; font-weight: bold !important;"'
            
        # 2. 对 content 里的文本进行包装：
        # 如果 content 里已经含有 span 标签，我们把 span 标签内的 style 里的 font-weight 强行替换为 font-weight: bold !important;
        # 并且将 font-size 也强行注入/替换为对应的标题大小，确保微信不会覆盖 span 内的字体大小！
        if '<span' in content.lower():
            def repl_span_style(m_span):
                span_attrs = m_span.group(1) or ""
                span_content = m_span.group(2)
                if 'style=' in span_attrs or 'style =' in span_attrs:
                    def repl_inner_style(m_inner):
                        inner_style = m_inner.group(2)
                        inner_style = re.sub(r'font-weight\s*:\s*[^;]+;?', '', inner_style).strip()
                        inner_style = re.sub(r'font-size\s*:\s*[^;]+;?', '', inner_style).strip()
                        return f'style="{inner_style}; font-weight: bold !important; font-size: {size_val} !important;"'
                    new_span_attrs = re.sub(r'style\s*=\s*(["\'])(.*?)(\1)', repl_inner_style, span_attrs)
                else:
                    new_span_attrs = span_attrs + f' style="font-weight: bold !important; font-size: {size_val} !important;"'
                return f'<span{new_span_attrs}>{span_content}</span>'
            
            new_content = re.sub(r'<span(\s+[^>]*?)?>(.*?)</span>', repl_span_style, content, flags=re.IGNORECASE | re.DOTALL)
            return f'<{tag}{new_attrs}>{new_content}</{tag}>'
        else:
            c = t_color if tag == 'h1' else accent_color
            return f'<{tag}{new_attrs}><span style="color: {c}; font-size: {size_val} !important; font-weight: bold !important; text-decoration: none !important;">{content}</span></{tag}>'

    html_raw = re.sub(r'<(h[1-6])(\s+[^>]*?)?>(.*?)</\1>', wrap_headings_in_span, html_raw, flags=re.IGNORECASE | re.DOTALL)
    return html_raw

def convert_to_wechat_html(md_text, theme, mode_ui, api_key=None, font_size="medium", bg_type="none", chan_config=None, custom_prompt="", for_wechat_api=False):
    # 强制微信转换函数输出干净的包含原始本地路径的 HTML 格式数据，不在转换内部处理 base64
    for_wechat_api = True
    
    import subprocess
    import json
    import re
    import base64
    from openai import OpenAI
    
    input_path = os.path.join("tests", "temp_wechat_input.md")
    output_path = os.path.join("tests", "temp_wechat_output.html")
    
    if os.path.exists(output_path):
        try: os.remove(output_path)
        except Exception: pass
        
    with open(input_path, "w", encoding="utf-8") as f:
        f.write(md_text)
        
    mode = "api" if "API" in mode_ui else "ai"
    
    accent_color = "#d97758"  # 默认暖橙 (autumn-warm)
    title_color = "#4a413d"   # 默认大标题字色 (黑褐)
    if theme == "spring-fresh":
        accent_color = "#6b9b7a"  # 嫩绿
        title_color = "#3d4a3d"   # 深绿灰
    elif theme == "ocean-calm":
        accent_color = "#4a7c9b"  # 蔚蓝
        title_color = "#3a4150"   # 深蓝灰
    
    # 提取首个标题并进行安全长度截断（不超过 32 个字符）
    title = "技术分析报告"
    for line in md_text.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            parsed_title = re.sub(r'^#+\s*', '', line).strip()
            if parsed_title:
                title = parsed_title
                break
    if len(title) > 10:
        title = title[:7] + "..."
        
    cmd = [
        NPX_CMD, "md2wechat", "convert", input_path, 
        "--mode", mode, 
        "--theme", theme, 
        "--preview", 
        "--output", output_path, 
        "--json",
        "--title", title,
        "--author", "AI",
        "--digest", "技术观察与趋势分析"
    ]

    
    if mode == "api" and api_key:
        cmd.extend(["--api-key", api_key])
        cmd.extend(["--font-size", font_size])
        cmd.extend(["--background-type", bg_type])
        
    if mode == "ai" and custom_prompt and theme == "custom":
        cmd.extend(["--custom-prompt", custom_prompt])
        
    # 临时 debug 日志记录
    try:
        with open("tests/wechat_cli_debug.log", "a", encoding="utf-8") as debug_f:
            debug_f.write("=== convert_to_wechat_html CLI CALL ===\n")
            debug_f.write(f"Title: {repr(title)}\n")
            debug_f.write(f"Cmd: {repr(cmd)}\n")
    except Exception:
        pass

    res = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
    
    try:
        with open("tests/wechat_cli_debug.log", "a", encoding="utf-8") as debug_f:
            debug_f.write(f"Returncode: {res.returncode}\n")
            debug_f.write(f"Stdout: {repr(res.stdout)}\n")
            debug_f.write(f"Stderr: {repr(res.stderr)}\n")
    except Exception:
        pass
    
    try:
        m_res = re.search(r'\{.*\}', res.stdout, re.DOTALL)
        if m_res:
            res_data = json.loads(m_res.group(0))
        else:
            res_data = {"success": False, "message": f"无法解析 CLI 返回的 JSON 数据。原始输出: {res.stdout}"}
    except Exception as e:
        res_data = {"success": False, "message": f"解析异常: {str(e)}"}
        
    html_raw = ""
    if mode == "api":
        if not res_data.get("success"):
            err_msg = res_data.get("message", "API 转换失败。")
            debug_info = f"<br><b>[Debug Info]</b><br>Cmd: <code>{re.sub(r'--api-key\s+\S+', '--api-key ******', ' '.join(cmd))}</code><br>Stdout: <pre>{res.stdout}</pre><br>Stderr: <pre>{res.stderr}</pre>"
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 微信排版美化失败 (API 模式)</h3><p>{err_msg}</p>{debug_info}<p>请检查您的 md2wechat API Key 是否正确。</p></div>"
        
        if os.path.exists(output_path):
            with open(output_path, "r", encoding="utf-8") as f:
                html_raw = f.read()
        else:
            return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 转换成功但未生成 HTML 目标文件。</div>"
        
    else: # mode == "ai"
        if not res_data.get("success") or res_data.get("code") != "CONVERT_AI_REQUEST_READY":
            err_msg = res_data.get("message", "AI 模式初始化失败。")
            debug_info = f"<br><b>[Debug Info]</b><br>Cmd: <code>{' '.join(cmd)}</code><br>Stdout: <pre>{res.stdout}</pre><br>Stderr: <pre>{res.stderr}</pre>"
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 微信排版美化失败 (AI 模式初始化)</h3><p>{err_msg}</p>{debug_info}</div>"
            
        ai_prompt = res_data.get("data", {}).get("prompt", "")
        if not ai_prompt:
            return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 转换初始化失败：未获取到 AI 排版 Prompt 提示词。</div>"
            
        if not chan_config or not chan_config.get("api_key"):
            return "<div style='color:red;padding:20px;font-family:sans-serif;'><h3>⚠️ 大模型未配置</h3><p>AI 模式需要使用您的 OpenAI/Gemini 大模型 API Key 进行样式翻译。请在左侧全局配置中填写 API Key！</p></div>"
            
        try:
            def find_anchors_for_image(md_text, target_index):
                lines = md_text.split('\n')
                img_pattern = re.compile(r'!\[.*?\]\(.*?\)')
                
                current_img_idx = 0
                target_line_idx = -1
                
                for idx, line in enumerate(lines):
                    if img_pattern.search(line):
                        if current_img_idx == target_index:
                            target_line_idx = idx
                            break
                        current_img_idx += 1
                        
                if target_line_idx == -1:
                    return None, None
                    
                # 向上寻找非空非图行作为前置锚点
                pre_anchor = None
                for i in range(target_line_idx - 1, -1, -1):
                    line = lines[i].strip()
                    if line and not img_pattern.search(line) and not line.startswith('---'):
                        cleaned = re.sub(r'[#\*_`\-\>\+]', '', line).strip()
                        if len(cleaned) > 2:
                            pre_anchor = cleaned
                            break
                            
                # 向下寻找非空非图行作为后置锚点
                post_anchor = None
                for i in range(target_line_idx + 1, len(lines)):
                    line = lines[i].strip()
                    if line and not img_pattern.search(line) and not line.startswith('---'):
                        cleaned = re.sub(r'[#\*_`\-\>\+]', '', line).strip()
                        if len(cleaned) > 2:
                            post_anchor = cleaned
                            break
                            
                return pre_anchor, post_anchor

            # 强化 ai_prompt 中的图片和排版规则
            if ai_prompt:
                # 0. 规范化并明确图片索引从 0 开始计数，防止部分主题（如 spring-fresh）因为缺少示例导致 LLM 从 1 开始索引
                ai_prompt = ai_prompt.replace(
                    "图片使用占位符格式：<!-- IMG:index -->，例如第一张图用 <!-- IMG:0 -->",
                    "图片使用占位符格式：<!-- IMG:index -->"
                )
                ai_prompt = ai_prompt.replace(
                    "图片使用占位符格式：<!-- IMG:index -->",
                    "图片使用占位符格式：<!-- IMG:index -->，其中索引从 0 开始（即第一张图用 <!-- IMG:0 -->，第二张图用 <!-- IMG:1 -->，依此类推）"
                )

                # 1. 物理剥离 ![说明文字](<!-- IMG:index -->) 的 alt 说明文字，防止生成配图卡片，仅保留无文本图片节点
                ai_prompt = re.sub(
                    r'!\[.*?\]\(<!--\s*IMG:(\d+)\s*-->\)', 
                    r'<div class="wechat-img-placeholder" data-index="\1"></div>', 
                    ai_prompt
                )
                
                # 2. 替换漏网的 <!-- IMG:index --> 为自定义标签
                ai_prompt = re.sub(
                    r'<!--\s*IMG:(\d+)\s*-->', 
                    r'<div class="wechat-img-placeholder" data-index="\1"></div>', 
                    ai_prompt
                )

                # 根据不同主题自适应主标题的强调色、文字色以及悬浮卡片样式
                accent_color = "#d97758"  # 默认暖橙 (autumn-warm)
                eyebrow_color = "#8a7e72" # 默认灰褐
                title_color = "#4a413d"   # 默认黑褐
                card_style = "background-color: #ffffff; border: 1px solid rgba(0, 0, 0, 0.05); box-shadow: 0 10px 30px rgba(0, 0, 0, 0.04), 0 0 15px rgba(217, 119, 88, 0.1); border-radius: 18px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"
                
                if theme == "spring-fresh":
                    accent_color = "#6b9b7a"  # 嫩绿
                    eyebrow_color = "#5c735c" # 绿灰
                    title_color = "#3d4a3d"   # 深绿灰
                    card_style = "background-color: #ffffff; border: 1px solid rgba(107, 155, 122, 0.1); box-shadow: 0 8px 24px rgba(74, 128, 88, 0.08), 0 0 12px rgba(107, 155, 122, 0.1); border-radius: 16px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"
                elif theme == "ocean-calm":
                    accent_color = "#4a7c9b"  # 蔚蓝
                    eyebrow_color = "#556e80" # 石蓝
                    title_color = "#3a4150"   # 深蓝灰
                    card_style = "background-color: #ffffff; border: 1px solid rgba(74, 124, 155, 0.08); box-shadow: 0 8px 28px rgba(58, 65, 80, 0.06), 0 0 16px rgba(74, 124, 155, 0.1); border-radius: 14px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"

                image_rule_enhancement = f"""
【重要：主标题 (H1) 固定排版规范（核心必做）】
1. 一级大标题（H1，即 `# 标题内容`）代表文章的主标题。你必须将其设计成一个极其典雅、浮动卡片式（Floating Card Layout）的头部区域（Header Region），位于网页最顶部。
2. 你必须将主标题 H1 的所有内容，包裹在一个独立的白色卡片容器中，其内联样式（style）必须严格设为：`{card_style}`，使其在视觉上呈现出与下文卡片相同的“悬浮卡片”质感。
3. 卡片容器内的元素必须严格按照以下顺序排列，并保证完美居中对齐：
   - 顶部分类眉标（Eyebrow）：在标题正上方放置一行小字，内容为“行业深度解析 | INDUSTRY ANALYSIS”，样式为：`font-size: 13px; color: {eyebrow_color}; letter-spacing: 2px; text-transform: uppercase; margin-bottom: 12px; font-weight: bold; text-align: center;`，居中。
   - 中间大标题：在眉标下方。使用 `<h1>` 标签，样式为：`font-size: 26px; line-height: 1.4; color: {title_color}; font-weight: bold; margin: 10px 0; font-family: 'Georgia', 'Source Serif Pro', serif; text-align: center; text-decoration: none;`，居中对齐。
   - 底部装饰短线：在标题正下方，居中放置一个装饰短横线：`<div style="margin: 15px auto 5px auto; width: 50px; height: 2px; background-color: {accent_color};"></div>`，用于与标题外部的其他正文卡片进行优雅的视觉区隔。
4. 确保 H1 卡片的宽度自适应，且其外层没有任何其他的嵌套卡片容器包裹。

【重要：图片与防溢出排版补充要求】
1. 在转换内容中，你必须原样保留图片占位符标签 `<div class="wechat-img-placeholder" data-index="index"></div>` 并在生成的 HTML 对应位置呈现出来。
2. 绝对禁止修改、删除或包裹此标签属性（如 class 或 data-index），严禁在此标签内写入任何子节点、文字、空格或 HTML 注释。它必须作为单独一行的独立块级节点。
3. 绝对禁止为图片生成任何多余的配图说明卡片或提示性文字（例如不要输出类似“【此处是手机芯片图片】”或说明该图用意、图片标题的字样）。
4. 最外层主背景 <div> 容器及所有卡片必须使用自适应宽度（如 `width: 100%;`），并指定 `box-sizing: border-box;`。
5. 所有一级、二级、三级标题以及段落等，如果使用了带边框或背景颜色的卡片，请确保内容不超出卡片的视觉边缘。
"""
                # 将补充规则插入到 md 转换内容之前
                if "请转换以下 Markdown内容：" in ai_prompt:
                    ai_prompt = ai_prompt.replace("请转换以下 Markdown内容：", image_rule_enhancement + "\n\n请转换以下 Markdown内容：")
                else:
                    ai_prompt = ai_prompt + "\n\n" + image_rule_enhancement

            client = OpenAI(api_key=chan_config.get("api_key"), base_url=chan_config.get("base_url") if chan_config.get("base_url") else None)
            response = llm.call_chat_completion(
                client,
                chan_config.get("selected_model"),
                [
                    {"role": "system", "content": (
                        "You are a professional WeChat push designer. You must return only a clean HTML fragment (or body content only) with inline CSS. "
                        "Strictly do not wrap it in any comments or explanations, just raw HTML or markdown-fenced HTML code blocks.\n"
                        "【排版设计美化强力指示（至关重要）】:\n"
                        "1. 必须使用丰富的内联样式（inline CSS）为推文进行色彩与板块美化设计。严格禁止使用任何 <style> 标签或 CSS class 选择器，所有样式必须直接写在每个元素的 style=\"...\" 内联属性中，因为微信公众号后台会剥离外部 style 样式表。\n"
                        "2. 必须包含大量精心美化的组件与样式元素，例如：\n"
                        "   - 带有柔和浅色背景色（如浅橘、淡绿、浅蓝等）、圆角（border-radius: 8px或12px）和阴影（box-shadow: 0 4px 12px rgba(0,0,0,0.05)）的卡片式展示盒子（cards）。\n"
                        "   - 带有左侧粗色条修饰（border-left: 4px solid ...）和淡背景色的引用块（blockquote）。\n"
                        "   - 带有特色标签、图标或彩色背景框的列表项（li）和表格（table）。\n"
                        "   - 加粗并以主题主色高亮的关键字或句段。\n"
                        "3. 第一个一级大标题（H1）必须设计得极具视觉冲击力，可使用较大字号、居中对齐、网页最顶部，并包含 text-decoration: none !important; 样式以防止微信默认下划线干扰。\n"
                        "4. 所有一级、二级、三级标题（h1/h2/h3）的 style 属性中，必须包含 text-decoration: none !important;，防止微信默认的下划虚线或实线影响排版美观。\n"
                        "5. 绝不能输出 html, head, body 等包装标签，直接返回最外层美化用 <div> 容器包裹的网页内容片段。\n"
                        "6. 绝对禁止以任何形式增删、改写或缩减原文的段落或字词内容！你的职责仅仅是对文章做 HTML 排版和样式美化包装，必须保留全部文字细节与配图占位符。\n"
                        "7. 严格禁止在排版生成的 HTML 中出现任何 Emoji 表情符号。如果检测到原文中带有 Emoji，请将其滤除或用等价的文字描述替代。"
                    )},
                    {"role": "user", "content": ai_prompt}
                ],
                chan_config=chan_config
            )
            
            if not response.choices or not response.choices[0].message.content:
                return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 大模型未返回 HTML 排版内容。</div>"
                
            html_raw = response.choices[0].message.content.strip()
            
            if "```html" in html_raw:
                html_raw = html_raw.split("```html")[1].split("```")[0].strip()
            elif "```" in html_raw:
                html_raw = html_raw.split("```")[1].split("```")[0].strip()

            # 物理清洗 HTML 标签，只保留正文部分以防微信 API 剥离外部 style 导致无样式
            # 1. 剥离 style 和 head 块
            html_raw = re.sub(r'<style[^>]*>.*?</style>', '', html_raw, flags=re.DOTALL | re.IGNORECASE)
            html_raw = re.sub(r'<head[^>]*>.*?</head>', '', html_raw, flags=re.DOTALL | re.IGNORECASE)
            # 2. 提取 body 内正文
            body_match = re.search(r'<body[^>]*>(.*?)</body>', html_raw, re.DOTALL | re.IGNORECASE)
            if body_match:
                html_raw = body_match.group(1).strip()
            else:
                # 3. 如果没有 body，去除 html, doctype 等外层包装标签
                html_raw = re.sub(r'<!DOCTYPE[^>]*>', '', html_raw, flags=re.IGNORECASE)
                html_raw = re.sub(r'<html[^>]*>', '', html_raw, flags=re.IGNORECASE)
                html_raw = re.sub(r'</html>', '', html_raw, flags=re.IGNORECASE)
                html_raw = html_raw.strip()

                
            # Replace image placeholders
            images_list = res_data.get("data", {}).get("images") or []
            replaced_indices = set()
            
            # 首先遍历图片列表进行占位符替换
            for idx, img in enumerate(images_list):
                index = img.get("Index") if "Index" in img else img.get("index", idx)
                original = img.get("Original") if "Original" in img else img.get("original", "")
                
                src = original
                if not for_wechat_api and src and not src.startswith("http") and os.path.exists(src):
                    try:
                        with open(src, "rb") as f_img:
                            img_b64 = base64.b64encode(f_img.read()).decode()
                        ext = os.path.splitext(src)[1].lower().replace(".", "")
                        if ext == "jpg": ext = "jpeg"
                        src = f"data:image/{ext};base64,{img_b64}"
                    except Exception:
                        pass
                
                img_tag = f'<img src="{src}" style="width: 100%; max-width: 100%; border-radius: 8px; margin: 16px 0; display: block; height: auto;" />'
                
                # 匹配类似 <div class="wechat-img-placeholder" data-index="0"></div> 的结构，允许颠倒顺序和自闭合
                placeholder_regex = re.compile(
                    rf'<div\s+[^>]*(?:class=["\']wechat-img-placeholder["\']|data-index=["\']{index}["\'])[^>]+(?:class=["\']wechat-img-placeholder["\']|data-index=["\']{index}["\'])[^>]*>(?:\s*</div>)?',
                    re.IGNORECASE
                )
                
                if placeholder_regex.search(html_raw):
                    html_raw = placeholder_regex.sub(img_tag, html_raw)
                    replaced_indices.add(index)
                    
            # 其次，针对大模型丢失占位符的情况，启动双向段落锚点兜底匹配强插机制
            for idx, img in enumerate(images_list):
                index = img.get("Index") if "Index" in img else img.get("index", idx)
                if index in replaced_indices:
                    continue
                    
                original = img.get("Original") if "Original" in img else img.get("original", "")
                src = original
                if not for_wechat_api and src and not src.startswith("http") and os.path.exists(src):
                    try:
                        with open(src, "rb") as f_img:
                            img_b64 = base64.b64encode(f_img.read()).decode()
                        ext = os.path.splitext(src)[1].lower().replace(".", "")
                        if ext == "jpg": ext = "jpeg"
                        src = f"data:image/{ext};base64,{img_b64}"
                    except Exception:
                        pass
                        
                img_tag = f'<img src="{src}" style="width: 100%; max-width: 100%; border-radius: 8px; margin: 16px 0; display: block; height: auto;" />'
                
                pre_anchor, post_anchor = find_anchors_for_image(md_text, index)
                inserted = False
                
                if pre_anchor:
                    pre_clean = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', pre_anchor).strip()
                    pos = html_raw.find(pre_clean[:25]) if len(pre_clean) > 25 else html_raw.find(pre_clean)
                    if pos != -1:
                        end_tag_pos = html_raw.find('>', pos)
                        if end_tag_pos != -1:
                            html_raw = html_raw[:end_tag_pos + 1] + "\n" + img_tag + "\n" + html_raw[end_tag_pos + 1:]
                            inserted = True
                            
                if not inserted and post_anchor:
                    post_clean = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', post_anchor).strip()
                    pos = html_raw.find(post_clean[:25]) if len(post_clean) > 25 else html_raw.find(post_clean)
                    if pos != -1:
                        start_tag_pos = html_raw.rfind('<', 0, pos)
                        if start_tag_pos != -1:
                            html_raw = html_raw[:start_tag_pos] + "\n" + img_tag + "\n" + html_raw[start_tag_pos:]
                            inserted = True
                            
                if not inserted:
                    if "</div>" in html_raw:
                        last_div = html_raw.rfind("</div>")
                        html_raw = html_raw[:last_div] + "\n" + img_tag + "\n" + html_raw[last_div:]
                    else:
                        html_raw = html_raw + "\n" + img_tag
                        
        except Exception as e:
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 大模型调用渲染失败</h3><p>{str(e)}</p></div>"

    # For both API and AI modes, convert any local image paths in src="..." to base64 data URIs
    # to bypass Streamlit iframe sandbox restrictions
    if not for_wechat_api:
        def replace_src(match):
            prefix = match.group(1)
            src = match.group(2)
            if not src.startswith("http") and not src.startswith("data:") and os.path.exists(src):
                try:
                    with open(src, "rb") as f_f:
                        b64 = base64.b64encode(f_f.read()).decode()
                    ext = os.path.splitext(src)[1].lower().replace(".", "")
                    if ext == "jpg": ext = "jpeg"
                    return f'{prefix}="data:image/{ext};base64,{b64}"'
                except Exception:
                    pass
            return match.group(0)

        html_raw = re.sub(r'(src)=["\']([^"\']+)["\']', replace_src, html_raw, flags=re.IGNORECASE)

    # 注入全局兼容性与防溢出样式，解决 450px viewport 长图导出及微信预览时发生的内容拉宽、背景托不住内容的问题
    compat_css = """
<style id="wechat-compat-styles">
  /* 强制全局 box-sizing 以便正确计算内边距 */
  *, *:before, *:after {
    box-sizing: border-box !important;
  }
  /* 限制最大宽度，防止子元素溢出主背景 */
  html, body {
    margin: 0;
    padding: 0;
    width: 100% !important;
    max-width: 100% !important;
    overflow-x: hidden !important;
  }
  /* 特别防止各种块级组件和图片溢出 */
  section, div, p, span, img, table, pre, code {
    max-width: 100% !important;
    word-wrap: break-word !important;
    word-break: break-word !important;
  }
  /* 代码块和表格在太宽时允许横向滚动，不可撑破父容器 */
  pre, code, table {
    overflow-x: auto !important;
    white-space: pre-wrap !important;
  }
  /* 图片自适应限制 */
  img {
    max-width: 100% !important;
    height: auto !important;
    display: block !important;
  }
</style>
"""
    if "</head>" in html_raw:
        html_raw = html_raw.replace("</head>", f"{compat_css}</head>")
    elif "<body>" in html_raw:
        html_raw = html_raw.replace("<body>", f"<body>{compat_css}")
    else:
        html_raw = compat_css + html_raw
        
    # 物理删除所有 Emoji 表情符号
    emoji_pattern = re.compile(r'[\U00010000-\U0010ffff]', flags=re.UNICODE)
    html_raw = emoji_pattern.sub('', html_raw)
    
    html_raw = post_process_wechat_html(html_raw, theme, title_color=title_color)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_raw)

    return html_raw

def validate_wechat_article_content(text: str) -> tuple[bool, str]:
    """
    检查生成的微信推文/源文本是否包含 AI 拒答/错误占位词汇，或者有效字数过少。
    返回 (is_valid, reason_message)
    """
    if not text or not text.strip():
        return False, "文章内容为空"
    
    cleaned = text.strip()
    if len(cleaned) < 120:
        return False, f"文章有效字数过少 (仅 {len(cleaned)} 字)，疑似异常占位符"
        
    invalid_keywords = [
        "作为AI语言模型",
        "作为AI，我无法",
        "作为AI我无法",
        "抱歉，我无法",
        "很抱歉，我无法",
        "无法为您生成",
        "无法为您提供",
        "由于未获取到任何动态",
        "跳过分析：未找到",
        "非法请求",
        "抱歉，我不能",
        "无法提供此类内容",
        "由于缺少可用的行业数据，无法",
        "无法完成分析任务",
        "请提供真实的",
        "占位提示：",
        "SYSTEM NOTICE:",
    ]
    
    for kw in invalid_keywords:
        if kw in cleaned:
            return False, f"检测到 AI 拒答/错误占位关键词: '{kw}'"
            
    return True, "校验通过"

