import streamlit as st
import os
import base64
import time
import httpx
import json
import glob
import markdown
import subprocess
import fitz  # PyMuPDF
import re
from datetime import datetime
from docx import Document
from openai import OpenAI
from playwright.sync_api import sync_playwright
from apscheduler.schedulers.background import BackgroundScheduler
import sys
import wechat_publisher
import importlib
importlib.reload(wechat_publisher)

NPX_CMD = "npx.cmd" if sys.platform == "win32" else "npx"
DREAMINA_CMD = "dreamina.exe" if sys.platform == "win32" else "./dreamina"

# --- Config & Constants ---
IMAGE_OUTPUT_DIR = os.path.join("outputs", "images")
SCRIPT_OUTPUT_DIR = os.path.join("outputs", "scripts")
INDICATOR_DOCS_DIR = os.path.join("outputs", "indicator_docs")
WECHAT_OUTPUT_DIR = os.path.join("outputs", "wechat")
WECHAT_IMAGES_DIR = os.path.join("outputs", "wechat", "images")
CONFIG_FILE = "config.json"
os.makedirs(IMAGE_OUTPUT_DIR, exist_ok=True)
os.makedirs(SCRIPT_OUTPUT_DIR, exist_ok=True)
os.makedirs(INDICATOR_DOCS_DIR, exist_ok=True)
os.makedirs(WECHAT_OUTPUT_DIR, exist_ok=True)
os.makedirs(WECHAT_IMAGES_DIR, exist_ok=True)

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_config(data):
    tmp_file = CONFIG_FILE + ".tmp"
    with open(tmp_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    os.replace(tmp_file, CONFIG_FILE)

INDICATORS_FILE = "indicators.json"

PROMPT_INDICATOR_STANDARD = """你是一个顶级金融技术分析专家，请根据提供的指标源码，严格按照以下结构输出技术观察文档：

【绝对红线：禁止金融预测与违规交易指令】
你的生成内容必须 100% 合规。严禁出现任何具有明确方向预测、交易暗示或主观情绪的词汇。
1. **违禁词库（绝对不可使用）**：抄底、逃顶、建仓、满仓、空仓、加仓、减仓、买入、卖出、暴涨、暴跌、必涨、必跌、主力、庄家、洗盘、拉升、砸盘、黑马、牛股、底部确立。
2. **强制替换规则**：
   - 将所有带有方向色彩的词替换为中立词汇，例如："抄底" -> "低位关注" 或 "潜在企稳"；"逃顶" -> "高位风险"；"买入/卖出" -> "信号触发/条件满足"；"趋势线<11" 只能表述为 "进入低位观察区"。
3. **安全缓冲用语**：所有描述必须加上“可能”、“潜在”、“或将”、“观察”等缓冲词，必须明确强调指标只是历史数据的概率统计，不构成对未来确定性的预测。

【绝对红线：禁止暴露代码规则（面向小白）】
用户是没有任何编程基础的金融小白，他们在使用软件时只能看到最终的指标图形（如：红绿柱体、黄白曲线、买卖文字提示），绝对看不到底层的代码逻辑。
1. **禁止暴露代码**：绝对禁止在文档中出现任何原始代码段、函数名（如 FILTER、CROSS、EMA、SMA）、代码变量名（如 C、H、L、O）或计算公式。
2. **强制图形化翻译**：必须将所有代码逻辑全部翻译为“盘面图形的视觉表现大白话”。例如：
   - 错误：“C上穿均线” / 正确：“当K线实体向上突破参考线”
   - 错误：“FILTER(...,15)” / 正确：“信号触发后会进行一段时间的观察，过滤掉频繁闪烁的无效信号”
   - 错误：“最低价+波动幅度*0.5” / 正确：“基于近期波动幅度测算的低位参考线”
   - 错误：“元素：`支撑`” / 正确：“元素：底部参考线”

【强制输出排版与色彩规范】
1. 必须使用 Markdown 层级（## 二级标题、### 三级标题）。核心结论必须使用 `>` 引用块（以总结性的大白话呈现）。
2. 绝对禁止在全文中使用任何 Emoji 表情符号（坚决删除所有 Emoji）。
3. **强制色彩标注（系统会做拦截处理，请严格按此格式写）**：
   - 遇到【风险、顶部、空头、压力、预警、超买】等词汇，必须使用 `:green[词汇]` 或 `:orange[词汇]` 进行包裹。
   - 遇到【底部、多头、支撑、低位、企稳、机会】等词汇，必须使用 `:red[词汇]` 进行包裹。
   - 强调性的核心逻辑或公式、重要参数，请直接使用纯文字或 **加粗**。
   - ⚠️ 警告：千万不要把色彩标签和加粗标签嵌套使用（绝不能写 `**:red[词汇]**` 或 `:red[**词汇**]`），色彩标签必须独立使用！且冒号必须是半角英文字符！
4. **面向小白用户**：内容必须尽量通俗易懂，结论前置，然后再用大白话补充解释原因。

【强制输出章节】
## 一、指标定位
## 二、盘面视觉与核心逻辑
## 三、合规化技术观察标签解读
## 四、信号解析与实战观察流程
## 五、特别教学与重要风险提示"""

PROMPT_INDICATOR_COMMUNITY = """你现在是一位幽默、接地气且拥有多年实盘经验的资深交易技术讲师，主要面向社群用户互动。
请根据以下标准技术文档，将其转化为高互动的社群语境教学文档：

【绝对红线：禁止金融预测与违规交易指令】
你的生成内容必须 100% 合规。严禁出现任何具有明确方向预测、交易暗示或主观情绪的词汇。
1. **违禁词库（绝对不可使用）**：抄底、逃顶、建仓、满仓、空仓、加仓、减仓、买入、卖出、暴涨、暴跌、必涨、必跌、主力、庄家、洗盘、拉升、砸盘、黑马、牛股、底部确立。
2. **强制替换规则**：
   - 将所有带有方向色彩的词替换为中立词汇，例如："抄底" -> "低位关注" 或 "潜在企稳"；"逃顶" -> "高位风险"；"买入/卖出" -> "信号触发/条件满足"。
3. **安全缓冲用语**：所有描述必须加上“可能”、“潜在”、“或将”、“观察”等缓冲词，必须明确强调指标只是历史数据的概率统计，不构成对未来确定性的预测。

【绝对红线：禁止暴露代码规则（面向小白）】
用户是没有任何编程基础的金融小白，他们在使用软件时只能看到最终的指标图形（如：红绿柱体、黄白曲线、买卖文字提示），绝对看不到底层的代码逻辑。
1. **禁止暴露代码**：绝对禁止在文档中出现任何原始代码段、函数名（如 FILTER、CROSS、EMA、SMA）、代码变量名（如 C、H、L、O）或计算公式。
2. **强制图形化翻译**：必须将所有代码逻辑全部翻译为“盘面图形的视觉表现大白话”。例如：不要写“C上穿均线”，要写“当K线实体向上突破参考线”；不要写“FILTER(...,15)”，要写“系统过滤掉了频繁闪烁的无效信号”。

【强制输出排版与色彩规范】
1. 必须保留 Markdown 层级标题。核心心法必须使用 `>` 引用块（以总结性的大白话呈现）。
2. 绝对禁止在全文中使用任何 Emoji 表情符号（坚决删除所有 Emoji）。
3. **强制色彩标注（系统会做拦截处理，请严格按此格式写）**：
   - 负面/预警类（顶部/空头/压力）：必须使用 `:green[词汇]` 或 `:orange[词汇]`。
   - 正面/机会类（底部/多头/支撑）：必须使用 `:red[词汇]`。
   - 强调核心逻辑、公式使用纯文字或 **加粗**。
   - ⚠️ 警告：千万不要把色彩标签和加粗标签嵌套使用！且冒号必须是半角英文字符！
4. **面向小白用户**：通俗易懂，结论前置，然后再用大白话补充解释原因。

【内容结构指令】
1. 引导语：以散户常见的技术观察盲区产生共鸣作为开头。
2. 话题互动：设置一个轻互动话题，引导客观复盘。
3. 知识讲解：采用《指标小课堂》系列形式（如标明：第1期），结合盘面图形表现，深入浅出讲解。
4. **必须生成表格**：生成一个名为“进阶形态与风控应对预案”的 Markdown 表格，严格包含三列：`| 盘面图形特征 | 技术观察含义 | 合规化应对策略建议 |`。
5. 结尾：发起学习打卡或客观技术小结任务，鼓励大家留言互动。"""

def load_indicators():
    if os.path.exists(INDICATORS_FILE):
        try:
            with open(INDICATORS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_indicator(name, code):
    data = load_indicators()
    data[name] = {
        "code": code,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    tmp_file = INDICATORS_FILE + ".tmp"
    with open(tmp_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    os.replace(tmp_file, INDICATORS_FILE)

def delete_indicator(name):
    data = load_indicators()
    if name in data:
        del data[name]
        tmp_file = INDICATORS_FILE + ".tmp"
        with open(tmp_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        os.replace(tmp_file, INDICATORS_FILE)

if "config" not in st.session_state:
    st.session_state.config = load_config()

config = st.session_state.config

def markdown_to_docx_file(md_text, filepath, indicator_name="本指标"):
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    # 标题颜色分级体系
    HEADING_COLORS = {
        1: RGBColor(0, 63, 114),    # 深靛蓝 — 一级标题
        2: RGBColor(0, 82, 148),    # 靛蓝 — 二级标题
        3: RGBColor(34, 107, 156),  # 钢蓝 — 三级标题
    }
    
    doc = Document()
    lines = md_text.split('\n')
    
    in_table = False
    table = None
    in_code_block = False
    prev_was_blank = False
    
    def _render_inline(p, text, is_quote=False, is_th=False):
        # 清洗可能存在的全角冒号及嵌套加粗
        text = text.replace("：green[", ":green[").replace("：red[", ":red[").replace("：orange[", ":orange[")
        # 脱掉颜色标签外层的加粗标记 (例如 **:red[升]** -> :red[升])
        text = re.sub(r'\*\*\s*(:(?:green|red|orange)\[.*?\])\s*\*\*', r'\1', text)
        
        pattern = r'(:(?:green|red|orange)\[.*?\]|\*\*.*?\*\*)'
        parts = re.split(pattern, text)
        for part in parts:
            if not part: continue
            run = p.add_run()
            if is_quote:
                run.bold = True
                run.font.color.rgb = RGBColor(64, 64, 64)
            
            if is_th: run.bold = True
                
            if part.startswith(':green[') and part.endswith(']'):
                run.text = part[7:-1]
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.bold = True
            elif part.startswith(':red[') and part.endswith(']'):
                run.text = part[5:-1]
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
            elif part.startswith(':orange[') and part.endswith(']'):
                run.text = part[8:-1]
                run.font.color.rgb = RGBColor(255, 165, 0)
                run.bold = True
            elif part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            else:
                run.text = part
                
    def _set_heading_color(h, level):
        color = HEADING_COLORS.get(level, RGBColor(0, 82, 148))
        for run in h.runs:
            run.font.color.rgb = color
            run.bold = True
                
    for line in lines:
        stripped = line.strip()
        
        # 处理代码块标记
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            prev_was_blank = False
            continue
        
        # 代码块内容：等宽缩进段落
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
            pf = p.paragraph_format
            pf.left_indent = Pt(24)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            prev_was_blank = False
            continue
        
        # 过滤水平分隔线 --- / *** / ___
        if re.match(r'^[-*_]{3,}$', stripped):
            prev_was_blank = False
            continue
        
        # 空行处理：直接跳过，避免生成无意义的大缝隙空行
        if not stripped:
            if in_table:
                in_table = False
            continue
        
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                row_cells = table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells): _render_inline(row_cells[i].paragraphs[0], cell_text, is_th=True)
            else:
                if all(re.match(r'^[-:\s]+$', c) for c in cells): continue 
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells): _render_inline(row_cells[i].paragraphs[0], cell_text)
            continue
        else:
            if in_table: in_table = False
        
        if stripped.startswith('#### '):
            h = doc.add_heading(level=4)
            _render_inline(h, stripped[5:])
            _set_heading_color(h, 3)
        elif stripped.startswith('### '): 
            h = doc.add_heading(level=3)
            _render_inline(h, stripped[4:])
            _set_heading_color(h, 3)
        elif stripped.startswith('## '): 
            h = doc.add_heading(level=2)
            _render_inline(h, stripped[3:])
            _set_heading_color(h, 2)
        elif stripped.startswith('# '): 
            h = doc.add_heading(level=1)
            _render_inline(h, stripped[2:])
            _set_heading_color(h, 1)
        else:
            is_quote = stripped.startswith('>')
            is_bullet = stripped.startswith('- ') or stripped.startswith('* ')
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                text = stripped[2:]
            else:
                p = doc.add_paragraph()
                if is_quote:
                    text = stripped[1:].strip()
                else:
                    text = stripped
            if text:
                _render_inline(p, text, is_quote)

    # 动态追加美化版的合规免责声明（与指标名称绑定）
    doc.add_paragraph()
    dtbl = doc.add_table(rows=1, cols=1)
    dcell = dtbl.cell(0, 0)
    
    tcPr = dcell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF0F0')
    tcPr.append(shd)
    
    tcBorders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), 'C00000')
    tcBorders.append(left)
    for border_name in ['top', 'right', 'bottom']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    dp1 = dcell.paragraphs[0]
    r1 = dp1.add_run("重要合规化声明：\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(192, 0, 0)
    
    dp2 = dcell.add_paragraph()
    r2 = dp2.add_run(f"以上关于《{indicator_name}》的所有信号提示和区间标注，均是基于历史收盘价等统计数据的技术测算展示，不保证任何未来走势预测的准确性。\n\n")
    r2.font.color.rgb = RGBColor(192, 0, 0)
    
    r3 = dp2.add_run(f"本手册中涉及的所有技术观察标签和计算逻辑仅供技术分析学习与参考，不构成任何形式的投资建议或操作指令。投资者据此操作风险自担，市场有风险，投资需谨慎。请务必结合自身风险承受能力，严格执行止盈止损纪律，独立做出判断。")
    r3.font.color.rgb = RGBColor(192, 0, 0)

    doc.save(filepath)

def generate_jimeng_image(prompt_text, retries=1):
    import re
    import time
    for attempt in range(retries + 1):
        try:
            # Call Dreamina CLI to generate image. --poll=120 will wait up to 120 seconds.
            # Ensure we output JSON or parse text carefully.
            res = subprocess.run([DREAMINA_CMD, "text2image", f"--prompt={prompt_text}", "--poll=120"], capture_output=True, text=True, encoding='utf-8')
            # Expecting output containing URL or similar, since dreamina text2image might output a result message
            # We look for https://... link to the generated image
            # Let's extract URLs from the output
            urls = re.findall(r'https?://[^\s\"\'\)]+', res.stdout)
            img_url = None
            for u in urls:
                if "tos-" in u or "image" in u or ".png" in u or ".jpg" in u or ".jpeg" in u or ".webp" in u:
                    img_url = u
                    break
            if not img_url and urls:
                img_url = urls[-1] # fallback to the last URL found
                
            if img_url:
                # Download the image
                import httpx
                import uuid
                img_resp = httpx.get(img_url, timeout=30)
                if img_resp.status_code == 200:
                    filename = f"jimeng_{uuid.uuid4().hex[:8]}.jpg"
                    filepath = os.path.join(WECHAT_IMAGES_DIR, filename)
                    with open(filepath, "wb") as f:
                        f.write(img_resp.content)
                    return filepath
        except Exception:
            pass
        if attempt < retries:
            time.sleep(2)
    return None

def generate_gemini_image(prompt_text, api_key, model_name="imagen-4.0-generate-001", aspect_ratio="1:1", retries=1):
    import base64
    import time
    import httpx
    import uuid
    
    if not api_key:
        return None
        
    for attempt in range(retries + 1):
        try:
            if "gemini" in model_name:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                payload = {
                    "contents": [{"parts": [{"text": prompt_text}]}],
                    "generationConfig": {
                        "responseModalities": ["IMAGE"]
                    }
                }
                resp = httpx.post(url, headers=headers, json=payload, timeout=60)
                if resp.status_code == 200:
                    resp_json = resp.json()
                    parts = resp_json.get("candidates", [{}])[0].get("content", {}).get("parts", [])
                    img_data = None
                    for p in parts:
                        if "inlineData" in p:
                            b64 = p["inlineData"]["data"]
                            img_data = base64.b64decode(b64)
                            break
                    if img_data:
                        filename = f"gemini_{uuid.uuid4().hex[:8]}.jpg"
                        filepath = os.path.join(WECHAT_IMAGES_DIR, filename)
                        with open(filepath, "wb") as f:
                            f.write(img_data)
                        return filepath
                else:
                    print(f"Gemini Image Gen failed, status code: {resp.status_code}, response: {resp.text}")
            else:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:predict?key={api_key}"
                headers = {"Content-Type": "application/json"}
                payload = {
                    "instances": [{"prompt": prompt_text}],
                    "parameters": {
                        "sampleCount": 1,
                        "aspectRatio": aspect_ratio
                    }
                }
                resp = httpx.post(url, headers=headers, json=payload, timeout=60)
                if resp.status_code == 200:
                    resp_json = resp.json()
                    predictions = resp_json.get("predictions", [])
                    if predictions:
                        img_b64 = predictions[0].get("bytesBase64Encoded")
                        if img_b64:
                            img_data = base64.b64decode(img_b64)
                            filename = f"gemini_{uuid.uuid4().hex[:8]}.jpg"
                            filepath = os.path.join(WECHAT_IMAGES_DIR, filename)
                            with open(filepath, "wb") as f:
                                f.write(img_data)
                            return filepath
                else:
                    print(f"Gemini Image Gen failed, status code: {resp.status_code}, response: {resp.text}")
        except Exception as e:
            print(f"Gemini Image Gen error: {str(e)}")
            
        if attempt < retries:
            time.sleep(2)
    return None

def markdown_to_wechat_docx_bytes(md_text):
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from io import BytesIO
    import re
    import httpx

    doc = Document()
    lines = md_text.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        # 匹配图片 ![alt](url)
        img_match = re.search(r'!\[.*?\]\((.*?)\)', stripped)
        if img_match:
            img_path = img_match.group(1)
            try:
                if img_path.startswith("http"):
                    resp = httpx.get(img_path, timeout=15)
                    if resp.status_code == 200:
                        doc.add_picture(BytesIO(resp.content), width=Inches(6.0))
                else:
                    if os.path.exists(img_path):
                        doc.add_picture(img_path, width=Inches(6.0))
            except Exception:
                pass
            continue
            
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            h = doc.add_heading(level=level)
            h.add_run(stripped[level:].strip())
        elif stripped.startswith('>'):
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:].strip())
            r.font.color.rgb = RGBColor(100, 100, 100)
            r.bold = True
        else:
            p = doc.add_paragraph()
            # 简单处理加粗
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(12)
        
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def render_wechat_preview(md_text):
    import re
    # split text by markdown images
    parts = re.split(r'(!\[.*?\]\(.*?\))', md_text)
    for part in parts:
        if part.startswith("!["):
            m = re.match(r'!\[(.*?)\]\((.*?)\)', part)
            if m:
                img_path = m.group(2)
                if img_path.startswith("http"):
                    st.image(img_path)
                elif os.path.exists(img_path):
                    st.image(img_path)
                else:
                    st.markdown(part)
        else:
            if part.strip():
                st.markdown(part, unsafe_allow_html=True)

def convert_to_wechat_html(md_text, theme, mode_ui, api_key=None, font_size="medium", bg_type="none", chan_config=None, custom_prompt="", for_wechat_api=False):
    import subprocess
    import json
    import re
    import base64
    from openai import OpenAI
    
    input_path = os.path.join("tests", "temp_wechat_input.md")
    output_path = os.path.join("tests", "temp_wechat_output.html")
    
    if os.path.exists(output_path):
        try: os.remove(output_path)
        except: pass
        
    with open(input_path, "w", encoding="utf-8") as f:
        f.write(md_text)
        
    mode = "api" if "API" in mode_ui else "ai"
    
    accent_color = "#d97758"  # 默认暖橙 (autumn-warm)
    if theme == "spring-fresh":
        accent_color = "#6b9b7a"  # 嫩绿
    elif theme == "ocean-calm":
        accent_color = "#4a7c9b"  # 蔚蓝
    
    # 提取首个标题并进行安全长度截断（不超过 32 个字符）
    title = "技术分析报告"
    for line in md_text.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            parsed_title = re.sub(r'^#+\s*', '', line).strip()
            if parsed_title:
                title = parsed_title
                break
    if len(title) > 10:
        title = title[:7] + "..."
        
    cmd = [
        NPX_CMD, "md2wechat", "convert", input_path, 
        "--mode", mode, 
        "--theme", theme, 
        "--preview", 
        "--output", output_path, 
        "--json",
        "--title", title,
        "--author", "AI",
        "--digest", "技术观察与趋势分析"
    ]

    
    if mode == "api" and api_key:
        cmd.extend(["--api-key", api_key])
        cmd.extend(["--font-size", font_size])
        cmd.extend(["--background-type", bg_type])
        
    if mode == "ai" and custom_prompt and theme == "custom":
        cmd.extend(["--custom-prompt", custom_prompt])
        
    # 临时 debug 日志记录
    try:
        with open("tests/wechat_cli_debug.log", "a", encoding="utf-8") as debug_f:
            debug_f.write("=== convert_to_wechat_html CLI CALL ===\n")
            debug_f.write(f"Title: {repr(title)}\n")
            debug_f.write(f"Cmd: {repr(cmd)}\n")
    except:
        pass

    res = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
    
    try:
        with open("tests/wechat_cli_debug.log", "a", encoding="utf-8") as debug_f:
            debug_f.write(f"Returncode: {res.returncode}\n")
            debug_f.write(f"Stdout: {repr(res.stdout)}\n")
            debug_f.write(f"Stderr: {repr(res.stderr)}\n")
    except:
        pass
    
    try:
        m_res = re.search(r'\{.*\}', res.stdout, re.DOTALL)
        if m_res:
            res_data = json.loads(m_res.group(0))
        else:
            res_data = {"success": False, "message": f"无法解析 CLI 返回的 JSON 数据。原始输出: {res.stdout}"}
    except Exception as e:
        res_data = {"success": False, "message": f"解析异常: {str(e)}"}
        
    html_raw = ""
    if mode == "api":
        if not res_data.get("success"):
            err_msg = res_data.get("message", "API 转换失败。")
            debug_info = f"<br><b>[Debug Info]</b><br>Cmd: <code>{re.sub(r'--api-key\s+\S+', '--api-key ******', ' '.join(cmd))}</code><br>Stdout: <pre>{res.stdout}</pre><br>Stderr: <pre>{res.stderr}</pre>"
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 微信排版美化失败 (API 模式)</h3><p>{err_msg}</p>{debug_info}<p>请检查您的 md2wechat API Key 是否正确。</p></div>"
        
        if os.path.exists(output_path):
            with open(output_path, "r", encoding="utf-8") as f:
                html_raw = f.read()
        else:
            return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 转换成功但未生成 HTML 目标文件。</div>"
        
    else: # mode == "ai"
        if not res_data.get("success") or res_data.get("code") != "CONVERT_AI_REQUEST_READY":
            err_msg = res_data.get("message", "AI 模式初始化失败。")
            debug_info = f"<br><b>[Debug Info]</b><br>Cmd: <code>{' '.join(cmd)}</code><br>Stdout: <pre>{res.stdout}</pre><br>Stderr: <pre>{res.stderr}</pre>"
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 微信排版美化失败 (AI 模式初始化)</h3><p>{err_msg}</p>{debug_info}</div>"
            
        ai_prompt = res_data.get("data", {}).get("prompt", "")
        if not ai_prompt:
            return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 转换初始化失败：未获取到 AI 排版 Prompt 提示词。</div>"
            
        if not chan_config or not chan_config.get("api_key"):
            return "<div style='color:red;padding:20px;font-family:sans-serif;'><h3>⚠️ 大模型未配置</h3><p>AI 模式需要使用您的 OpenAI/Gemini 大模型 API Key 进行样式翻译。请在左侧全局配置中填写 API Key！</p></div>"
            
        try:
            def find_anchors_for_image(md_text, target_index):
                lines = md_text.split('\n')
                img_pattern = re.compile(r'!\[.*?\]\(.*?\)')
                
                current_img_idx = 0
                target_line_idx = -1
                
                for idx, line in enumerate(lines):
                    if img_pattern.search(line):
                        if current_img_idx == target_index:
                            target_line_idx = idx
                            break
                        current_img_idx += 1
                        
                if target_line_idx == -1:
                    return None, None
                    
                # 向上寻找非空非图行作为前置锚点
                pre_anchor = None
                for i in range(target_line_idx - 1, -1, -1):
                    line = lines[i].strip()
                    if line and not img_pattern.search(line) and not line.startswith('---'):
                        cleaned = re.sub(r'[#\*_`\-\>\+]', '', line).strip()
                        if len(cleaned) > 2:
                            pre_anchor = cleaned
                            break
                            
                # 向下寻找非空非图行作为后置锚点
                post_anchor = None
                for i in range(target_line_idx + 1, len(lines)):
                    line = lines[i].strip()
                    if line and not img_pattern.search(line) and not line.startswith('---'):
                        cleaned = re.sub(r'[#\*_`\-\>\+]', '', line).strip()
                        if len(cleaned) > 2:
                            post_anchor = cleaned
                            break
                            
                return pre_anchor, post_anchor

            # 强化 ai_prompt 中的图片和排版规则
            if ai_prompt:
                # 0. 规范化并明确图片索引从 0 开始计数，防止部分主题（如 spring-fresh）因为缺少示例导致 LLM 从 1 开始索引
                ai_prompt = ai_prompt.replace(
                    "图片使用占位符格式：<!-- IMG:index -->，例如第一张图用 <!-- IMG:0 -->",
                    "图片使用占位符格式：<!-- IMG:index -->"
                )
                ai_prompt = ai_prompt.replace(
                    "图片使用占位符格式：<!-- IMG:index -->",
                    "图片使用占位符格式：<!-- IMG:index -->，其中索引从 0 开始（即第一张图用 <!-- IMG:0 -->，第二张图用 <!-- IMG:1 -->，依此类推）"
                )

                # 1. 物理剥离 ![说明文字](<!-- IMG:index -->) 的 alt 说明文字，防止生成配图卡片，仅保留无文本图片节点
                ai_prompt = re.sub(
                    r'!\[.*?\]\(<!--\s*IMG:(\d+)\s*-->\)', 
                    r'<div class="wechat-img-placeholder" data-index="\1"></div>', 
                    ai_prompt
                )
                
                # 2. 替换漏网的 <!-- IMG:index --> 为自定义标签
                ai_prompt = re.sub(
                    r'<!--\s*IMG:(\d+)\s*-->', 
                    r'<div class="wechat-img-placeholder" data-index="\1"></div>', 
                    ai_prompt
                )

                # 根据不同主题自适应主标题的强调色、文字色以及悬浮卡片样式
                accent_color = "#d97758"  # 默认暖橙 (autumn-warm)
                eyebrow_color = "#8a7e72" # 默认灰褐
                title_color = "#4a413d"   # 默认黑褐
                card_style = "background-color: #ffffff; border: 1px solid rgba(0, 0, 0, 0.05); box-shadow: 0 10px 30px rgba(0, 0, 0, 0.04), 0 0 15px rgba(217, 119, 88, 0.1); border-radius: 18px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"
                
                if theme == "spring-fresh":
                    accent_color = "#6b9b7a"  # 嫩绿
                    eyebrow_color = "#5c735c" # 绿灰
                    title_color = "#3d4a3d"   # 深绿灰
                    card_style = "background-color: #ffffff; border: 1px solid rgba(107, 155, 122, 0.1); box-shadow: 0 8px 24px rgba(74, 128, 88, 0.08), 0 0 12px rgba(107, 155, 122, 0.1); border-radius: 16px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"
                elif theme == "ocean-calm":
                    accent_color = "#4a7c9b"  # 蔚蓝
                    eyebrow_color = "#556e80" # 石蓝
                    title_color = "#3a4150"   # 深蓝灰
                    card_style = "background-color: #ffffff; border: 1px solid rgba(74, 124, 155, 0.08); box-shadow: 0 8px 28px rgba(58, 65, 80, 0.06), 0 0 16px rgba(74, 124, 155, 0.1); border-radius: 14px; padding: 25px; margin-bottom: 30px; display: block; box-sizing: border-box; width: 100%;"

                image_rule_enhancement = f"""
【重要：主标题 (H1) 固定排版规范（核心必做）】
1. 一级大标题（H1，即 `# 标题内容`）代表文章的主标题。你必须将其设计成一个极其典雅、浮动卡片式（Floating Card Layout）的头部区域（Header Region），位于网页最顶部。
2. 你必须将主标题 H1 的所有内容，包裹在一个独立的白色卡片容器中，其内联样式（style）必须严格设为：`{card_style}`，使其在视觉上呈现出与下文卡片相同的“悬浮卡片”质感。
3. 卡片容器内的元素必须严格按照以下顺序排列，并保证完美居中对齐：
   - 顶部分类眉标（Eyebrow）：在标题正上方放置一行小字，内容为“行业深度解析 | INDUSTRY ANALYSIS”，样式为：`font-size: 13px; color: {eyebrow_color}; letter-spacing: 2px; text-transform: uppercase; margin-bottom: 12px; font-weight: bold; text-align: center;`，居中。
   - 中间大标题：在眉标下方。使用 `<h1>` 标签，样式为：`font-size: 26px; line-height: 1.4; color: {title_color}; font-weight: bold; margin: 10px 0; font-family: 'Georgia', 'Source Serif Pro', serif; text-align: center;`，居中对齐。
   - 底部装饰短线：在标题正下方，居中放置一个装饰短横线：`<div style="margin: 15px auto 5px auto; width: 50px; height: 2px; background-color: {accent_color};"></div>`，用于与标题外部的其他正文卡片进行优雅的视觉区隔。
4. 确保 H1 卡片的宽度自适应，且其外层没有任何其他的嵌套卡片容器包裹。

【重要：图片与防溢出排版补充要求】
1. 在转换内容中，你必须原样保留图片占位符标签 `<div class="wechat-img-placeholder" data-index="index"></div>` 并在生成的 HTML 对应位置呈现出来。
2. 绝对禁止修改、删除或包裹此标签属性（如 class 或 data-index），严禁在此标签内写入任何子节点、文字、空格或 HTML 注释。它必须作为单独一行的独立块级节点。
3. 绝对禁止为图片生成任何多余的配图说明卡片或提示性文字（例如不要输出类似“【此处是手机芯片图片】”或说明该图用意、图片标题的字样）。
4. 最外层主背景 <div> 容器及所有卡片必须使用自适应宽度（如 `width: 100%;`），并指定 `box-sizing: border-box;`。
5. 所有一级、二级、三级标题以及段落等，如果使用了带边框或背景颜色的卡片，请确保内容不超出卡片的视觉边缘。
"""
                # 将补充规则插入到 md 转换内容之前
                if "请转换以下 Markdown内容：" in ai_prompt:
                    ai_prompt = ai_prompt.replace("请转换以下 Markdown内容：", image_rule_enhancement + "\n\n请转换以下 Markdown内容：")
                else:
                    ai_prompt = ai_prompt + "\n\n" + image_rule_enhancement

            client = OpenAI(api_key=chan_config.get("api_key"), base_url=chan_config.get("base_url") if chan_config.get("base_url") else None)
            response = call_chat_completion(
                client,
                chan_config.get("selected_model"),
                [
                    {"role": "system", "content": (
                        "You are a professional WeChat push designer. You must return only a valid, complete HTML page. "
                        "Strictly do not wrap it in any comments or explanations, just raw HTML or markdown-fenced HTML code blocks.\n"
                        "【排版设计美化强力指示（至关重要）】:\n"
                        "1. 必须使用丰富的内联样式（inline CSS）为推文进行色彩与板块美化设计。\n"
                        "2. 必须包含大量精心美化的组件与样式元素，例如：\n"
                        "   - 带有柔和浅色背景色（如浅橘、淡绿、浅蓝等）、圆角（border-radius: 8px或12px）和阴影（box-shadow: 0 4px 12px rgba(0,0,0,0.05)）的卡片式展示盒子（cards）。\n"
                        "   - 带有左侧粗色条修饰（border-left: 4px solid ...）和淡背景色的引用块（blockquote）。\n"
                        "   - 带有特色标签、图标或彩色背景框的列表项（li）和表格（table）。\n"
                        "   - 加粗并以主题主色高亮的关键字或句段。\n"
                        "3. 第一个一级大标题（H1）必须设计得极具视觉冲击力，可使用较大字号、居中对齐、彩色背景横带或精致下划线装饰。\n"
                        "4. 绝不能输出素面朝天、简陋或完全没有美化样式的纯文本 HTML！\n"
                        "5. 绝对禁止以任何形式增删、改写或缩减原文的段落或字词内容！你的职责仅仅是对文章做 HTML 排版和样式美化包装，必须保留全部文字细节与配图占位符。\n"
                        "6. 严格禁止在排版生成的 HTML 中出现任何 Emoji 表情符号。如果检测到原文中带有 Emoji，请将其滤除或用等价的文字描述替代。"
                    )},
                    {"role": "user", "content": ai_prompt}
                ],
                chan_config=chan_config
            )
            
            if not response.choices or not response.choices[0].message.content:
                return "<div style='color:red;padding:20px;font-family:sans-serif;'>❌ 大模型未返回 HTML 排版内容。</div>"
                
            html_raw = response.choices[0].message.content.strip()
            
            if "```html" in html_raw:
                html_raw = html_raw.split("```html")[1].split("```")[0].strip()
            elif "```" in html_raw:
                html_raw = html_raw.split("```")[1].split("```")[0].strip()
                
            # Replace image placeholders
            images_list = res_data.get("data", {}).get("images") or []
            replaced_indices = set()
            
            # 首先遍历图片列表进行占位符替换
            for idx, img in enumerate(images_list):
                index = img.get("Index") if "Index" in img else img.get("index", idx)
                original = img.get("Original") if "Original" in img else img.get("original", "")
                
                src = original
                if not for_wechat_api and src and not src.startswith("http") and os.path.exists(src):
                    try:
                        with open(src, "rb") as f_img:
                            img_b64 = base64.b64encode(f_img.read()).decode()
                        ext = os.path.splitext(src)[1].lower().replace(".", "")
                        if ext == "jpg": ext = "jpeg"
                        src = f"data:image/{ext};base64,{img_b64}"
                    except:
                        pass
                
                img_tag = f'<img src="{src}" style="width: 100%; max-width: 100%; border-radius: 8px; margin: 16px 0; display: block; height: auto;" />'
                
                # 匹配类似 <div class="wechat-img-placeholder" data-index="0"></div> 的结构，允许颠倒顺序和自闭合
                placeholder_regex = re.compile(
                    rf'<div\s+[^>]*(?:class=["\']wechat-img-placeholder["\']|data-index=["\']{index}["\'])[^>]+(?:class=["\']wechat-img-placeholder["\']|data-index=["\']{index}["\'])[^>]*>(?:\s*</div>)?',
                    re.IGNORECASE
                )
                
                if placeholder_regex.search(html_raw):
                    html_raw = placeholder_regex.sub(img_tag, html_raw)
                    replaced_indices.add(index)
                    
            # 其次，针对大模型丢失占位符的情况，启动双向段落锚点兜底匹配强插机制
            for idx, img in enumerate(images_list):
                index = img.get("Index") if "Index" in img else img.get("index", idx)
                if index in replaced_indices:
                    continue
                    
                original = img.get("Original") if "Original" in img else img.get("original", "")
                src = original
                if not for_wechat_api and src and not src.startswith("http") and os.path.exists(src):
                    try:
                        with open(src, "rb") as f_img:
                            img_b64 = base64.b64encode(f_img.read()).decode()
                        ext = os.path.splitext(src)[1].lower().replace(".", "")
                        if ext == "jpg": ext = "jpeg"
                        src = f"data:image/{ext};base64,{img_b64}"
                    except:
                        pass
                        
                img_tag = f'<img src="{src}" style="width: 100%; max-width: 100%; border-radius: 8px; margin: 16px 0; display: block; height: auto;" />'
                
                pre_anchor, post_anchor = find_anchors_for_image(md_text, index)
                inserted = False
                
                if pre_anchor:
                    pre_clean = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', pre_anchor).strip()
                    pos = html_raw.find(pre_clean[:25]) if len(pre_clean) > 25 else html_raw.find(pre_clean)
                    if pos != -1:
                        end_tag_pos = html_raw.find('>', pos)
                        if end_tag_pos != -1:
                            html_raw = html_raw[:end_tag_pos + 1] + "\n" + img_tag + "\n" + html_raw[end_tag_pos + 1:]
                            inserted = True
                            
                if not inserted and post_anchor:
                    post_clean = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', post_anchor).strip()
                    pos = html_raw.find(post_clean[:25]) if len(post_clean) > 25 else html_raw.find(post_clean)
                    if pos != -1:
                        start_tag_pos = html_raw.rfind('<', 0, pos)
                        if start_tag_pos != -1:
                            html_raw = html_raw[:start_tag_pos] + "\n" + img_tag + "\n" + html_raw[start_tag_pos:]
                            inserted = True
                            
                if not inserted:
                    if "</div>" in html_raw:
                        last_div = html_raw.rfind("</div>")
                        html_raw = html_raw[:last_div] + "\n" + img_tag + "\n" + html_raw[last_div:]
                    else:
                        html_raw = html_raw + "\n" + img_tag
                        
        except Exception as e:
            return f"<div style='color:red;padding:20px;font-family:sans-serif;'><h3>❌ 大模型调用渲染失败</h3><p>{str(e)}</p></div>"

    # For both API and AI modes, convert any local image paths in src="..." to base64 data URIs
    # to bypass Streamlit iframe sandbox restrictions
    if not for_wechat_api:
        def replace_src(match):
            prefix = match.group(1)
            src = match.group(2)
            if not src.startswith("http") and not src.startswith("data:") and os.path.exists(src):
                try:
                    with open(src, "rb") as f_f:
                        b64 = base64.b64encode(f_f.read()).decode()
                    ext = os.path.splitext(src)[1].lower().replace(".", "")
                    if ext == "jpg": ext = "jpeg"
                    return f'{prefix}="data:image/{ext};base64,{b64}"'
                except:
                    pass
            return match.group(0)

        html_raw = re.sub(r'(src)=["\']([^"\']+)["\']', replace_src, html_raw, flags=re.IGNORECASE)

    # 注入全局兼容性与防溢出样式，解决 450px viewport 长图导出及微信预览时发生的内容拉宽、背景托不住内容的问题
    compat_css = """
<style id="wechat-compat-styles">
  /* 强制全局 box-sizing 以便正确计算内边距 */
  *, *:before, *:after {
    box-sizing: border-box !important;
  }
  /* 限制最大宽度，防止子元素溢出主背景 */
  html, body {
    margin: 0;
    padding: 0;
    width: 100% !important;
    max-width: 100% !important;
    overflow-x: hidden !important;
  }
  /* 特别防止各种块级组件和图片溢出 */
  section, div, p, span, img, table, pre, code {
    max-width: 100% !important;
    word-wrap: break-word !important;
    word-break: break-word !important;
  }
  /* 代码块和表格在太宽时允许横向滚动，不可撑破父容器 */
  pre, code, table {
    overflow-x: auto !important;
    white-space: pre-wrap !important;
  }
  /* 图片自适应限制 */
  img {
    max-width: 100% !important;
    height: auto !important;
    display: block !important;
  }
</style>
"""
    if "</head>" in html_raw:
        html_raw = html_raw.replace("</head>", f"{compat_css}</head>")
    elif "<body>" in html_raw:
        html_raw = html_raw.replace("<body>", f"<body>{compat_css}")
    else:
        html_raw = compat_css + html_raw
        
    # 物理删除所有 Emoji 表情符号
    emoji_pattern = re.compile(r'[\U00010000-\U0010ffff]', flags=re.UNICODE)
    html_raw = emoji_pattern.sub('', html_raw)
    
    # 重点强调 (strong/b) 字色强制高亮为主题强调色
    def highlight_strong(match):
        tag = match.group(1)
        attrs = match.group(2) or ""
        content = match.group(3)
        if 'style=' in attrs or 'style =' in attrs:
            def repl_style(m):
                style_content = m.group(2)
                # 滤除原先可能带有的任何 color 声明
                style_content = re.sub(r'color\s*:\s*[^;]+;?', '', style_content).strip()
                # 拼装新的主题色
                return f'style="{style_content}; color: {accent_color}; font-weight: bold;"'
            new_attrs = re.sub(r'style\s*=\s*(["\'])(.*?)(\1)', repl_style, attrs)
        else:
            new_attrs = attrs + f' style="color: {accent_color}; font-weight: bold;"'
        return f'<{tag}{new_attrs}>{content}</{tag}>'

    html_raw = re.sub(r'<(strong|b)(\s+[^>]*?)?>(.*?)</\1>', highlight_strong, html_raw, flags=re.IGNORECASE)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_raw)

    return html_raw

def generate_wechat_long_image(html_content):
    import os
    from datetime import datetime
    from playwright.sync_api import sync_playwright
    
    output_filename = f"wechat_long_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
    output_path = os.path.join(WECHAT_IMAGES_DIR, output_filename)
    
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={"width": 450, "height": 800})
        page.set_content(html_content)
        page.wait_for_load_state("networkidle")
        page.wait_for_timeout(1000)
        page.screenshot(path=output_path, full_page=True)
        browser.close()
        
    return output_path

# Default Config Setup
defaults = {
    "user_groups": {},
    "groups": {"默认群组": ""},
    "platform": "自定义/OpenAI",
    "mode": "常规总结",
    "selected_group": "默认群组",
    "auto_run": False,
    "run_time": "08:00",
    "channel_configs": {
        "自定义/OpenAI": {"api_key": "", "base_url": "https://api.openai.com/v1", "selected_model": "gpt-4o", "available_models": ["gpt-4o", "gpt-3.5-turbo"]},
        "火山方舟 (Volcengine)": {"api_key": "", "base_url": "https://ark.cn-beijing.volces.com/api/v3", "selected_model": "", "available_models": []},
        "魔塔 (ModelScope)": {"api_key": "", "base_url": "https://api.modelscope.cn/v1", "selected_model": "", "available_models": []},
        "DeepSeek": {"api_key": "", "base_url": "https://api.deepseek.com", "selected_model": "deepseek-chat", "available_models": ["deepseek-chat", "deepseek-reasoner"], "enable_thinking": True, "reasoning_effort": "high"}
    },
    "custom_prompts": {"合规化处理": "请对以上内容进行合规化处理：1. 隐藏具体的人名和联系方式；2. 增加‘以上内容仅供参考，不构成投资建议’的免责声明；3. 语气调整为客观中立的中台视角。"}
}
for key, val in defaults.items():
    if key not in config:
        config[key] = val

# Ensure all default channels exist in channel_configs
if "channel_configs" not in config:
    config["channel_configs"] = defaults["channel_configs"].copy()
else:
    for plat, plat_cfg in defaults["channel_configs"].items():
        if plat not in config["channel_configs"]:
            config["channel_configs"][plat] = plat_cfg.copy()

# Migrate old config to channel_configs specifically, and clean up polluted platforms
if "api_key" in config and config["api_key"]:
    legacy_key = config.get("api_key", "")
    legacy_url = config.get("base_url", "")
    legacy_model = config.get("selected_model", "")
    legacy_models = config.get("available_models", [])
    
    # Check if the legacy URL belongs to ModelScope
    is_modelscope = "modelscope.cn" in legacy_url
    target_plat = "魔塔 (ModelScope)" if is_modelscope else "自定义/OpenAI"
    
    # Migrate Specifically to target_plat if it doesn't already have an api_key
    target_cfg = config["channel_configs"].setdefault(target_plat, {})
    if not target_cfg.get("api_key"):
        target_cfg["api_key"] = legacy_key
        target_cfg["base_url"] = legacy_url if legacy_url else defaults["channel_configs"][target_plat]["base_url"]
        target_cfg["selected_model"] = legacy_model if legacy_model else defaults["channel_configs"][target_plat]["selected_model"]
        target_cfg["available_models"] = legacy_models if legacy_models else defaults["channel_configs"][target_plat]["available_models"]
        
    # Reset other channels if they were polluted by legacy_key
    for plat, default_cfg in defaults["channel_configs"].items():
        if plat in config["channel_configs"] and plat != target_plat:
            chan_cfg = config["channel_configs"][plat]
            if chan_cfg.get("api_key") == legacy_key:
                config["channel_configs"][plat] = default_cfg.copy()
                
    # Remove all legacy root keys
    for k in ["api_key", "base_url", "selected_model", "available_models"]:
        config.pop(k, None)
    save_config(config)

if "available_models" not in st.session_state:
    st.session_state.available_models = config["channel_configs"][config.get("platform", "自定义/OpenAI")].get("available_models", [])

@st.cache_resource
def get_scheduler():
    sched = BackgroundScheduler()
    sched.start()
    return sched

if "scheduler" not in st.session_state:
    st.session_state.scheduler = get_scheduler()

if "virtual_history" not in st.session_state:
    st.session_state.virtual_history = []

def adjust_markdown_images_placement(md_text):
    import re
    lines = md_text.split('\n')
    sections = []
    current_section = []
    
    for line in lines:
        if line.strip().startswith('#'):
            if current_section:
                sections.append(current_section)
            current_section = [line]
        else:
            current_section.append(line)
    if current_section:
        sections.append(current_section)
        
    new_sections = []
    for sec in sections:
        if not sec:
            continue
        if sec[0].strip().startswith('#'):
            heading = sec[0]
            content_lines = sec[1:]
            
            img_lines = []
            other_lines = []
            img_pattern = re.compile(r'!\[.*?\]\((?:\[IMAGE_GENERATE:.*?\]|.*?)\)')
            
            for line in content_lines:
                if img_pattern.search(line):
                    img_lines.append(line)
                else:
                    other_lines.append(line)
            
            if img_lines:
                new_sec = [heading]
                for img in img_lines:
                    new_sec.append(img)
                while other_lines and not other_lines[0].strip():
                    other_lines.pop(0)
                new_sec.extend(other_lines)
                new_sections.append(new_sec)
            else:
                new_sections.append(sec)
        else:
                        new_sections.append(sec)
            
    flat_lines = []
    for sec in new_sections:
        flat_lines.extend(sec)
    return '\n'.join(flat_lines)

def get_wechat_system_prompt():
    return f"""你是一个顶级的科技/金融类微信公众号爆款作者。
请根据以下 AI 深度分析的结论，将其重写为一篇适合微信公众号发布的推文。
【核心要求】：
【时间感知】：当前系统物理时间是 {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}，请在撰写内容时以此为基准判断时效性与时态表述。
1. 结构化排版：最开头必须包含吸睛的一级大标题（格式为：`# 标题内容`，必须使用单个 `#` 号表示一级标题）、引言、逻辑清晰的分层正文。文章最后必须包含正式的总结部分，且总结部分必须以二级标题开始（格式为：`## 写在最后：[富有深度或代入感的总结标题]`），以生成独立的悬浮卡片样式。
2. 语言风格与排版美化：
   - 严禁 Emoji 表情：严格禁止在文章的任何地方（包括大标题、分标题、正文、免责声明等）使用任何 Emoji 表情符号（如 🚀, 🔥, 📊, 💡, ⚠️ 等）。必须通过纯文字以及段落本身的条理性和文采来吸引读者。
   - 提升可读性：段落宜短不宜长，每段最多包含 3-4 行文字，段落之间要多留白。
   - 多使用列表：多采用无序列表（- 列表项）或有序列表（1. 列表项）对复杂要点和行业逻辑进行条理清晰的归纳，以便读者在手机屏幕上快速扫读。
   - 重点加粗：对于行业洞察中的关键词、核心结论句段，应使用双星号加粗（如 `**重点内容**`），这将在后续被渲染为亮眼的主题强调色。
   - 语言应通俗易懂，大白话，有极强的情绪价值 and 代入感（就像和朋友面对面聊天）。如果用户提供了个性化要求，必须优先满足。
3. 金融合规红线（必须绝对遵守）：
   - 严禁对第三方公司进行主观点评或说他好坏。只描述客观物理现实、事实逻辑和上下游合作关系，禁止带有任何主观褒贬色彩。
   - 严禁点名个股股份表现描述/预测：绝对禁止直接点名某只具体个股并搭配股价表现描述（如大涨、大跌、暴涨、崩盘、突破新高）或对其未来价格进行预测与判断。一旦出现“个股名称/股票代码 + 股价表现/走势判断”的组合即视为违规。
   - 使用行业/板块代替个股描述：若现实中某只股票出现大幅波动，请将其抽象为对该股票所处的行业或板块的整体走势进行描述，并使用相近的合规词汇。例如：将“中芯国际股价大涨”替换为“半导体制造板块集体走强”；将“中际旭创暴跌”替换为“光通信板块走势趋弱”；将“大洋电机暴涨”替换为“微电机板块表现活跃”。
4. 智能配图（核心硬性红线）：每一篇文章都必须做到图文并茂。你必须为你撰写出的【每一个标题】（包括文章的大标题 H1，以及每一个类似“## 一、”、“## 二、”的二级标题 H2）的正下方、第一句正文内容之前，立即插入且仅插入 1 张相关配图。
   - 严禁遗漏任何标题！如果文章包含 1 个大标题和 4 个分级小标题，你必须输出恰好 5 张配图。少图即为违规。
   - 图片位置：必须在标题的正下方、第一句正文内容之前。严禁将图片放在段落结尾、小节末尾或内容下方。
   - 插入图片 Markdown 格式严格为：![图片说明]([IMAGE_GENERATE:即梦大模型高质量英文提示词])。
   - ⚠️ 警告：中括号内必须填入具体的、具象的**英文视觉描述词**（绝对不能用抽象词汇），侧重于光影、构图、写实摄影或精美插画，例如 ![原油大跌]([IMAGE_GENERATE:A realistic photography of oil fields in the Middle East with dramatic lighting, cinematic, 8k resolution])。
5. 结尾固定免责声明：在文章的最后，必须一字不差地加上以下免责声明：
> 本文内容及数据均基于公开市场资料与行业研报，仅作产业趋势分析与逻辑梳理之用，旨在探讨技术发展方向与产业格局变迁，不构成任何具体的投资建议或操作指引。文中提及的企业及产品仅作为产业案例分析，不构成推荐。投资有风险，入市需谨慎。请您基于自身独立判断做出决策。"""

def write_cron_log(msg):
    log_path = os.path.join("outputs", "cron_execution.log")
    os.makedirs(os.path.dirname(log_path), exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {msg}\n")

def run_scheduled_wechat_publish(draft_file_path: str):
    """定时群发公众号推文的核心执行逻辑（被 APScheduler 异步调用）"""
    import os
    import json
    import datetime
    import traceback
    
    write_cron_log(f"==================================================")
    write_cron_log(f"[定时群发公众号] 定时群发任务启动，关联文件: {draft_file_path}")
    
    if not os.path.exists(draft_file_path):
        write_cron_log(f"[定时群发公众号] 错误：找不到草稿关联 JSON 文件 {draft_file_path}")
        return
        
    try:
        with open(draft_file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        media_id = data.get("media_id")
        appid = data.get("appid")
        secret = data.get("secret")
        url = data.get("url")
        status = data.get("status")
        
        if not media_id or not appid or not secret:
            write_cron_log(f"[定时群发公众号] 错误：JSON 数据不全 (media_id: {media_id}, appid: {appid})")
            return
            
        if status == "published":
            write_cron_log(f"[定时群发公众号] 提示：该文章此前已发布，跳过本次群发。")
            return
            
        # 1. 获取微信凭证
        write_cron_log(f"[定时群发公众号] 正在获取 Access Token (AppID: {appid[:6]}...)")
        token = wechat_publisher.get_access_token(appid, secret)
        
        # 2. 正式发布
        write_cron_log(f"[定时群发公众号] 正在正式发布 MediaID: {media_id} ...")
        pub_id = wechat_publisher.publish_draft(token, media_id)
        
        # 3. 更新状态
        data["status"] = "published"
        data["publish_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(draft_file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
            
        write_cron_log(f"[定时群发公众号] 成功：已正式发布！发布ID: {pub_id}")
        
    except Exception as e:
        write_cron_log(f"[定时群发公众号] 异常崩溃：{str(e)}\n{traceback.format_exc()}")

def run_scheduled_deep_analysis():
    cfg = load_config()
    da_sched = cfg.get("schedulers", {}).get("AI 深度分析", {})
    if not da_sched.get("auto_run", False): return
    ui_state = da_sched.get("ui_state", {})
    if not ui_state:
        write_cron_log("[AI 深度分析] 未找到保存的配置，跳过执行。")
        return
        
    write_cron_log("==================================================")
    write_cron_log("[AI 深度分析] 定时任务启动")
    try:
        status_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8')
        user_id = ""
        m = re.search(r'\{.*\}', status_res.stdout, re.DOTALL)
        if m:
            auth_data = json.loads(m.group(0))
            if auth_data.get("ok") and auth_data.get("data", {}).get("loggedIn"):
                user_id = auth_data["data"].get("userId", "")
        if not user_id:
            write_cron_log("[AI 深度分析] 错误: 知识星球未登录/未授权，任务终止。")
            return
            
        selected_group = ui_state.get("selected_group", cfg.get("selected_group", "默认群组"))
        group_id = cfg.get("user_groups", {}).get(user_id, {}).get(selected_group)
        if not group_id:
            write_cron_log(f"[AI 深度分析] 错误: 未找到选中群组 '{selected_group}' 的 ID，任务终止。")
            return
            
        scope_ui = ui_state.get("scope_ui", "最新总结 (话题+文件)")
        s_key = "all" if "最新" in scope_ui else "files"
        l_limit = ui_state.get("l_limit", 3)
        a_mode = ui_state.get("a_mode", "常规总结")
        
        write_cron_log(f"[AI 深度分析] 正在获取知识星球动态 (群组: {selected_group}, 范围: {scope_ui}, 数量: {l_limit})...")
        raw, _, _ = fetch_zsxq(group_id, limit=l_limit, scope=s_key)
        if raw.startswith("获取失败") or raw.startswith("获取异常"):
            write_cron_log(f"[AI 深度分析] 知识星球数据获取失败: {raw}")
            return
            
        plat = cfg.get("platform", "自定义/OpenAI")
        plat_cfg = cfg.get("channel_configs", {}).get(plat, {})
        selected_model = plat_cfg.get("selected_model", "")
        
        write_cron_log(f"[AI 深度分析] 正在向大模型发起深度分析 (平台: {plat}, 模型: {selected_model}, 模式: {a_mode})...")
        res_main = generate_summary(raw, plat_cfg.get("api_key"), plat_cfg.get("base_url"), selected_model, a_mode, chan_config=plat_cfg)
        if res_main.startswith("AI 总结失败") or res_main.startswith("未提供"):
            write_cron_log(f"[AI 深度分析] 分析失败: {res_main}")
            return
            
        final_res = res_main
        use_p_ui = ui_state.get("use_p_ui", False)
        cp_text = ui_state.get("cp_text", "")
        if use_p_ui and cp_text:
            write_cron_log("[AI 深度分析] 正在执行个性化二次加工...")
            final_res = generate_summary(res_main, plat_cfg.get("api_key"), plat_cfg.get("base_url"), selected_model, a_mode, custom_prompt=cp_text, chan_config=plat_cfg)
            if final_res.startswith("AI 总结失败"):
                write_cron_log(f"[AI 深度分析] 二次加工失败: {final_res}")
                return
                
        use_wechat = ui_state.get("use_wechat", False)
        also_generate_report = ui_state.get("also_generate_report", False)
        
        if use_wechat:
            write_cron_log("[AI 深度分析] 正在生成微信公众号推文...")
            wechat_system_prompt = get_wechat_system_prompt()
            wechat_prompt = ui_state.get("wechat_prompt", "")
            wechat_user_content = f"【基础分析总结】\n{final_res}\n"
            if wechat_prompt.strip():
                wechat_user_content += f"\n【用户个性化要求】\n{wechat_prompt}"
                
            from openai import OpenAI
            client = OpenAI(api_key=plat_cfg.get("api_key"), base_url=plat_cfg.get("base_url") if plat_cfg.get("base_url") else None)
            wc_response = call_chat_completion(client, selected_model, [{"role": "system", "content": wechat_system_prompt}, {"role": "user", "content": wechat_user_content}], chan_config=plat_cfg)
            if wc_response.choices and wc_response.choices[0].message.content:
                raw_wechat = wc_response.choices[0].message.content
                raw_wechat = adjust_markdown_images_placement(raw_wechat)
                def replace_img(match):
                    kw = match.group(1)
                    image_engine = cfg.get("image_generator", "即梦 (Dreamina)")
                    if image_engine == "Google Gemini (Imagen 3)":
                        write_cron_log(f"[AI 深度分析] 🖼️ 正在调用 Gemini 生成配图: `{kw}`")
                        img_path = generate_gemini_image(kw, cfg.get("google_api_key", ""), model_name=cfg.get("gemini_image_model", "imagen-4.0-generate-001"), aspect_ratio=cfg.get("image_aspect_ratio", "1:1"))
                    else:
                        write_cron_log(f"[AI 深度分析] 🖼️ 正在调用即梦生成配图: `{kw}`")
                        img_path = generate_jimeng_image(kw)
                    if img_path:
                        img_path = img_path.replace("\\", "/")
                        return f"({img_path})"
                    return "(https://dummyimage.com/800x400/ffebee/d32f2f.png&text=Image+Generate+Failed)"
                wechat_res = re.sub(r'\(\[IMAGE_GENERATE:(.*?)\]\)', replace_img, raw_wechat)
                
                wc_filename = f"wechat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                wc_path = os.path.join(WECHAT_OUTPUT_DIR, wc_filename)
                with open(wc_path, "w", encoding="utf-8") as f:
                    f.write(wechat_res)
                    
                write_cron_log("[AI 深度分析] 🎨 正在使用 md2wechat 进行样式美化排版...")
                theme = ui_state.get("wechat_theme", cfg.get("wechat_theme", "spring-fresh"))
                mode_ui = ui_state.get("wechat_mode", cfg.get("wechat_mode", "AI 模式 (免费)"))
                api_key = cfg.get("md2wechat_api_key", "")
                font_size = cfg.get("wechat_font_size", "medium")
                bg_type = cfg.get("wechat_background_type", "none")
                custom_prompt = cfg.get("wechat_custom_prompt", "")
                
                html_res = convert_to_wechat_html(wechat_res, theme, mode_ui, api_key=api_key, font_size=font_size, bg_type=bg_type, chan_config=plat_cfg, custom_prompt=custom_prompt)
                html_filename = wc_filename.replace(".md", ".html")
                html_path = os.path.join(WECHAT_OUTPUT_DIR, html_filename)
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html_res)
                write_cron_log(f"[AI 深度分析] 微信推文生成成功: {html_path}")
                
                # 定时任务微信自动发布逻辑
                pub_mode = ui_state.get("wechat_publish_mode", "仅生成本地文件 (不上传)")
                if pub_mode != "仅生成本地文件 (不上传)":
                    write_cron_log(f"[AI 深度分析] 触发定时微信自动发布，发布模式: {pub_mode}")
                    
                    accounts = wechat_publisher.load_accounts()
                    selected_acc_name = ui_state.get("wechat_publish_account", "")
                    active_account = next((a for a in accounts if a["name"] == selected_acc_name), None)
                    if not active_account and accounts:
                        active_account = accounts[0]
                        write_cron_log(f"[AI 深度分析] 未找到指定微信账号 '{selected_acc_name}'，默认使用第一个账号 '{active_account['name']}'")
                        
                    if active_account:
                        try:
                            pub_res = wechat_publisher.auto_publish_to_wechat(
                                wc_path, 
                                html_path, 
                                active_account, 
                                pub_mode
                            )
                            write_cron_log(f"[AI 深度分析] 微信自动发布成功！MediaID: {pub_res.get('media_id')}")
                            if pub_res.get("publish_id"):
                                write_cron_log(f"[AI 深度分析] 微信已提交正式发布，PublishID: {pub_res.get('publish_id')}")
                        except Exception as pub_err:
                            write_cron_log(f"[AI 深度分析] 微信自动发布失败 (不重试): {str(pub_err)}")
                    else:
                        write_cron_log("[AI 深度分析] 错误: 微信自动发布失败，未配置任何微信公众号账号")
            else:
                write_cron_log("[AI 深度分析] 警告: 微信推文生成返回内容为空。")
                
        if not use_wechat or also_generate_report:
            write_cron_log("[AI 深度分析] 正在渲染专业分析报表...")
            img_p = render_to_image(final_res, a_mode)
            write_cron_log(f"[AI 深度分析] 分析报表已生成并保存到: {img_p}")
        write_cron_log("[AI 深度分析] 定时任务执行完毕。")
    except Exception as e:
        import traceback
        write_cron_log(f"[AI 深度分析] 异常终止: {str(e)}\n{traceback.format_exc()}")

def run_scheduled_video_script():
    cfg = load_config()
    vs_sched = cfg.get("schedulers", {}).get("视频脚本制作器", {})
    if not vs_sched.get("auto_run", False): return
    ui_state = vs_sched.get("ui_state", {})
    if not ui_state:
        write_cron_log("[视频脚本制作器] 未找到保存的配置，跳过执行。")
        return
        
    write_cron_log("==================================================")
    write_cron_log("[视频脚本制作器] 定时任务启动")
    try:
        plat = cfg.get("platform", "自定义/OpenAI")
        plat_cfg = cfg.get("channel_configs", {}).get(plat, {})
        selected_model = plat_cfg.get("selected_model", "")
        
        prompt_input = ui_state.get("prompt_input", "")
        virtual_history = ui_state.get("virtual_history", [])
        
        history_texts = []
        for vf in virtual_history:
            history_texts.append(f"【历史脚本（追加）：{vf['name']}】\n{vf['text']}")
        history_context = "\n\n".join(history_texts)
        
        system_prompt = "你是一个专业的金融/交易类视频脚本编导。请严格学习并模仿用户提供的历史脚本的文案风格、语气和排版格式。\n\n【重要排版指令】：请必须使用 Markdown 语法进行排版输出。为了作为提词器使用时的重音提示，请务必对文案中的核心观点、金句或转折词使用**加粗**（如 `**重点内容**`）或引用块（如 `> 核心金句`）进行高亮。"
        user_content = ""
        if history_context:
            user_content += f"以下是你需要学习参考的历史脚本序列：\n\n{history_context}\n\n====================\n\n"
        user_content += f"请为我创作一期全新的视频脚本，要求如下：\n"
        if prompt_input:
            user_content += f"\n在续写时，请必须结合以下新素材或要求：\n[新素材与要求]：\n{prompt_input}"
        else:
            user_content += "\n请注意：由于我没有提供新素材，请直接根据历史上下文的逻辑推演，自动拟定下一期主题并生成完整的视频文案！"
            
        write_cron_log(f"[视频脚本制作器] 正在向大模型发起请求 (模型: {selected_model})...")
        from openai import OpenAI
        client = OpenAI(api_key=plat_cfg.get("api_key"), base_url=plat_cfg.get("base_url") if plat_cfg.get("base_url") else None)
        response = call_chat_completion(client, selected_model, [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_content}], chan_config=plat_cfg)
        if response.choices and response.choices[0].message.content:
            script_content = response.choices[0].message.content
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = os.path.join(SCRIPT_OUTPUT_DIR, f"script_{timestamp}.md")
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(script_content)
            write_cron_log(f"[视频脚本制作器] 视频脚本已生成并保存到: {save_path}")
        else:
            write_cron_log("[视频脚本制作器] 错误: 大模型返回空数据。")
        write_cron_log("[视频脚本制作器] 定时任务执行完毕。")
    except Exception as e:
        import traceback
        write_cron_log(f"[视频脚本制作器] 异常终止: {str(e)}\n{traceback.format_exc()}")

def run_scheduled_indicator_docs():
    cfg = load_config()
    id_sched = cfg.get("schedulers", {}).get("指标文档制作", {})
    if not id_sched.get("auto_run", False): return
    ui_state = id_sched.get("ui_state", {})
    if not ui_state:
        write_cron_log("[指标文档制作] 未找到保存的配置，跳过执行。")
        return
        
    write_cron_log("==================================================")
    write_cron_log("[指标文档制作] 定时任务启动")
    try:
        plat = cfg.get("platform", "自定义/OpenAI")
        plat_cfg = cfg.get("channel_configs", {}).get(plat, {})
        selected_model = plat_cfg.get("selected_model", "")
        
        selected_indicator = ui_state.get("selected_indicator", "")
        indicators = load_indicators()
        if not selected_indicator or selected_indicator not in indicators:
            write_cron_log(f"[指标文档制作] 错误: 未找到选中的指标 '{selected_indicator}'，任务终止。")
            return
            
        user_content = f"【指标名称】：{selected_indicator}\n【指标源码】：\n{indicators[selected_indicator]['code']}"
        write_cron_log(f"[指标文档制作] 正在向大模型发起标准合规分析 (指标: {selected_indicator}, 模型: {selected_model})...")
        from openai import OpenAI
        client = OpenAI(api_key=plat_cfg.get("api_key"), base_url=plat_cfg.get("base_url") if plat_cfg.get("base_url") else None)
        response = call_chat_completion(client, selected_model, [{"role": "system", "content": PROMPT_INDICATOR_STANDARD}, {"role": "user", "content": user_content}], chan_config=plat_cfg)
        if response.choices and response.choices[0].message.content:
            doc_content = response.choices[0].message.content
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
            md_filename = f"《{selected_indicator}》---合规化_{ts}.md"
            docx_filename = f"《{selected_indicator}》---合规化_{ts}.docx"
            md_path = os.path.join(INDICATOR_DOCS_DIR, md_filename)
            docx_path = os.path.join(INDICATOR_DOCS_DIR, docx_filename)
            
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(doc_content)
            markdown_to_docx_file(doc_content, docx_path, indicator_name=selected_indicator)
            write_cron_log(f"[指标文档制作] 标准合规分析文档已生成并保存到: {md_path}")
        else:
            write_cron_log("[指标文档制作] 错误: 大模型返回空数据。")
        write_cron_log("[指标文档制作] 定时任务执行完毕。")
    except Exception as e:
        import traceback
        write_cron_log(f"[指标文档制作] 异常终止: {str(e)}\n{traceback.format_exc()}")

def update_scheduler():
    from apscheduler.triggers.cron import CronTrigger
    sched = st.session_state.scheduler
    
    # 清理所有 auto_job 任务
    for job in sched.get_jobs():
        if job.id.startswith("auto_job"):
            sched.remove_job(job.id)
            
    cfg = load_config()
    schedulers = cfg.get("schedulers", {})
    
    # 1. AI 深度分析
    da_sched = schedulers.get("AI 深度分析", {})
    if da_sched.get("auto_run", False):
        try:
            cron_expr = da_sched.get("cron_expr", "0 8 * * *")
            sched.add_job(run_scheduled_deep_analysis, CronTrigger.from_crontab(cron_expr), id="auto_job_deep_analysis")
        except Exception as e:
            pass
            
    # 2. 视频脚本制作器
    vs_sched = schedulers.get("视频脚本制作器", {})
    if vs_sched.get("auto_run", False):
        try:
            cron_expr = vs_sched.get("cron_expr", "0 8 * * *")
            sched.add_job(run_scheduled_video_script, CronTrigger.from_crontab(cron_expr), id="auto_job_video_script")
        except Exception as e:
            pass
            
    # 3. 指标文档制作
    id_sched = schedulers.get("指标文档制作", {})
    if id_sched.get("auto_run", False):
        try:
            cron_expr = id_sched.get("cron_expr", "0 8 * * *")
            sched.add_job(run_scheduled_indicator_docs, CronTrigger.from_crontab(cron_expr), id="auto_job_indicator_docs")
        except Exception as e:
            pass

    # 4. 微信公众号定时群发任务自愈重载机制
    # 移除现有 scheduler 中所有以 schedule_publish_ 开头的推文定时群发 Job，重新扫描载入
    for job in sched.get_jobs():
        if job.id.startswith("schedule_publish_"):
            try:
                sched.remove_job(job.id)
            except:
                pass
                
    wechat_dir = os.path.join("outputs", "wechat")
    if os.path.exists(wechat_dir):
        import glob
        import datetime as dt_module
        draft_files = glob.glob(os.path.join(wechat_dir, "*.draft.json"))
        for draft_file in draft_files:
            try:
                with open(draft_file, "r", encoding="utf-8") as f_draft:
                    draft_data = json.load(f_draft)
                    
                if draft_data.get("status") == "scheduled":
                    media_id = draft_data.get("media_id")
                    scheduled_time_str = draft_data.get("scheduled_time")
                    
                    if media_id and scheduled_time_str:
                        scheduled_dt = dt_module.datetime.strptime(scheduled_time_str, "%Y-%m-%d %H:%M:%S")
                        now_dt = dt_module.datetime.now()
                        
                        if scheduled_dt > now_dt:
                            # 目标时间在未来，重新注册
                            job_id = f"schedule_publish_{media_id}"
                            sched.add_job(
                                run_scheduled_wechat_publish,
                                'date',
                                run_date=scheduled_dt,
                                args=[os.path.abspath(draft_file)],
                                id=job_id
                            )
                            write_cron_log(f"[定时发布自愈] 成功重载任务 {job_id}，时间：{scheduled_time_str}")
                        else:
                            # 时间已过期，出于安全考虑将其重置为草稿状态并记录警告
                            draft_data["status"] = "draft"
                            with open(draft_file, "w", encoding="utf-8") as f_write:
                                json.dump(draft_data, f_write, ensure_ascii=False, indent=2)
                            write_cron_log(f"[定时发布自愈] 警告：任务已过期且未成功发布，已重置为 draft 草稿状态：{draft_file}")
                            
            except Exception as ex:
                write_cron_log(f"[定时发布自愈] 异常：载入草稿 {draft_file} 失败: {str(ex)}")

update_scheduler()

@st.dialog("✨ 动态 Cron 表达式配置器")
def cron_configurator_dialog(page_selection, curr_cron):
    st.info("💡 提示：留空表示不限制（即 *）。侧边栏支持直接手写诸如 `*/5` 这种高阶语法。")
    
    parts = curr_cron.split()
    if len(parts) >= 5:
        c_min, c_hour, c_dom, c_mon, c_dow = parts[0], parts[1], parts[2], parts[3], parts[4]
    else:
        c_min, c_hour, c_dom, c_mon, c_dow = "0", "8", "*", "*", "*"

    def safe_parse(val, opts):
        return [v for v in val.split(',') if v in opts] if val != '*' else []
    
    min_opts = [str(i) for i in range(60)]
    hour_opts = [str(i) for i in range(24)]
    dom_opts = [str(i) for i in range(1, 32)]
    mon_opts = [str(i) for i in range(1, 13)]
    dow_opts = [str(i) for i in range(7)]

    col1, col2 = st.columns(2)
    with col1:
        sel_min = st.multiselect("分钟 (多选)", min_opts, default=safe_parse(c_min, min_opts), placeholder="留空为 *")
        sel_dom = st.multiselect("日期 (多选)", dom_opts, default=safe_parse(c_dom, dom_opts), placeholder="留空为 *")
        sel_dow = st.multiselect("星期 (多选, 0为周日)", dow_opts, default=safe_parse(c_dow, dow_opts), placeholder="留空为 *")
    with col2:
        sel_hour = st.multiselect("小时 (多选)", hour_opts, default=safe_parse(c_hour, hour_opts), placeholder="留空为 *")
        sel_mon = st.multiselect("月份 (多选)", mon_opts, default=safe_parse(c_mon, mon_opts), placeholder="留空为 *")
        
    n_min = ",".join(sel_min) if sel_min else "*"
    n_hour = ",".join(sel_hour) if sel_hour else "*"
    n_dom = ",".join(sel_dom) if sel_dom else "*"
    n_mon = ",".join(sel_mon) if sel_mon else "*"
    n_dow = ",".join(sel_dow) if sel_dow else "*"
    
    new_cron = f"{n_min} {n_hour} {n_dom} {n_mon} {n_dow}"
    
    st.markdown(f"**当前生成的表达式:** `{new_cron}`")
    
    try:
        from apscheduler.triggers.cron import CronTrigger
        from datetime import datetime
        trigger = CronTrigger.from_crontab(new_cron)
        now = datetime.now()
        runs = []
        from datetime import timedelta
        for _ in range(5):
            now = trigger.get_next_fire_time(None, now)
            runs.append(now.strftime("%Y-%m-%d %H:%M:%S"))
            now = now + timedelta(seconds=1)
        st.success("🔮 **未来 5 次执行时间预演:**\n\n" + "\n".join([f"- {r}" for r in runs]))
        is_valid = True
    except Exception as e:
        st.error(f"❌ 表达式无效或无法解析，请检查输入格式。({str(e)})")
        is_valid = False

    if st.button("✅ 确定使用此配置", use_container_width=True, disabled=not is_valid):
        st.session_state[f"cron_input_{page_selection}"] = new_cron
        st.rerun()


# --- Helpers ---

def call_chat_completion(client, model, messages, chan_config=None):
    kwargs = {
        "model": model,
        "messages": messages
    }
    if chan_config and isinstance(chan_config, dict):
        if chan_config.get("enable_thinking", False):
            kwargs["extra_body"] = {"thinking": {"type": "enabled"}}
            effort = chan_config.get("reasoning_effort", "high")
            if effort in ["high", "max"]:
                kwargs["reasoning_effort"] = effort
    return client.chat.completions.create(**kwargs)

def parse_uploaded_file(uploaded_file):
    if uploaded_file.name.endswith('.md'):
        return uploaded_file.getvalue().decode('utf-8')
    elif uploaded_file.name.endswith('.docx'):
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(uploaded_file.getvalue()))
        return "\n".join([p.text for p in doc.paragraphs])
    elif uploaded_file.name.endswith('.csv'):
        import pandas as pd
        df = pd.read_csv(uploaded_file)
        return df.to_csv(index=False)
    elif uploaded_file.name.endswith('.xlsx'):
        import pandas as pd
        df = pd.read_excel(uploaded_file)
        return df.to_csv(index=False)
    return ""

def extract_text_from_pdf(content):
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"[PDF 提取失败: {str(e)}]"

def generate_summary(text, api_key, base_url, model, mode, custom_prompt=None, chan_config=None):
    if not api_key: return "未提供 API Key。"
    system_prompt = (
        "你是一个顶级金融与行业研究专家。请严格按照研报级格式输出最终成品，剔除所有 AI 味，禁止使用‘综上所述’、‘首先/其次’等呆板结构。\n"
        "【排版强制要求】\n"
        "1. 必须使用丰富的 Markdown 语法（多级标题、加粗、引用块、无序/有序列表等）构建极具结构美感的排版。\n"
        "2. 在关键数据对比、产业链剖析、多维度评估等环节，**必须大量使用 Markdown 表格**来归纳，拒绝大段纯文本堆砌。\n"
        "3. 在列举要点时，必须使用带高亮短标题的列表结构，例如：`- **核心看点**：详细解析...`。\n"
        "4. 直接输出正文，绝对不要输出“好的”、“这是报告”等任何废话。\n"
        "【内容合规与风险隔离要求（至关重要）】\n"
        "1. **禁止公司点评**：严禁对第三方公司进行主观点评或说他好坏。只描述客观物理现实、事实逻辑和上下游合作关系，禁止带有任何主观褒贬色彩。\n"
        "2. **严禁个股+股价表现组合**：针对各个股票，即使是正在发生的现实情况，也绝不能直接点名“个股名称/股票代码 + 股价表现描述（如大涨、大跌、暴涨、崩盘、突破新高）或未来走势预测”。一旦出现此种组合搭配即为违规。\n"
        "3. **板块与行业代替**：必须将个股的股价大幅波动描述抽象为对该个股所处的行业或板块的整体走势进行描述，并替换为相近的合规化词语。例如：将“中芯国际股价大涨”替换为“半导体制造板块集体走强”；将“中际旭创暴跌”替换为“光通信板块走势趋弱”；将“大洋电机暴涨”替换为“微电机板块表现活跃”。\n"
        "4. **剔除投资暗示**：必须保持中立客观的第三方行业观察视角。绝不能出现任何带有“买入、建议建仓、看好、潜力巨大”等带有诱导性投资建议或价格预测的表述。\n"
        "5. **柔化绝对性用词**：禁止使用“一定”、“必然”、“绝对”等肯定用词，必须将其替换为“可能”、“有望”、“或将”、“呈现XX趋势”等客观柔性的学术化表述，极大降低合规风险。\n"
    )
    if mode == "个股分析": 
        system_prompt += "【当前任务】：深度个股价值分析，必须涵盖核心逻辑、业务拆解、估值及风险。重点指标必须用表格解构。"
    elif mode == "行业分析": 
        system_prompt += "【当前任务】：深度行业宏观趋势分析，涵盖宏观驱动力、产业链上下游剖析、市场竞争格局及展望。产业链和竞争格局环节必须强制使用表格排版。"
    else: 
        system_prompt += "【当前任务】：详尽总结分析，提取核心要点，将散乱的信息重构成逻辑极为清晰、带有丰富表格和加粗高亮的深度简报。"
    
    user_content = f"指令: {custom_prompt}\n\n待处理内容:\n{text}" if custom_prompt else text
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url if base_url else None)
        response = call_chat_completion(client, model, [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_content}], chan_config=chan_config)
        if not response.choices:
            return f"AI 总结失败: 接口返回空数据，可能是该模型({model})暂不支持或网络限流。详情: {response}"
        return response.choices[0].message.content
    except Exception as e: return f"AI 总结失败: {str(e)}"

def fetch_zsxq(group_id, limit=3, scope="all", progress_callback=None):
    if not group_id:
        return "请提供知识星球的 Group ID。", [], []
    
    processed_files = []
    processed_topics_brief = []
    
    # 1. Auth check
    if progress_callback: progress_callback({"type": "info", "msg": "正在检查授权状态..."})
    status_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8')
    is_logged = False
    try:
        m = re.search(r'\{.*\}', status_res.stdout, re.DOTALL)
        if m:
            auth_data = json.loads(m.group(0))
            if auth_data.get("ok") and auth_data.get("data", {}).get("loggedIn"):
                is_logged = True
    except: pass
    
    if not is_logged:
        return "获取失败: 知识星球未授权。请在左侧边栏获取授权链接。", [], []

    # 2. Fetch topics
    if progress_callback: progress_callback({"type": "info", "msg": "正在向知识星球请求最新动态..."})
    # 如果是仅限附件模式，加大获取数量以供滚动筛选
    fetch_limit = max(20, limit * 3) if scope == "files" else limit
    max_retries = 3
    import time
    for attempt in range(max_retries):
        result = subprocess.run([
            NPX_CMD, "zsxq-cli", "group", "+topics", "--group-id", str(group_id), "--limit", str(fetch_limit), "--json"
        ], capture_output=True, text=True, encoding='utf-8', timeout=30)
        
        if result.returncode == 0:
            break
            
        err_msg = result.stderr.strip() if result.stderr else (result.stdout.strip() if result.stdout else "CLI 内部错误")
        if any(keyword in err_msg.upper() for keyword in ["EOF", "TRANSPORT", "TIMEOUT", "CONNECTION", "CLOSED"]):
            if attempt < max_retries - 1:
                sleep_sec = 3 * (attempt + 1)
                if progress_callback: 
                    progress_callback({"type": "info", "msg": f"⚠️ 网络请求异常 ({err_msg})，将在 {sleep_sec} 秒后进行第 {attempt+2} 次重试..."})
                time.sleep(sleep_sec)
                continue
                
        m_err = re.search(r'(error:.*?)(?:\r|\n|$)', err_msg)
        if m_err: err_msg = m_err.group(1)
        return f"获取失败: {err_msg}", [], []

    try:
        m = re.search(r'\{.*\}', result.stdout, re.DOTALL)
        if not m: return "获取失败: 无法解析返回数据", [], []
        
        data = json.loads(m.group(0))
        if not (data.get("success") or data.get("ok")):
            return f"获取失败: {data.get('message', '未知错误')}", [], []
        
        topics = data.get("data", {}).get("topics", []) or data.get("topics_brief", [])
        
        print("\n=== [DEBUG] 开始遍历 Topic ===")
        content_list = []
        valid_topics_count = 0
        
        for topic in topics:
            if valid_topics_count >= limit:
                break
                
            # 1. 修正提取路径：直接读取并防止 NoneType
            files_in_topic = topic.get('files') or []
            
            topic_text = ""
            if 'talk' in topic and isinstance(topic['talk'], dict): topic_text = topic['talk'].get('text', '')
            elif 'question' in topic and isinstance(topic['question'], dict): topic_text = f"提问: {topic['question'].get('text', '')}\n回答: {topic.get('answer', {}).get('text', '')}"
            elif 'article' in topic and isinstance(topic['article'], dict): topic_text = f"文章: {topic['article'].get('title', '')}"
            elif 'content' in topic: topic_text = topic.get('content', '')
            elif 'title' in topic or 'text' in topic: topic_text = topic.get('title', '') or topic.get('text', '')
            
            # 更精准的日志打印
            print(f"\n--- [DEBUG] Topic ID: {topic.get('topic_id')} ---")
            print(f"Content Preview: {topic_text[:20]}...")
            print(f"Files raw data: {str(files_in_topic)}")
            
            topic_preview = (topic_text[:20].replace('\n', ' ') + "...") if topic_text else "无正文内容"
            
            if scope == "files":
                has_potential = False
                for f in files_in_topic:
                    fname = f.get('name') or f.get('file_name') or f.get('title') or ''
                    if fname.lower().endswith(('.pdf', '.docx', '.xlsx', '.csv', '.md')):
                        has_potential = True
                        break
                if not has_potential:
                    continue
                    
            if progress_callback: progress_callback({"type": "topic_start", "topic_id": topic.get('topic_id'), "preview": topic_preview})
            
            # 添加异常数据诊断信息
            diagnostic_info = f"TopicID: {topic.get('topic_id', 'Unknown')}, Keys: {list(topic.keys())}"
            if files_in_topic:
                diagnostic_info += f", 找到附件数量: {len(files_in_topic)}"
            else:
                diagnostic_info += ", 未找到任何附件节点"
            processed_topics_brief.append(diagnostic_info)
            
            current_content = ""
            if topic_text and topic_text != "「文件」":
                current_content += topic_text

            has_valid_file_content = False
            for f in files_in_topic:
                # 2. 兼容提取文件名属性
                fname = f.get('name') or f.get('file_name') or f.get('title') or ''
                fid = f.get('file_id') or f.get('id') or ''
                if not fid: continue
                
                # 明确剔除 MP3 等不需要的音频和富媒体
                if fname.lower().endswith(('.mp3', '.mp4', '.zip', '.rar', '.jpg', '.png')):
                    if progress_callback: progress_callback({"type": "topic_log", "msg": f"⏭️ `{fname}` 属于过滤文件，已跳过。"})
                    continue
                
                # 3. 统统保留并提取支持的文档格式
                if fname.lower().endswith(('.pdf', '.docx', '.xlsx', '.csv', '.md')):
                    if progress_callback: progress_callback({"type": "topic_log", "msg": f"⏳ 正在下载并解析附件: `{fname}` (可能需要一些时间...)"})
                    try:
                        dl_res = None
                        for dl_attempt in range(3):
                            dl_res = subprocess.run([NPX_CMD, "zsxq-cli", "api", "call", "call_zsxq_api", "--params", json.dumps({"method": "GET", "path": f"/v2/files/{fid}/download_url"})], capture_output=True, text=True, encoding='utf-8', timeout=10)
                            if dl_res.returncode == 0:
                                break
                            time.sleep(2 * (dl_attempt + 1))
                        dl_m = re.search(r'\{.*\}', dl_res.stdout, re.DOTALL)
                        if dl_m:
                            dl_url = json.loads(dl_m.group(0)).get('download_url') or json.loads(dl_m.group(0)).get('body', {}).get('download_url')
                            if dl_url:
                                f_raw = httpx.get(dl_url, timeout=30).content
                                extracted = ""
                                if fname.lower().endswith('.pdf'):
                                    extracted = extract_text_from_pdf(f_raw)
                                elif fname.lower().endswith('.docx'):
                                    from io import BytesIO
                                    from docx import Document
                                    extracted = "\n".join([p.text for p in Document(BytesIO(f_raw)).paragraphs])
                                elif fname.lower().endswith('.xlsx'):
                                    import pandas as pd
                                    from io import BytesIO
                                    extracted = pd.read_excel(BytesIO(f_raw)).to_csv(index=False)
                                elif fname.lower().endswith('.csv'):
                                    import pandas as pd
                                    from io import BytesIO
                                    extracted = pd.read_csv(BytesIO(f_raw)).to_csv(index=False)
                                elif fname.lower().endswith('.md'):
                                    extracted = f_raw.decode('utf-8')
                                
                                if extracted.strip():
                                    current_content += f"\n[附件内容: {fname}]:\n{extracted}"
                                    processed_files.append(fname)
                                    has_valid_file_content = True
                                    if progress_callback: progress_callback({"type": "topic_log", "msg": f"✅ `{fname}` 解析成功，提取字数: {len(extracted)}"})
                                else:
                                    if progress_callback: progress_callback({"type": "topic_log", "msg": f"⚠️ `{fname}` 提取内容为空"})
                    except Exception as e:
                        if progress_callback: progress_callback({"type": "topic_log", "msg": f"❌ 解析 `{fname}` 失败: {str(e)}"})

            # 如果是“仅限附件”模式，且本条动态没有任何可以解析成功的附件，则放弃本条动态，继续寻找下一条
            if scope == "files" and not has_valid_file_content:
                if progress_callback: progress_callback({"type": "topic_end", "success": False, "preview": topic_preview, "reason": "未提取到有效附件内容"})
                continue

            if current_content.strip(): 
                content_list.append(current_content)
                valid_topics_count += 1
                if progress_callback: progress_callback({"type": "topic_end", "success": True, "preview": topic_preview})
            else:
                if progress_callback: progress_callback({"type": "topic_end", "success": False, "preview": topic_preview, "reason": "动态内容为空"})

        print("===============================\n")

        # 4. 修复 NoneType 崩溃 Bug: 在没有匹配结果时直接返回中断标识
        if not content_list:
            return "跳过分析：未找到有效文档附件", processed_files, processed_topics_brief
        return "\n\n---\n\n".join(content_list), processed_files, processed_topics_brief
        
    except Exception as e:
        return f"获取异常: {str(e)}", [], []



def render_to_image(summary_text, mode_name):
    html_content = markdown.markdown(summary_text, extensions=['tables', 'fenced_code', 'nl2br'])
    
    logo_html = ""
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            logo_b64 = base64.b64encode(f.read()).decode()
        logo_html = f'<img src="data:image/png;base64,{logo_b64}" class="footer-logo" alt="logo">'

    html_template = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;500;700;900&display=swap');
            :root {{ --primary: #c19b52; --brand: #1d2d50; --gold: #c19b52; --bg: #ffffff; }}
            body {{ font-family: 'Noto Sans SC', sans-serif; background: #f0f2f5; margin: 0; padding: 20px; display: flex; justify-content: center; }}
            .container {{ background: var(--bg); width: 800px; padding: 50px 60px; box-shadow: 0 4px 20px rgba(0,0,0,0.08); position: relative; }}
            .header-banner {{ 
                background: linear-gradient(135deg, #1d2d50 0%, #2b4170 100%); 
                padding: 40px 60px 35px 60px; margin: -50px -60px 45px -60px; color: white; position: relative; border-bottom: 4px solid var(--gold);
                display: flex; justify-content: space-between; align-items: flex-end;
            }}
            .header-title {{ font-size: 34px; font-weight: 900; margin: 0; letter-spacing: 2px; color: #ffffff; text-shadow: 0 2px 10px rgba(0,0,0,0.3); }}
            .header-time {{ font-size: 14px; color: #e2e8f0; opacity: 0.9; margin-bottom: 5px; letter-spacing: 1px; font-weight: 500; }}
            
            h1 {{ font-size: 28px; color: var(--brand); margin-top: 40px; border-bottom: 2px solid #f0f0f0; padding-bottom: 15px; font-weight: 900; }}
            h2 {{ 
                font-size: 22px; color: var(--brand); margin-top: 35px; margin-bottom: 20px; 
                position: relative; padding-bottom: 10px; font-weight: 800;
            }}
            h2::after {{ content: ''; position: absolute; left: 0; bottom: 0; width: 50px; height: 4px; background: var(--gold); }}
            h3 {{ font-size: 18px; color: #333; margin-top: 25px; border-left: 5px solid var(--gold); padding-left: 15px; font-weight: 700; }}
            
            p {{ font-size: 16px; color: #353535; line-height: 1.8; margin-bottom: 18px; text-align: justify; }}
            ul, ol {{ padding-left: 20px; margin-bottom: 20px; }}
            li {{ font-size: 15.5px; color: #353535; line-height: 1.7; margin-bottom: 8px; }}
            
            table {{ width: 100%; border-collapse: collapse; margin: 25px 0; border-radius: 8px; overflow: hidden; border: 1px solid #e5e5e5; }}
            th {{ background: #f8f9fa; color: var(--brand); font-weight: 700; text-align: left; padding: 12px 15px; border-bottom: 2px solid var(--gold); }}
            td {{ padding: 12px 15px; border-bottom: 1px solid #f0f0f0; color: #353535; font-size: 14.5px; }}
            
            strong {{ color: var(--brand); font-weight: 700; }}
            .disclaimer {{ margin-top: 60px; padding: 20px; background: #f8f9fa; border-radius: 8px; font-size: 13.5px; color: #666; line-height: 1.7; border-left: 4px solid var(--gold); text-align: justify; }}
            .footer-logo {{ height: 35px; width: auto; object-fit: contain; }}
            .footer {{ margin-top: 30px; display: flex; justify-content: space-between; align-items: center; font-size: 12px; color: #999; border-top: 1px solid #eee; padding-top: 20px; letter-spacing: 1px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header-banner">
                <div class="header-title">深度分析报告</div>
                <div class="header-time">{datetime.now().strftime('%Y年%m月%d日')}</div>
            </div>
            <div class="content">{html_content}</div>
            <div class="disclaimer">
                <strong>免责声明：</strong>本文内容及数据均基于公开市场资料与行业研报，仅作逻辑梳理与行业趋势分析之用，旨在探讨投资理念与方法，不构成任何具体的投资建议或操作指引。文中提及的企业及产品仅作为产业案例分析，不构成推荐。投资有风险，入市需谨慎。请您基于自身独立判断做出决策。
            </div>
            <div class="footer">
                <div>{logo_html}</div>
                <div>DEEP SUMMARY PRO · 生成于 {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
            </div>
        </div>
    </body>
    </html>
    """
    output_path = os.path.join(IMAGE_OUTPUT_DIR, f"summary_{datetime.now().strftime('%Y%m%d%H%M%S')}.png")
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_template)
        page.wait_for_load_state("networkidle")
        page.locator(".container").screenshot(path=output_path)
        browser.close()
    return output_path

# --- Streamlit UI ---
st.set_page_config(page_title="AI 多功能控制台", layout="wide")

if 'show_success_toast' in st.session_state:
    st.toast(st.session_state['show_success_toast'], icon="🎉")
    st.balloons()
    del st.session_state['show_success_toast']

with st.sidebar:
    st.header("🧭 导航")
    page_selection = st.radio("选择功能模块", ["AI 深度分析", "视频脚本制作器", "指标文档制作"], label_visibility="collapsed")
    st.markdown("---")
    
    st.header("⚙️ 全局配置")
    
    @st.cache_data(ttl=300, show_spinner=False)
    def get_zsxq_auth_status():
        res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8')
        return res.stdout

    @st.cache_data(ttl=300, show_spinner=False)
    def get_dreamina_credit_status():
        res = subprocess.run([DREAMINA_CMD, "user_credit"], capture_output=True, text=True, encoding='utf-8')
        return res.returncode, res.stdout
        
    def clear_auth_cache():
        get_zsxq_auth_status.clear()
        get_dreamina_credit_status.clear()
        
    if hasattr(st, "dialog") and "show_log_dialog" not in globals():
        @st.dialog("📋 任务执行日志", width="large")
        def show_log_dialog(log_content, page_selection="AI 深度分析"):
            try:
                cfg = load_config()
                da_sched = cfg.get("schedulers", {}).get(page_selection, {})
                ui_state = da_sched.get("ui_state", {})
                if ui_state:
                    st.markdown(f"**📢 当前定时任务配置 ({page_selection})：**")
                    if page_selection == "AI 深度分析":
                        c1, c2, c3, c4 = st.columns(4)
                        with c1:
                            st.caption(f"👥 **星球群组**：{ui_state.get('selected_group', '默认')}")
                            st.caption(f"⏰ **调度周期**：{da_sched.get('cron_expr', '未配置')}")
                        with c2:
                            st.caption(f"📝 **分析模式**：{ui_state.get('a_mode', '常规')}")
                            st.caption(f"🔢 **拉取篇数**：{ui_state.get('l_limit', 3)} 篇")
                        with c3:
                            st.caption(f"📱 **微信推文**：{'开启' if ui_state.get('use_wechat') else '关闭'}")
                            st.caption(f"🎨 **排版模式**：{ui_state.get('wechat_mode', 'AI模式')}")
                        with c4:
                            st.caption(f"🍁 **排版主题**：{ui_state.get('wechat_theme', 'spring-fresh')}")
                            st.caption(f"📐 **字号/背景**：{ui_state.get('wechat_font_size', 'medium')} / {ui_state.get('wechat_background_type', 'none')}")
                    else:
                        st.json(ui_state)
                    st.divider()
            except Exception as e:
                st.caption(f"无法加载配置快照: {str(e)}")
            st.text_area("", value=log_content, height=500, disabled=True)
            
    if hasattr(st, "dialog") and "show_wechat_publish_dialog" not in globals():
        @st.dialog("📲 微信公众号一键发布与预览", width="large")
        def show_wechat_publish_dialog():
            import urllib.parse
            
            # 自动从关联 of .draft.json 中加载草稿状态，防止页面刷新丢失
            filepath = st.session_state.get('latest_wechat_file_path')
            draft_info = wechat_publisher.load_draft_info(filepath) if filepath else {}
            
            is_published = (draft_info.get("status") == "published")
            is_scheduled = (draft_info.get("status") == "scheduled")
            
            if is_published:
                st.error(f"⚠️ **防重群发拦截**：该文章已于 `{draft_info.get('publish_time', '此前')}` 正式群发到公众号，不可重复操作！")
            elif is_scheduled:
                st.info(f"📅 **定时发布中**：该文章已排期于 `{draft_info.get('scheduled_time', '未来')}` 自动定时群发。")
            
            if "wechat_draft_media_id" not in st.session_state or st.session_state.wechat_draft_media_id is None:
                st.session_state.wechat_draft_media_id = draft_info.get("media_id")
            if "wechat_draft_url" not in st.session_state or st.session_state.wechat_draft_url is None:
                st.session_state.wechat_draft_url = draft_info.get("url")
            if "wechat_publish_result" not in st.session_state:
                st.session_state.wechat_publish_result = None
            if "wechat_preview_status" not in st.session_state:
                st.session_state.wechat_preview_status = None
                
            st.markdown("### 1. 公众号账号选择与管理")
            accounts = wechat_publisher.load_accounts()
            
            account_names = [a["name"] for a in accounts]
            selected_acc_name = st.selectbox(
                "选择当前操作的公众号", 
                account_names + ["➕ 新增公众号账号配置..."], 
                index=0 if account_names else 0
            )
            
            active_account = None
            if selected_acc_name == "➕ 新增公众号账号配置...":
                with st.form("add_wechat_acc_form"):
                    st.write("🔑 **添加公众号开发者凭证 (AppID / AppSecret)**")
                    new_name = st.text_input("公众号名称 (例如：慧峰金融)", placeholder="请输入公众号名称")
                    new_appid = st.text_input("开发者 ID (AppID)", placeholder="请输入 AppID")
                    new_secret = st.text_input("开发者密码 (AppSecret)", type="password", placeholder="请输入 AppSecret")
                    
                    submitted = st.form_submit_button("💾 保存配置", use_container_width=True)
                    if submitted:
                        if new_name and new_appid and new_secret:
                            accounts = [a for a in accounts if a["name"] != new_name.strip()]
                            accounts.append({
                                "name": new_name.strip(),
                                "appid": new_appid.strip(),
                                "secret": new_secret.strip()
                            })
                            wechat_publisher.save_accounts(accounts)
                            st.success(f"公众号「{new_name}」配置保存成功！")
                            st.rerun()
                        else:
                            st.error("❌ 所有字段均为必填项！")
            elif selected_acc_name:
                active_account = next((a for a in accounts if a["name"] == selected_acc_name), None)
                if active_account:
                    col_det, col_del = st.columns([5, 1.5])
                    with col_det:
                        st.caption(f"🛡️ AppID: `{active_account['appid'][:6]}******` | 密钥已妥善加密保存")
                    with col_del:
                        if st.button("🗑️ 删除该公众号", key="del_wechat_acc_btn", use_container_width=True):
                            accounts = [a for a in accounts if a["name"] != selected_acc_name]
                            wechat_publisher.save_accounts(accounts)
                            st.success("账号已删除")
                            st.rerun()
                            
            if not active_account:
                st.info("💡 请先添加或选择一个有效的微信公众号。")
                return
                
            st.divider()
            
            st.markdown("### 2. IP 白名单安全配置")
            with st.spinner("正在获取服务器外网出口 IP..."):
                server_ip = wechat_publisher.get_server_ip()
            st.info(
                f"🖥️ **当前服务器公网 IP：`{server_ip}`**\n\n"
                "⚠️ **重要提示**：请登录「微信公众平台 -> 开发接口管理 -> IP白名单」，把上述公网 IP 加进去。否则接口调用会返回 `ip not in whitelist` 报错。"
            )
            
            st.divider()
            
            st.markdown("### 3. 推文封面图及基本元数据")
            
            latest_wechat = st.session_state.get('latest_wechat', '')
            parsed_title = "未命名推文"
            for line in latest_wechat.split("\n"):
                line = line.strip()
                if line.startswith("#"):
                    t = re.sub(r'^#+\s*', '', line).strip()
                    if t:
                        parsed_title = t
                        break
            if len(parsed_title) > 32:
                parsed_title = parsed_title[:32]
                
            meta_title = st.text_input("推文标题 (微信要求 ≤ 32字)", value=parsed_title)
            meta_author = st.text_input("作者名称 (建议 ≤ 8字)", value="AI 架构师")
            
            # 强健地过滤各种 markdown 图片格式（兼容空格及换行）
            clean_text = re.sub(r'!\[[^\]]*\]\s*\([^)]*\)', '', latest_wechat)
            # 强健地过滤 markdown 链接格式（兼容空格）
            clean_text = re.sub(r'\[[^\]]*\]\s*\([^)]*\)', '', clean_text)
            # 过滤任何可能存在的 HTML 标签
            clean_text = re.sub(r'<[^>]*>', '', clean_text)
            # 安全防线：彻底清除可能泄露的本地图片路径特征，确保摘要纯净
            clean_text = re.sub(r'outputs[/\\]wechat[/\\]images[/\\][^\s]+', '', clean_text)
            clean_text = re.sub(r'\(\s*outputs[/\\]wechat[/\\]images[/\\][^\s]*\)', '', clean_text)
            # 清理特殊排版字符及多余空格
            clean_text = re.sub(r'[#\*_`\-\>\+\n\r\t]', ' ', clean_text)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            default_digest = clean_text[:80] + "..." if len(clean_text) > 80 else clean_text
            meta_digest = st.text_area("推文摘要 (微信要求 ≤ 128字)", value=default_digest, height=70)
            
            st.markdown("**封面图预览与设置**")
            local_images = []
            # 强健识别包含空格在内的所有配图路径，支持 WebP
            img_pattern = re.compile(r'!\[.*?\]\s*\(\s*([^)\s]+\.(?:jpg|png|jpeg|webp))\s*\)', re.IGNORECASE)
            found_imgs = img_pattern.findall(latest_wechat)
            for img in found_imgs:
                img_path = img.replace("\\", "/").strip()
                # 兼容相对路径、绝对路径、前导斜杠以及不同系统的路径前缀
                paths_to_check = [img_path]
                if img_path.startswith("/"):
                    paths_to_check.append(img_path[1:])
                else:
                    paths_to_check.append("/" + img_path)
                
                resolved_path = None
                for p in paths_to_check:
                    if os.path.exists(p) and not os.path.isdir(p):
                        resolved_path = p
                        break
                
                if resolved_path and resolved_path not in local_images:
                    local_images.append(resolved_path)
                    
            cover_mode_sel = st.radio("选择封面图来源", ["从推文已生成的配图中选择", "手动上传自定义封面图片"], horizontal=True)
            
            selected_cover_path = None
            if cover_mode_sel == "从推文已生成的配图中选择":
                if local_images:
                    selected_cover_path = st.selectbox("选择配图", local_images, format_func=lambda x: os.path.basename(x))
                    if selected_cover_path:
                        st.image(selected_cover_path, width=240, caption="选定的封面图")
                else:
                    st.warning("⚠️ 推文中未检测到本地配图，请选择手动上传封面。")
            else:
                uploaded_cover = st.file_uploader("上传 JPG/PNG 格式封面 (建议宽高比 2.35:1)", type=["jpg", "png", "jpeg"])
                if uploaded_cover:
                    os.makedirs("tests", exist_ok=True)
                    temp_cover_path = os.path.join("tests", f"temp_upload_cover_{int(time.time())}.jpg")
                    with open(temp_cover_path, "wb") as f_cover:
                        f_cover.write(uploaded_cover.getbuffer())
                    selected_cover_path = temp_cover_path
                    st.image(uploaded_cover, width=240, caption="上传的自定义封面")
                    
            st.divider()
            
            st.markdown("### 4. 预览与正式发布控制台")
            
            c_action1, c_action2, c_action3 = st.columns(3)
            
            with c_action1:
                if st.button("1️⃣ 生成草稿箱文章", use_container_width=True, type="primary"):
                    if not selected_cover_path:
                        st.error("❌ 请先设置封面图！")
                        return
                    with st.spinner("🚀 正在上传所有配图、上传封面图、并在微信后台生成草稿..."):
                        try:
                            token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                            
                            raw_html = convert_to_wechat_html(
                                st.session_state['latest_wechat'],
                                st.session_state.get('latest_wechat_rendered_theme', 'spring-fresh'),
                                st.session_state.get('latest_wechat_rendered_mode', 'AI 模式 (免费)'),
                                api_key=config.get("md2wechat_api_key", ""),
                                font_size=config.get("wechat_font_size", "medium"),
                                bg_type=config.get("wechat_background_type", "none"),
                                chan_config=current_chan_config,
                                custom_prompt=config.get("wechat_custom_prompt", ""),
                                for_wechat_api=True
                            )
                            
                            final_html = wechat_publisher.replace_local_images_with_wechat_urls(raw_html, token)
                            cover_media_id = wechat_publisher.upload_cover_image(selected_cover_path, token)
                            draft_media_id = wechat_publisher.create_draft(
                                token, 
                                meta_title, 
                                meta_author, 
                                meta_digest, 
                                final_html, 
                                cover_media_id
                            )
                            draft_url = wechat_publisher.get_draft_url(token, draft_media_id)
                            
                            st.session_state.wechat_draft_media_id = draft_media_id
                            st.session_state.wechat_draft_url = draft_url
                            st.session_state.wechat_publish_result = None
                            st.session_state.wechat_preview_status = None
                            
                            # 将草稿状态持久化到文件，防止刷新丢失
                            filepath = st.session_state.get('latest_wechat_file_path')
                            if filepath:
                                wechat_publisher.save_draft_info(
                                    filepath, 
                                    draft_media_id, 
                                    draft_url, 
                                    status="draft", 
                                    appid=active_account["appid"], 
                                    secret=active_account["secret"]
                                )
                            st.success("🎉 微信草稿文章创建成功！")
                            st.rerun()
                        except Exception as e:
                            err_str = str(e)
                            m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                            if m_ip:
                                ip_addr = m_ip.group(1)
                                st.error(f"❌ 操作失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                st.warning(f"💡 **解决方法**：请登录「微信公众平台 -> 开发 -> 开发设置 -> IP白名单」，把 IP **`{ip_addr}`** 加进去，然后重新尝试。")
                            else:
                                st.error(f"❌ 操作失败: {err_str}")
                            
            if st.session_state.wechat_draft_url:
                st.markdown("#### 📱 微信扫码预览草稿")
                col_qr, col_info = st.columns([1.2, 2])
                with col_qr:
                    qr_api_url = f"https://api.qrserver.com/v1/create-qr-code/?size=200x200&data={urllib.parse.quote(st.session_state.wechat_draft_url)}"
                    st.image(qr_api_url, width=180, caption="扫码即可在手机微信预览草稿")
                with col_info:
                    st.markdown("**草稿箱 MediaID：**")
                    st.code(st.session_state.wechat_draft_media_id)
                    st.markdown("**临时预览链接 (在微信内打开有效)：**")
                    st.markdown(f"[👉 点击直接跳转预览]({st.session_state.wechat_draft_url})")
                    st.caption("注：微信临时预览链接有短期时效限制，过期后需重新点击「生成草稿箱文章」获取新链接。")
                
                with c_action2:
                    st.write("**手机预览推送**")
                    preview_wx_id = st.text_input("接收人的微信号", placeholder="如: my_wechat_id")
                    if st.button("2️⃣ 发送预览推送", use_container_width=True, disabled=not preview_wx_id):
                        with st.spinner("正在向指定微信号发送预览..."):
                            try:
                                token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                                wechat_publisher.send_preview(token, st.session_state.wechat_draft_media_id, preview_wx_id)
                                st.session_state.wechat_preview_status = f"✅ 预览成功发送至微信号: {preview_wx_id}，请检查手机微信通知！"
                                st.rerun()
                            except Exception as e:
                                err_str = str(e)
                                m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                                if m_ip:
                                    ip_addr = m_ip.group(1)
                                    st.error(f"❌ 发送失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                    st.warning(f"💡 **解决方法**：请登录「微信公众平台」，把 IP **`{ip_addr}`** 加到 IP 白名单中，然后重新尝试。")
                                else:
                                    st.error(f"❌ 发送失败: {err_str}")
                                
                if st.session_state.wechat_preview_status:
                    st.info(st.session_state.wechat_preview_status)
                    
                with c_action3:
                    st.write("**正式发布**")
                    st.warning("⚠️ 正式发布为不可逆的群发操作，将对所有关注者可见！")
                    
                    enable_schedule = st.checkbox("定时群发（不勾选则立即群发）", key="wechat_enable_schedule", disabled=is_published)
                    
                    sched_dt = None
                    if enable_schedule:
                        col_d, col_t = st.columns(2)
                        with col_d:
                            import datetime as dt_module
                            sched_date = st.date_input("发布日期", min_value=dt_module.date.today(), key="wechat_sched_date")
                        with col_t:
                            sched_time = st.time_input("发布时间", key="wechat_sched_time")
                        
                        sched_dt = dt_module.datetime.combine(sched_date, sched_time)
                        
                    confirm_publish = st.checkbox("我已确认扫码预览无误，同意正式群发", disabled=is_published)
                    
                    btn_label = "⏰ 安排定时发布" if enable_schedule else "3️⃣ 正式群发到公众号"
                    
                    if st.button(btn_label, use_container_width=True, type="primary", disabled=(not confirm_publish) or is_published):
                        if enable_schedule and sched_dt:
                            import datetime as dt_module
                            now_dt = dt_module.datetime.now()
                            if sched_dt <= now_dt:
                                st.error("❌ 定时发布时间必须是未来的时间！")
                            else:
                                with st.spinner("正在安排定时发布..."):
                                    try:
                                        # 1. 保存状态为 scheduled 并记入 scheduled_time
                                        wechat_publisher.save_draft_info(
                                            filepath,
                                            st.session_state.wechat_draft_media_id,
                                            st.session_state.wechat_draft_url,
                                            status="scheduled",
                                            scheduled_time=sched_dt.strftime("%Y-%m-%d %H:%M:%S"),
                                            appid=active_account["appid"],
                                            secret=active_account["secret"]
                                        )
                                        # 2. 注册定时任务到 BackgroundScheduler
                                        job_id = f"schedule_publish_{st.session_state.wechat_draft_media_id}"
                                        
                                        # 移除可能已存在的同 ID 任务
                                        try:
                                            st.session_state.scheduler.remove_job(job_id)
                                        except:
                                            pass
                                            
                                        # 计算绝对路径
                                        draft_json_path = os.path.abspath(filepath).replace(".md", ".draft.json")
                                        st.session_state.scheduler.add_job(
                                            run_scheduled_wechat_publish,
                                            'date',
                                            run_date=sched_dt,
                                            args=[draft_json_path],
                                            id=job_id
                                        )
                                        st.session_state.wechat_publish_result = f"📅 定时群发任务已成功设置！安排在 `{sched_dt.strftime('%Y-%m-%d %H:%M:%S')}` 自动执行。请保持系统后台运行。"
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"❌ 设置定时发布失败: {str(e)}")
                        else:
                            # 立即群发
                            with st.spinner("正在群发中，请稍候..."):
                                try:
                                    token = wechat_publisher.get_access_token(active_account["appid"], active_account["secret"])
                                    pub_id = wechat_publisher.publish_draft(token, st.session_state.wechat_draft_media_id)
                                    st.session_state.wechat_publish_result = f"🚀 文章已正式群发！发布 ID: `{pub_id}`。您可在公众平台后台查看群发状态。"
                                    
                                    # 更新为已发布状态
                                    if filepath:
                                        wechat_publisher.save_draft_info(
                                            filepath, 
                                            st.session_state.wechat_draft_media_id, 
                                            st.session_state.wechat_draft_url, 
                                            status="published"
                                        )
                                    st.rerun()
                                except Exception as e:
                                    err_str = str(e)
                                    m_ip = re.search(r'invalid ip ([\d\.]+)', err_str, re.IGNORECASE)
                                    if m_ip:
                                        ip_addr = m_ip.group(1)
                                        st.error(f"❌ 发布失败：获取微信凭证失败。您的实际出口 IP **`{ip_addr}`** 未被微信白名单许可。")
                                        st.warning(f"💡 **解决方法**：请登录「微信公众平台」，把 IP **`{ip_addr}`** 加到 IP 白名单中，然后重新尝试。")
                                    else:
                                        st.error(f"❌ 发布失败: {err_str}")
                                
                if st.session_state.wechat_publish_result:
                    st.success(st.session_state.wechat_publish_result)

    # Pre-load current channel config so it is available globally
    current_chan_config = config["channel_configs"].get(config.get("platform", "自定义/OpenAI"), {})

    # Expander 1: 🔑 授权与星球管理
    with st.expander("🔑 授权与星球管理", expanded=False):
        if st.button("🔄 刷新授权状态", use_container_width=True):
            clear_auth_cache()
            st.rerun()
            
        st.markdown("**知识星球授权**")
        zsxq_stdout = get_zsxq_auth_status()
        logged_in = False
        user_id = ""
        user_name = ""
        try:
            m_auth = re.search(r'\{.*\}', zsxq_stdout, re.DOTALL)
            if m_auth:
                auth_data = json.loads(m_auth.group(0))
                if auth_data.get("ok") and auth_data.get("data", {}).get("loggedIn"):
                    logged_in = True
                    user_id = auth_data["data"].get("userId", "")
                    user_name = auth_data["data"].get("userName", "")
        except: pass
        
        if logged_in:
            st.success(f"✅ {user_name} (已授权)")
            if st.button("退出登录"): 
                subprocess.run([NPX_CMD, "zsxq-cli", "auth", "logout"])
                clear_auth_cache()
                st.rerun()
        else:
            st.warning("⚠️ 未授权")
            if st.button("🔗 获取授权链接"):
                login_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "login", "--json", "--no-browser", "--no-wait"], capture_output=True, text=True, encoding='utf-8')
                try:
                    m_login = re.search(r'\{.*\}', login_res.stdout, re.DOTALL)
                    if m_login:
                        d_login = json.loads(m_login.group(0))["data"]
                        st.session_state.zsxq_device_code = d_login["device_code"]
                        st.markdown(f"**[👉 点击授权]({d_login['verification_uri_complete']})**")
                        st.info(f"确认码：`{d_login['user_code']}`")
                except: st.error("无法获取链接")
            if "zsxq_device_code" in st.session_state and st.button("我已完成授权"):
                with st.spinner("验证中..."):
                    verify_res = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "login", "--device-code", st.session_state.zsxq_device_code, "--json"], capture_output=True, text=True, encoding='utf-8')
                    try:
                        m_verify = re.search(r'\{.*\}', verify_res.stdout, re.DOTALL)
                        if m_verify and json.loads(m_verify.group(0)).get("ok"):
                            st.success("授权成功！")
                            del st.session_state.zsxq_device_code
                            clear_auth_cache()
                            st.rerun()
                        else:
                            st.error("未检测到成功授权，请确认您已在手机端扫码并输入了确认码！")
                    except Exception:
                        st.error("验证失败：无法解析登录状态。")
        
        st.markdown("---")
        st.markdown("**🔍 故障诊断工具**")
        if st.button("🌐 诊断服务器网络与授权", use_container_width=True):
            with st.spinner("正在进行多维度网络连通性与授权诊断..."):
                st.markdown("#### 1. DNS 解析测试")
                import socket
                try:
                    ip = socket.gethostbyname("mcp.zsxq.com")
                    st.success(f"✅ DNS 解析成功: `mcp.zsxq.com` -> `{ip}`")
                except Exception as e:
                    st.error(f"❌ DNS 解析失败: {str(e)}")
                
                st.markdown("#### 2. Curl 详细连接握手诊断 (TLS/WAF)")
                curl_cmd = "curl.exe" if sys.platform == "win32" else "curl"
                try:
                    res = subprocess.run([curl_cmd, "-iv", "-X", "POST", "https://mcp.zsxq.com/topic/mcp"], capture_output=True, text=True, timeout=15)
                    st.text("执行命令: curl -iv -X POST https://mcp.zsxq.com/topic/mcp")
                    st.code(f"STDOUT:\n{res.stdout}\n\nSTDERR:\n{res.stderr}", language="bash")
                except Exception as e:
                    st.error(f"❌ Curl 执行异常: {str(e)}")
                    
                st.markdown("#### 3. CLI 授权状态诊断")
                try:
                    res_cli = subprocess.run([NPX_CMD, "zsxq-cli", "auth", "status", "--json"], capture_output=True, text=True, encoding='utf-8', timeout=15)
                    st.code(res_cli.stdout or res_cli.stderr, language="json")
                except Exception as e:
                    st.error(f"❌ CLI 诊断异常: {str(e)}")

        st.markdown("---")
        # Dynamic groups based on userId
        if logged_in and user_id:
            if "user_groups" not in config: config["user_groups"] = {}
            if user_id not in config["user_groups"]:
                config["user_groups"][user_id] = config.get("groups", {"默认群组": ""})
                save_config(config)
            groups_dict = config["user_groups"][user_id]
        else:
            groups_dict = {"默认群组": ""}
            
        group_keys = list(groups_dict.keys())
        if not group_keys:
            groups_dict = {"默认群组": ""}
            group_keys = ["默认群组"]
            
        if config.get("selected_group") not in group_keys:
            config["selected_group"] = group_keys[0]

        sel_g_name = st.selectbox("选择星球/群组", group_keys, index=group_keys.index(config["selected_group"]))
        if sel_g_name != config["selected_group"]:
            config["selected_group"] = sel_g_name
            save_config(config)
            
        curr_group_id = groups_dict[sel_g_name]
        
        if logged_in and user_id:
            st.markdown("---")
            st.markdown("**➕ 星群管理**")
            st.markdown("**新增或修改群组**")
            ng_name = st.text_input("星球名称", placeholder="例如：阿铭linux")
            ng_id = st.text_input("Group ID", placeholder="星球数字ID")
            if st.button("💾 保存/更新群组"):
                if ng_name and ng_id:
                    config["user_groups"][user_id][ng_name] = ng_id
                    config["selected_group"] = ng_name
                    save_config(config)
                    st.success("更新成功！")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("名称和ID不能为空！")
                    
            st.markdown("---")
            st.markdown("**删除群组**")
            del_g_name = st.selectbox("选择要删除的群组", group_keys)
            if st.button("🗑️ 删除选中群组"):
                if del_g_name in config["user_groups"][user_id]:
                    del config["user_groups"][user_id][del_g_name]
                    if config["selected_group"] == del_g_name:
                        config["selected_group"] = list(config["user_groups"][user_id].keys())[0] if config["user_groups"][user_id] else ""
                    save_config(config)
                    st.success("删除成功！")
                    time.sleep(0.5)
                    st.rerun()

    # Expander 2: 🎨 绘图引擎配置
    with st.expander("🎨 绘图引擎配置", expanded=False):
        img_gen_ops = ["即梦 (Dreamina)", "Google Gemini (Imagen 3)"]
        prev_img_gen = config.get("image_generator", "即梦 (Dreamina)")
        selected_img_gen = st.selectbox("图片生成引擎", img_gen_ops, index=img_gen_ops.index(prev_img_gen) if prev_img_gen in img_gen_ops else 0)
        
        if selected_img_gen != prev_img_gen:
            config["image_generator"] = selected_img_gen
            save_config(config)
            st.rerun()

        if selected_img_gen == "Google Gemini (Imagen 3)":
            google_api_key_val = config.get("google_api_key", "")
            new_google_key = st.text_input("Google API Key", value=google_api_key_val, type="password")
            
            model_ops = ["imagen-4.0-generate-001", "imagen-4.0-ultra-generate-001", "imagen-4.0-fast-generate-001", "gemini-3.1-flash-image"]
            prev_model = config.get("gemini_image_model", "imagen-4.0-generate-001")
            selected_model = st.selectbox("Gemini 绘图模型", model_ops, index=model_ops.index(prev_model) if prev_model in model_ops else 0)
            
            aspect_ratio_ops = ["1:1", "4:3", "16:9", "3:4", "9:16"]
            prev_aspect_ratio = config.get("image_aspect_ratio", "1:1")
            selected_ratio = st.selectbox("图片比例", aspect_ratio_ops, index=aspect_ratio_ops.index(prev_aspect_ratio) if prev_aspect_ratio in aspect_ratio_ops else 0)
            
            if new_google_key != google_api_key_val or selected_model != prev_model or selected_ratio != prev_aspect_ratio:
                config["google_api_key"] = new_google_key
                config["gemini_image_model"] = selected_model
                config["image_aspect_ratio"] = selected_ratio
                save_config(config)
                
            if new_google_key:
                st.success("✅ Google Gemini 已配置")
            else:
                st.warning("⚠️ Google Gemini 未配置 API Key")
        else:
            d_returncode, d_stdout = get_dreamina_credit_status()
            dreamina_logged_in = (d_returncode == 0)
            dreamina_credit_info = ""
            
            if dreamina_logged_in:
                try:
                    credit_data = json.loads(d_stdout)
                    dreamina_credit_info = f"剩余积分: {credit_data.get('total_credit', '未知')}"
                except Exception:
                    dreamina_credit_info = d_stdout.strip()
                st.success(f"✅ 即梦已授权 ({dreamina_credit_info})")
                if st.button("退出即梦登录"): 
                    subprocess.run([DREAMINA_CMD, "logout"])
                    clear_auth_cache()
                    st.rerun()
            else:
                st.warning("⚠️ 即梦未授权")
                if st.button("🔗 获取即梦授权链接"):
                    login_res = subprocess.run([DREAMINA_CMD, "login", "--headless"], capture_output=True, text=True, encoding='utf-8')
                    m_uri = re.search(r'verification_uri:\s*(https://\S+)', login_res.stdout)
                    m_code = re.search(r'user_code:\s*(\w+)', login_res.stdout)
                    m_device = re.search(r'device_code:\s*(\w+)', login_res.stdout)
                    if m_uri and m_code and m_device:
                        st.session_state.dreamina_device_code = m_device.group(1)
                        st.markdown(f"**[👉 点击授权]({m_uri.group(1)})**")
                        st.info(f"确认码：`{m_code.group(1)}`")
                    else:
                        st.error("无法获取链接")
                if "dreamina_device_code" in st.session_state and st.button("我已完成即梦授权"):
                    with st.spinner("验证即梦授权中（最多等待30秒）..."):
                        verify_res = subprocess.run(
                            [DREAMINA_CMD, "login", "checklogin",
                             f"--device_code={st.session_state.dreamina_device_code}",
                             "--poll=30"],
                            capture_output=True, text=True, encoding='utf-8'
                        )
                        combined_output = verify_res.stdout + verify_res.stderr
                        if verify_res.returncode == 0 or "LOGIN_SUCCESS" in combined_output or "登录成功" in combined_output or "Successfully" in combined_output:
                            st.success("授权成功！")
                            del st.session_state.dreamina_device_code
                            clear_auth_cache()
                            st.rerun()
                        else:
                            st.error("未检测到成功授权，请确认您已在网页端登录并确认！")

    # Expander 3: 🤖 AI 渠道配置
    with st.expander("🤖 AI 渠道配置", expanded=False):
        plat_ops = ["自定义/OpenAI", "火山方舟 (Volcengine)", "魔塔 (ModelScope)", "DeepSeek"]
        prev_platform = config.get("platform", "自定义/OpenAI")
        
        selected_platform = st.selectbox("AI 渠道", plat_ops, index=plat_ops.index(prev_platform) if prev_platform in plat_ops else 0)
        if selected_platform != prev_platform:
            config["platform"] = selected_platform
            save_config(config)
            st.rerun()
            
        current_chan_config = config["channel_configs"].get(selected_platform, {})
        
        new_api_key = st.text_input("API Key", value=current_chan_config.get("api_key", ""), type="password")
        new_base_url = st.text_input("Base URL", value=current_chan_config.get("base_url", ""))
        
        if new_api_key != current_chan_config.get("api_key") or new_base_url != current_chan_config.get("base_url"):
            config["channel_configs"][selected_platform]["api_key"] = new_api_key
            config["channel_configs"][selected_platform]["base_url"] = new_base_url
            save_config(config)
        
        if st.button("🔄 获取模型列表") and current_chan_config.get("api_key") and current_chan_config.get("base_url"):
            with st.spinner("获取中..."):
                try:
                    h = {"Authorization": f"Bearer {current_chan_config['api_key']}"}
                    r = httpx.get(f"{current_chan_config['base_url'].rstrip('/')}/models", headers=h, timeout=10)
                    if r.status_code == 200:
                        models = [m["id"] for m in r.json().get("data", [])]
                        config["channel_configs"][selected_platform]["available_models"] = models
                        save_config(config)
                        st.success(f"获取成功！共获取到 {len(models)} 个可用模型。")
                    else:
                        st.error(f"获取失败，HTTP 状态码: {r.status_code}")
                except Exception as e:
                    st.error(f"获取异常: {str(e)}")

        all_mods = list(dict.fromkeys(current_chan_config.get("available_models", []) + config.get("manual_models", [])))
        if not all_mods: all_mods = ["gpt-4o"]
        
        new_model = st.selectbox("当前模型", all_mods, index=all_mods.index(current_chan_config.get("selected_model")) if current_chan_config.get("selected_model") in all_mods else 0)
        if new_model != current_chan_config.get("selected_model"):
            config["channel_configs"][selected_platform]["selected_model"] = new_model
            save_config(config)

        if selected_platform == "DeepSeek":
            enable_thinking = st.toggle("开启思考模式 (Thinking Mode)", value=current_chan_config.get("enable_thinking", True))
            reasoning_effort = st.selectbox("思考强度 (Reasoning Effort)", ["high", "max"], index=0 if current_chan_config.get("reasoning_effort", "high") == "high" else 1)
            if enable_thinking != current_chan_config.get("enable_thinking", True) or reasoning_effort != current_chan_config.get("reasoning_effort", "high"):
                config["channel_configs"][selected_platform]["enable_thinking"] = enable_thinking
                config["channel_configs"][selected_platform]["reasoning_effort"] = reasoning_effort
                save_config(config)

    # Expander 4: ⏰ 定时任务 (Cron)
    # Expander 4: ⏰ 定时任务 (Cron)
    with st.expander("⏰ 定时任务 (Cron)", expanded=False):
        if "schedulers" not in config:
            config["schedulers"] = {}
            
        # Migrate old config if present
        if "auto_run" in config and "cron_expr" in config and not config["schedulers"]:
            config["schedulers"]["AI 深度分析"] = {
                "auto_run": config.get("auto_run", False),
                "cron_expr": config.get("cron_expr", "0 8 * * *"),
                "ui_state": {}
            }
            
        sched_config = config["schedulers"].setdefault(page_selection, {"auto_run": False, "cron_expr": "0 8 * * *", "ui_state": {}})
        
        curr_cron = sched_config.get("cron_expr", "0 8 * * *")
        new_cron = st.text_input("✏️ 手动输入 Cron 表达式", value=curr_cron, key=f"cron_input_{page_selection}")
        
        if st.button("✨ 打开 Cron 可视化配置器", use_container_width=True):
            cron_configurator_dialog(page_selection, new_cron)
            
        try:
            from apscheduler.triggers.cron import CronTrigger
            from datetime import datetime
            trigger = CronTrigger.from_crontab(new_cron)
            now = datetime.now()
            runs = []
            from datetime import timedelta
            for _ in range(5):
                now = trigger.get_next_fire_time(None, now)
                runs.append(now.strftime("%Y-%m-%d %H:%M:%S"))
                now = now + timedelta(seconds=1)
            st.info("🔮 **主界面实时预演 (未来 5 次):**\n\n" + "\n".join([f"- {r}" for r in runs]))
        except Exception:
            st.error("❌ 表达式无效，无法计算执行时间。")
            
        new_auto_run = st.toggle("开启定时自动运行", value=sched_config.get("auto_run", False), key=f"auto_run_{page_selection}")
        
        if new_auto_run:
            st.success("🟢 定时任务已激活并运行中")
        else:
            st.warning("⚪ 定时任务未开启")
            
        col_save, col_test, col_log = st.columns(3)
        with col_save:
            if st.button("💾 保存当前配置", use_container_width=True):
                # Capture UI state snapshot
                ui_state = {}
                if page_selection == "AI 深度分析":
                    ui_state = {
                        "scope_ui": st.session_state.get("da_scope_ui", "最新总结 (话题+文件)"),
                        "l_limit": st.session_state.get("da_l_limit", 3),
                        "a_mode": st.session_state.get("da_a_mode", "常规总结"),
                        "use_p_ui": st.session_state.get("da_use_p_ui", False),
                        "cp_text": st.session_state.get("da_cp_text", ""),
                        "use_wechat": st.session_state.get("da_use_wechat", False),
                        "also_generate_report": st.session_state.get("da_also_generate_report", False),
                        "wechat_mode": st.session_state.get("da_wechat_mode", "AI 模式 (免费)"),
                        "wechat_theme": st.session_state.get("da_wechat_theme", "spring-fresh"),
                        "wechat_custom_prompt": st.session_state.get("da_wechat_custom_prompt", ""),
                        "wechat_font_size": st.session_state.get("da_wechat_font_size", "medium"),
                        "wechat_background_type": st.session_state.get("da_wechat_background_type", "none"),
                        "wechat_prompt": st.session_state.get("da_wechat_prompt", ""),
                        "wechat_publish_mode": st.session_state.get("da_wechat_publish_mode", "仅生成本地文件 (不上传)"),
                        "wechat_publish_account": st.session_state.get("da_wechat_publish_account", ""),
                        "selected_group": config.get("selected_group", "默认群组")
                    }
                elif page_selection == "视频脚本制作器":
                    v_hist_serial = []
                    for vf in st.session_state.get("virtual_history", []):
                        v_hist_serial.append({"name": vf["name"], "text": vf["text"]})
                    ui_state = {
                        "script_mode": st.session_state.get("vs_script_mode", "仿写现有格式生成新脚本"),
                        "export_format": st.session_state.get("vs_export_format", ".docx"),
                        "prompt_input": st.session_state.get("vs_prompt_input", ""),
                        "virtual_history": v_hist_serial
                    }
                elif page_selection == "指标文档制作":
                    ui_state = {
                        "selected_indicator": st.session_state.get("ind_selected", "")
                    }
                    
                config["schedulers"][page_selection] = {
                    "auto_run": new_auto_run,
                    "cron_expr": new_cron,
                    "ui_state": ui_state
                }
                save_config(config)
                update_scheduler()
                st.success("配置与定时任务已保存！")
                time.sleep(1)
                st.rerun()
                
        with col_test:
            if st.button("⚡ 测试运行任务", use_container_width=True):
                import threading
                def test_runner():
                    if page_selection == "AI 深度分析":
                        run_scheduled_deep_analysis()
                    elif page_selection == "视频脚本制作器":
                        run_scheduled_video_script()
                    elif page_selection == "指标文档制作":
                        run_scheduled_indicator_docs()
                
                threading.Thread(target=test_runner, daemon=True).start()
                st.success("⚡ 测试任务已在后台启动！请稍后点击『查看执行日志』关注运行进展。")
                
        with col_log:
            if st.button("📋 查看执行日志", use_container_width=True):
                log_path = os.path.join("outputs", "cron_execution.log")
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8") as f:
                        log_content = f.read()
                else:
                    log_content = "暂无执行日志。"
                if len(log_content) > 30000:
                    log_content = "...(已截断历史日志)...\n" + log_content[-30000:]
                if hasattr(st, "dialog"):
                    show_log_dialog(log_content, page_selection)
                else:
                    st.session_state["show_log_fallback"] = log_content
                    st.rerun()

    if st.button("💾 手动保存所有配置", use_container_width=True): save_config(config); st.success("已保存")

# Main Area
if "show_log_fallback" in st.session_state:
    with st.container(border=True):
        st.subheader("📋 定时任务执行日志")
        st.text_area("", value=st.session_state["show_log_fallback"], height=450, disabled=True)
        if st.button("关闭日志"):
            del st.session_state["show_log_fallback"]
            st.rerun()

if page_selection == "AI 深度分析":
    st.title("🤖 AI 深度分析控制台")
    col1, col2 = st.columns([2, 1.3])
    with col1:
        st.subheader("📊 分析模式与数据源")
        scope_ui = st.radio("获取范围", ["最新总结 (话题+文件)", "文件总结 (仅限附件)"], horizontal=True, key="da_scope_ui")
        s_key = "all" if "最新" in scope_ui else "files"
        l_limit = st.number_input("获取近多少条消息", min_value=1, max_value=100, value=3, key="da_l_limit")
        a_mode = st.radio("分析模式", ["常规总结", "个股分析", "行业分析"], horizontal=True, key="da_a_mode")
        
        st.markdown("---")
        st.subheader("✨ 个性化输出处理")
        use_p_ui = st.checkbox("总结后进行个性化二次加工", key="da_use_p_ui")
        cp_text = ""
        if use_p_ui:
            cp_ops = ["自定义输入"] + list(config["custom_prompts"].keys())
            sp_name = st.selectbox("选择预设", cp_ops)
            
            if sp_name == "自定义输入":
                cp_text = st.text_area("指令", key="da_cp_text")
                np_name = st.text_input("预设名称")
                if st.button("💾 保存预设", use_container_width=True) and cp_text and np_name: 
                    config["custom_prompts"][np_name] = cp_text
                    save_config(config)
                    st.rerun()
            else:
                cp_text = st.text_area("编辑指令", value=config["custom_prompts"][sp_name], height=120, key="da_cp_text")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("🔄 更新当前预设", use_container_width=True):
                        config["custom_prompts"][sp_name] = cp_text
                        save_config(config)
                        st.success("更新成功！")
                        time.sleep(1)
                        st.rerun()
                with col_btn2:
                    if st.button("🗑️ 删除当前预设", use_container_width=True):
                        st.session_state['delete_confirm'] = sp_name
                
                if st.session_state.get('delete_confirm') == sp_name:
                    with st.container(border=True):
                        st.warning(f"⚠️ 确定要删除预设【{sp_name}】吗？")
                        del_col1, del_col2 = st.columns(2)
                        if del_col1.button("✅ 确认删除", type="primary", use_container_width=True):
                            del config["custom_prompts"][sp_name]
                            save_config(config)
                            del st.session_state['delete_confirm']
                            st.rerun()
                        if del_col2.button("❌ 取消", use_container_width=True):
                            del st.session_state['delete_confirm']
                            st.rerun()

        st.markdown("---")
        st.subheader("📱 微信公众号推文生成")
        use_wechat = st.checkbox("生成微信公众号推文 (跳过常规图表渲染加速)", value=config.get("use_wechat", False), key="da_use_wechat")
        if use_wechat != config.get("use_wechat", False):
            config["use_wechat"] = use_wechat
            save_config(config)
            st.rerun()

        also_generate_report = False
        if use_wechat:
            also_generate_report = st.checkbox("同时生成分析报告", value=config.get("also_generate_report", False), key="da_also_generate_report")
            if also_generate_report != config.get("also_generate_report", False):
                config["also_generate_report"] = also_generate_report
                save_config(config)
                st.rerun()
                
            # 定时自动发布模式
            wechat_pub_modes = ["仅生成本地文件 (不上传)", "自动保存至微信草稿箱", "自动保存草稿并正式发布"]
            prev_pub_mode = config.get("wechat_publish_mode", "仅生成本地文件 (不上传)")
            selected_pub_mode = st.selectbox(
                "定时任务微信发布模式",
                wechat_pub_modes,
                index=wechat_pub_modes.index(prev_pub_mode) if prev_pub_mode in wechat_pub_modes else 0,
                key="da_wechat_publish_mode"
            )
            
            # 微信发布目标账号
            accounts = wechat_publisher.load_accounts()
            account_names = [a["name"] for a in accounts]
            prev_pub_acc = config.get("wechat_publish_account", account_names[0] if account_names else "")
            
            selected_pub_acc = st.selectbox(
                "微信发布目标账号",
                account_names if account_names else ["无可用账号"],
                index=account_names.index(prev_pub_acc) if (account_names and prev_pub_acc in account_names) else 0,
                key="da_wechat_publish_account"
            )
            
            if selected_pub_mode != prev_pub_mode or selected_pub_acc != prev_pub_acc:
                config["wechat_publish_mode"] = selected_pub_mode
                config["wechat_publish_account"] = selected_pub_acc
                save_config(config)

        wechat_prompt = ""
        if use_wechat:
            col_mode, col_theme = st.columns(2)
            with col_mode:
                wechat_mode_ops = ["AI 模式 (免费)", "API 模式 (专业)"]
                prev_wechat_mode = config.get("wechat_mode", "AI 模式 (免费)")
                selected_wechat_mode = st.selectbox(
                    "排版模式", 
                    wechat_mode_ops, 
                    index=wechat_mode_ops.index(prev_wechat_mode) if prev_wechat_mode in wechat_mode_ops else 0,
                    key="da_wechat_mode"
                )
            
            is_api_mode = (selected_wechat_mode == "API 模式 (专业)")
            if is_api_mode:
                theme_ops = [
                    "default", "bytedance", "apple", "sspai-red", "wechat-native", 
                    "nyt-classic", "sunset-amber", "mint-fresh", "lavender-dream",
                    "elegant-gold", "elegant-green", "elegant-blue", "elegant-red",
                    "focus-gold", "focus-green", "focus-blue", "focus-red",
                    "minimal-gold", "minimal-green", "minimal-blue", "minimal-red",
                    "bold-gold", "bold-green", "bold-blue", "bold-red",
                    "chinese", "cyber", "sports"
                ]
            else:
                theme_ops = ["spring-fresh", "autumn-warm", "ocean-calm", "custom"]
                
            with col_theme:
                prev_wechat_theme = config.get("wechat_theme", "spring-fresh" if not is_api_mode else "default")
                if prev_wechat_theme not in theme_ops:
                    prev_wechat_theme = theme_ops[0]
                selected_wechat_theme = st.selectbox(
                    "排版主题",
                    theme_ops,
                    index=theme_ops.index(prev_wechat_theme),
                    key="da_wechat_theme"
                )

            md2wechat_api_key = config.get("md2wechat_api_key", "")
            if is_api_mode:
                md2wechat_api_key = st.text_input("md2wechat API Key", value=md2wechat_api_key, type="password", placeholder="填入 md2wechat.cn 专属 API Key")
            
            wechat_custom_prompt = config.get("wechat_custom_prompt", "")
            if not is_api_mode and selected_wechat_theme == "custom":
                wechat_custom_prompt = st.text_area("自定义 AI 排版 Prompt", value=wechat_custom_prompt, placeholder="输入您自定义的 CSS/HTML 设计提示词...", key="da_wechat_custom_prompt")

            with st.expander("⚙️ 微信高级排版选项"):
                prev_font_size = config.get("wechat_font_size", "medium")
                font_size_ops = ["medium", "small", "large"]
                selected_font_size = st.selectbox("正文字号", font_size_ops, index=font_size_ops.index(prev_font_size), key="da_wechat_font_size")
                
                prev_bg_type = config.get("wechat_background_type", "none")
                bg_type_ops = ["none", "default", "grid"]
                selected_bg_type = st.selectbox("背景类型", bg_type_ops, index=bg_type_ops.index(prev_bg_type), key="da_wechat_background_type")

            if (selected_wechat_mode != prev_wechat_mode or 
                selected_wechat_theme != prev_wechat_theme or 
                md2wechat_api_key != config.get("md2wechat_api_key", "") or
                wechat_custom_prompt != config.get("wechat_custom_prompt", "") or
                selected_font_size != prev_font_size or
                selected_bg_type != prev_bg_type):
                
                config["wechat_mode"] = selected_wechat_mode
                config["wechat_theme"] = selected_wechat_theme
                config["md2wechat_api_key"] = md2wechat_api_key
                config["wechat_custom_prompt"] = wechat_custom_prompt
                config["wechat_font_size"] = selected_font_size
                config["wechat_background_type"] = selected_bg_type
                save_config(config)
            
            wechat_prompt = st.text_area("推文个性化要求 (可选)", placeholder="例如：语气更加诙谐幽默，重点强调端侧存储，多用短句...", height=68, key="da_wechat_prompt")

        if st.button("🚀 立即开始深度分析", use_container_width=True, type="primary"):
            st.session_state['last_execution_log'] = [] # Reset old logs
            if not current_chan_config.get("api_key"):
                st.error("请先在左侧全局配置中填写 API Key！")
                st.stop()
                
            log_container = st.container()
            with log_container:
                st.write("### 🔄 任务执行日志")
                log_messages = []
                with st.status("🚀 深度分析任务执行中...", expanded=True) as main_status:
                    # Patch main_status.write to collect logs
                    original_write = main_status.write
                    def patched_write(msg):
                        original_write(msg)
                        log_messages.append(msg)
                    main_status.write = patched_write
                    
                    def update_progress(event):
                        if event['type'] == 'info':
                            main_status.write(f"ℹ️ {event['msg']}")
                        elif event['type'] == 'topic_start':
                            preview_text = re.sub(r'<[^>]+>|<[^>]*$', '', event['preview'])
                            main_status.write(f"⏳ **处理话题**: `{preview_text}`")
                        elif event['type'] == 'topic_log':
                            main_status.write(f"　 └─ {event['msg']}")
                        elif event['type'] == 'topic_end':
                            if event.get('success'):
                                main_status.write(f"　 └─ ✅ **完成**")
                            else:
                                reason = event.get('reason', '')
                                main_status.write(f"　 └─ ⏭️ **跳过** ({reason})")
                    
                    raw, f_list, briefs = fetch_zsxq(curr_group_id, limit=l_limit, scope=s_key, progress_callback=update_progress)
                    
                    if raw.startswith("获取失败") or raw.startswith("获取异常") or raw.startswith("跳过分析") or (not briefs and not f_list):
                        st.session_state['last_execution_log'] = log_messages
                        main_status.update(label="❌ 任务异常终止", state="error", expanded=True)
                        st.error(f"分析无法继续: {raw}")
                        st.stop()

                    main_status.write(f"🧠 **向大模型发起请求** (模型: `{current_chan_config.get('selected_model')}`)...")
                    main_status.write("　 └─ 正在构建上下文 Prompt...")
                    res_main = generate_summary(raw, current_chan_config.get("api_key"), current_chan_config.get("base_url"), current_chan_config.get("selected_model"), a_mode, chan_config=current_chan_config)
                    
                    if res_main.startswith("AI 总结失败") or res_main.startswith("未提供"):
                        st.session_state['last_execution_log'] = log_messages
                        main_status.update(label="❌ 任务异常终止", state="error", expanded=True)
                        st.error(res_main)
                        st.stop()

                    final_res = res_main
                    if use_p_ui and cp_text:
                        main_status.write("✨ **正在执行个性化二次美化加工**...")
                        final_res = generate_summary(res_main, current_chan_config.get("api_key"), current_chan_config.get("base_url"), current_chan_config.get("selected_model"), a_mode, custom_prompt=cp_text, chan_config=current_chan_config)
                        if final_res.startswith("AI 总结失败"):
                            st.session_state['last_execution_log'] = log_messages
                            main_status.update(label="❌ 二次美化加工异常", state="error", expanded=True)
                            st.error(final_res)
                            st.stop()
                            
                    if use_wechat:
                        main_status.write("📱 **正在创作微信公众号推文** (包含智能配图)...")
                        wechat_system_prompt = get_wechat_system_prompt()
                        
                        wechat_user_content = f"【基础分析总结】\n{final_res}\n"
                        if wechat_prompt.strip():
                            wechat_user_content += f"\n【用户个性化要求】\n{wechat_prompt}"
                            
                        try:
                            from openai import OpenAI
                            client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                            wc_response = call_chat_completion(client, current_chan_config.get("selected_model"), [{"role": "system", "content": wechat_system_prompt}, {"role": "user", "content": wechat_user_content}], chan_config=current_chan_config)
                            if wc_response.choices:
                                raw_wechat = wc_response.choices[0].message.content
                                raw_wechat = adjust_markdown_images_placement(raw_wechat)
                                import re
                                def replace_img(match):
                                    kw = match.group(1)
                                    image_engine = config.get("image_generator", "即梦 (Dreamina)")
                                    if image_engine == "Google Gemini (Imagen 3)":
                                        main_status.write(f"🖼️ 正在调用 Gemini 生成配图: `{kw}`")
                                        img_path = generate_gemini_image(
                                            kw, 
                                            config.get("google_api_key", ""), 
                                            model_name=config.get("gemini_image_model", "imagen-4.0-generate-001"),
                                            aspect_ratio=config.get("image_aspect_ratio", "1:1")
                                        )
                                    else:
                                        main_status.write(f"🖼️ 正在调用即梦生成配图: `{kw}`")
                                        img_path = generate_jimeng_image(kw)
                                    if img_path:
                                        # Convert path to forward slashes to ensure markdown compatibility
                                        img_path = img_path.replace("\\", "/")
                                        return f"({img_path})"
                                    else:
                                        return "(https://dummyimage.com/800x400/ffebee/d32f2f.png&text=Image+Generate+Failed)"
                                    
                                wechat_res = re.sub(r'\(\[IMAGE_GENERATE:(.*?)\]\)', replace_img, raw_wechat)
                                st.session_state['latest_wechat'] = wechat_res
                                if 'wechat_long_image' in st.session_state:
                                    del st.session_state['wechat_long_image']
                                # 切换/生成新文章时重置内存中的微信草稿状态，防止交叉污染
                                st.session_state.wechat_draft_media_id = None
                                st.session_state.wechat_draft_url = None
                                st.session_state.wechat_publish_result = None
                                st.session_state.wechat_preview_status = None
                                
                                # 保存历史记录
                                wc_filename = f"wechat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                                wc_path = os.path.join(WECHAT_OUTPUT_DIR, wc_filename)
                                with open(wc_path, "w", encoding="utf-8") as f:
                                    f.write(wechat_res)
                                    
                                st.session_state['latest_wechat_file_path'] = wc_path
                                    
                                # 调用 md2wechat 美化排版
                                main_status.write("🎨 **正在使用 md2wechat 进行样式美化排版**...")
                                theme = config.get("wechat_theme", "spring-fresh" if "AI" in config.get("wechat_mode", "AI 模式") else "default")
                                mode_ui = config.get("wechat_mode", "AI 模式 (免费)")
                                api_key = config.get("md2wechat_api_key", "")
                                font_size = config.get("wechat_font_size", "medium")
                                bg_type = config.get("wechat_background_type", "none")
                                custom_prompt = config.get("wechat_custom_prompt", "")
                                
                                html_res = convert_to_wechat_html(
                                    wechat_res, 
                                    theme, 
                                    mode_ui, 
                                    api_key=api_key, 
                                    font_size=font_size, 
                                    bg_type=bg_type, 
                                    chan_config=current_chan_config, 
                                    custom_prompt=custom_prompt
                                )
                                
                                st.session_state['latest_wechat_html'] = html_res
                                st.session_state['latest_wechat_rendered_theme'] = theme
                                st.session_state['latest_wechat_rendered_mode'] = mode_ui
                                
                                # 保存美化后的 HTML
                                html_filename = wc_filename.replace(".md", ".html")
                                html_path = os.path.join(WECHAT_OUTPUT_DIR, html_filename)
                                with open(html_path, "w", encoding="utf-8") as f:
                                    f.write(html_res)

                            else:
                                main_status.write("⚠️ 微信公众号推文生成返回为空")
                        except Exception as e:
                            main_status.write(f"⚠️ 微信公众号推文生成失败: {str(e)}")

                    also_generate_report = config.get("also_generate_report", False)
                    if not use_wechat or also_generate_report:
                        main_status.write("🎨 正在渲染专业分析报表...")
                        img_p = render_to_image(final_res, a_mode)
                        st.session_state['latest_img'] = img_p
                        
                    st.session_state['last_execution_log'] = log_messages
                    main_status.update(label="🎉 任务完成！", state="complete", expanded=True)

            if not use_wechat:
                st.success("分析报表已生成！请在右侧预览。")
            elif also_generate_report:
                st.success("微信公众号推文与分析报表已生成！请在右侧预览。")
            else:
                st.success("微信公众号推文已生成！请在右侧预览。")
            time.sleep(1)
            st.rerun()

        if st.session_state.get('last_execution_log'):
            st.write("---")
            with st.expander("📋 上次深度分析执行日志", expanded=True):
                for log_line in st.session_state['last_execution_log']:
                    st.markdown(log_line)

    with col2:
        tab1, tab2 = st.tabs(["📊 分析报告预览", "📱 公众号推文预览"])
        
        with tab1:
            if 'latest_img' in st.session_state and os.path.exists(st.session_state['latest_img']):
                st.image(st.session_state['latest_img'], use_container_width=True)
            
            st.write("---")
            st.subheader("📅 历史报告归档")
            hist_imgs = sorted(glob.glob(os.path.join(IMAGE_OUTPUT_DIR, "*.png")), reverse=True)
            if not hist_imgs: st.info("暂无历史记录")
            else:
                d_groups = {}
                for h_im in hist_imgs:
                    try:
                        ds = os.path.basename(h_im).split('_')[1][:8]
                        dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                    except: dfmt = "其他"
                    if dfmt not in d_groups: d_groups[dfmt] = []
                    d_groups[dfmt].append(h_im)
                
                for d_key, d_ims in d_groups.items():
                    with st.expander(f"📅 {d_key} ({len(d_ims)}份)"):
                        for single_im in d_ims:
                            st.image(single_im, caption=os.path.basename(single_im))
                            
        with tab2:
            if st.session_state.get('latest_wechat'):
                # 初始化 selectbox 状态以防止切换两次的冲突问题
                if 'preview_wechat_mode_select' not in st.session_state:
                    st.session_state['preview_wechat_mode_select'] = st.session_state.get('latest_wechat_rendered_mode', config.get("wechat_mode", "AI 模式 (免费)"))
                if 'preview_wechat_theme_select' not in st.session_state:
                    st.session_state['preview_wechat_theme_select'] = st.session_state.get('latest_wechat_rendered_theme', config.get("wechat_theme", "spring-fresh"))
                
                # 🎨 实时排版主题/模式切换面板
                st.markdown("### 🎨 实时排版主题切换")
                col_preview_mode, col_preview_theme = st.columns(2)
                with col_preview_mode:
                    preview_mode_ops = ["AI 模式 (免费)", "API 模式 (专业)"]
                    val_mode = st.session_state['preview_wechat_mode_select']
                    if val_mode not in preview_mode_ops:
                        val_mode = preview_mode_ops[0]
                    preview_mode = st.selectbox(
                        "预览排版模式", 
                        preview_mode_ops, 
                        key="preview_wechat_mode_select",
                        index=preview_mode_ops.index(val_mode)
                    )
                
                is_preview_api = "API" in preview_mode
                if is_preview_api:
                    preview_theme_ops = [
                        "default", "bytedance", "apple", "sspai-red", "wechat-native", 
                        "nyt-classic", "sunset-amber", "mint-fresh", "lavender-dream",
                        "elegant-gold", "elegant-green", "elegant-blue", "elegant-red",
                        "focus-gold", "focus-green", "focus-blue", "focus-red",
                        "minimal-gold", "minimal-green", "minimal-blue", "minimal-red",
                        "bold-gold", "bold-green", "bold-blue", "bold-red",
                        "chinese", "cyber", "sports"
                    ]
                else:
                    preview_theme_ops = ["spring-fresh", "autumn-warm", "ocean-calm", "custom"]
                    
                with col_preview_theme:
                    val_theme = st.session_state['preview_wechat_theme_select']
                    if val_theme not in preview_theme_ops:
                        val_theme = preview_theme_ops[0]
                    preview_theme = st.selectbox(
                        "预览排版主题",
                        preview_theme_ops,
                        key="preview_wechat_theme_select",
                        index=preview_theme_ops.index(val_theme)
                    )
                
                # 响应式渲染检查
                if st.session_state.get('just_loaded_history'):
                    st.session_state['latest_wechat_rendered_theme'] = preview_theme
                    st.session_state['latest_wechat_rendered_mode'] = preview_mode
                    st.session_state['just_loaded_history'] = False
                
                need_re_render = False
                if not st.session_state.get('latest_wechat_html'):
                    need_re_render = True
                elif st.session_state.get('latest_wechat_rendered_theme') != preview_theme:
                    need_re_render = True
                elif st.session_state.get('latest_wechat_rendered_mode') != preview_mode:
                    need_re_render = True
                    
                if need_re_render:
                    with st.spinner("🔄 正在为当前推文切换主题并重新排版..."):
                        font_size = config.get("wechat_font_size", "medium")
                        bg_type = config.get("wechat_background_type", "none")
                        api_key = config.get("md2wechat_api_key", "")
                        custom_prompt = config.get("wechat_custom_prompt", "")
                        
                        html_res = convert_to_wechat_html(
                            st.session_state['latest_wechat'], 
                            preview_theme, 
                            preview_mode, 
                            api_key=api_key, 
                            font_size=font_size, 
                            bg_type=bg_type, 
                            chan_config=current_chan_config, 
                            custom_prompt=custom_prompt
                        )
                        st.session_state['latest_wechat_html'] = html_res
                        st.session_state['latest_wechat_rendered_theme'] = preview_theme
                        st.session_state['latest_wechat_rendered_mode'] = preview_mode
                        
                        if 'wechat_long_image' in st.session_state:
                            del st.session_state['wechat_long_image']
                            
                        file_path = st.session_state.get('latest_wechat_file_path')
                        if file_path:
                            html_p = file_path.replace(".md", ".html")
                            try:
                                with open(html_p, "w", encoding="utf-8") as hf:
                                    hf.write(html_res)
                            except:
                                pass
                        
                        st.rerun()
                
                st.write("---")
                col_btn_docx, col_btn_img, col_btn_wechat = st.columns(3)
                with col_btn_docx:

                    st.download_button(
                        label="📄 导出推文为 Docx 文档",
                        data=markdown_to_wechat_docx_bytes(st.session_state['latest_wechat']),
                        file_name=f"微信公众号推文_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary"
                    )
                with col_btn_img:
                    if st.button("📸 导出推文为长图", use_container_width=True):
                        if st.session_state.get('latest_wechat_html'):
                            with st.spinner("正在启动 Playwright 渲染生成高保真长图..."):
                                try:
                                    img_path = generate_wechat_long_image(st.session_state['latest_wechat_html'])
                                    st.session_state['wechat_long_image'] = img_path
                                    st.success("长图生成成功！请在下方查看并下载。")
                                except Exception as e:
                                    st.error(f"长图生成失败: {str(e)}")
                        else:
                            st.warning("暂无已美化的推文 HTML 样式，请重新运行分析。")
                with col_btn_wechat:
                    if st.button("📲 公众号一键发布/预览", use_container_width=True):
                        if st.session_state.get('latest_wechat_html'):
                            show_wechat_publish_dialog()
                        else:
                            st.warning("暂无已美化的推文 HTML 样式，请重新运行分析。")
                
                if st.session_state.get('wechat_long_image') and os.path.exists(st.session_state['wechat_long_image']):
                    st.write("---")
                    st.subheader("🖼️ 已生成的推文长图")
                    st.image(st.session_state['wechat_long_image'], use_container_width=True)
                    with open(st.session_state['wechat_long_image'], "rb") as img_file:
                        st.download_button(
                            label="📥 下载超长排版图",
                            data=img_file.read(),
                            file_name=os.path.basename(st.session_state['wechat_long_image']),
                            mime="image/png",
                            use_container_width=True
                        )
                
                st.write("---")
                st.subheader("📱 排版预览 (微信内置视口)")
                
                if st.session_state.get('latest_wechat_html'):
                    import streamlit.components.v1 as components
                    components.html(st.session_state['latest_wechat_html'], height=800, scrolling=True)
                else:
                    st.markdown(
                        """
                        <style>
                        .wechat-container {
                            max-width: 480px;
                            margin: 16px auto;
                            background-color: white;
                            border-radius: 12px;
                            padding: 24px 16px;
                            box-shadow: 0 4px 12px rgba(0,0,0,0.1);
                            color: #333;
                            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                        }
                        .wechat-container img {
                            width: 100%;
                            border-radius: 8px;
                            margin: 16px 0;
                        }
                        .wechat-container h1, .wechat-container h2, .wechat-container h3 {
                            color: #1a1a1a;
                        }
                        .wechat-container blockquote {
                            border-left: 4px solid #07c160;
                            background: #f7f7f7;
                            margin: 16px 0;
                            padding: 12px;
                            color: #666;
                            font-size: 0.9em;
                        }
                        </style>
                        """, unsafe_allow_html=True
                    )
                    st.markdown('<div class="wechat-container">', unsafe_allow_html=True)
                    render_wechat_preview(st.session_state['latest_wechat'])
                    st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.info("暂无生成的公众号推文。请在左侧勾选「生成微信公众号推文 (跳过常规图表渲染加速)」并开始深度分析。")
            
            st.write("---")
            st.subheader("📅 历史推文归档")
            hist_mds = sorted(glob.glob(os.path.join(WECHAT_OUTPUT_DIR, "*.md")), reverse=True)
            if not hist_mds: st.info("暂无历史记录")
            else:
                d_groups = {}
                for h_md in hist_mds:
                    try:
                        ds = os.path.basename(h_md).split('_')[1][:8]
                        dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                    except: dfmt = "其他"
                    if dfmt not in d_groups: d_groups[dfmt] = []
                    d_groups[dfmt].append(h_md)
                
                for d_key, d_mds in d_groups.items():
                    with st.expander(f"📅 {d_key} ({len(d_mds)}篇)"):
                        for single_md in d_mds:
                            col_a, col_b = st.columns([3, 1])
                            try:
                                with open(single_md, "r", encoding="utf-8") as f:
                                    content = f.read()
                                    title = content.split('\n')[0].replace('#', '').strip()[:20]
                                    if not title: title = "无标题"
                            except: title = "读取失败"
                            
                            col_a.write(f"📄 {title}...")
                            if col_b.button("查看", key=f"view_{os.path.basename(single_md)}"):
                                st.session_state['latest_wechat'] = content
                                st.session_state['latest_wechat_file_path'] = single_md
                                if 'wechat_long_image' in st.session_state:
                                    del st.session_state['wechat_long_image']
                                # 加载历史文章时重置内存中的微信草稿状态，弹窗会自动重新读取对应文件的 .draft.json
                                st.session_state.wechat_draft_media_id = None
                                st.session_state.wechat_draft_url = None
                                st.session_state.wechat_publish_result = None
                                st.session_state.wechat_preview_status = None
                                    
                                html_p = single_md.replace(".md", ".html")
                                if os.path.exists(html_p):
                                    with open(html_p, "r", encoding="utf-8") as hf:
                                        st.session_state['latest_wechat_html'] = hf.read()
                                    st.session_state['just_loaded_history'] = True
                                else:
                                    st.session_state['latest_wechat_html'] = None
                                    st.session_state['just_loaded_history'] = False
                                    
                                st.rerun()

elif page_selection == "视频脚本制作器":
    st.title("🎥 视频脚本制作器")
    
    col1, col2 = st.columns([2, 1.3])
    with col1:
        st.subheader("📁 历史脚本库 (供AI学习风格/续写)")
        uploaded_files = st.file_uploader("支持多文件上传 (.docx, .md, .xlsx, .csv)", type=['docx', 'md', 'xlsx', 'csv'], accept_multiple_files=True)
        
        if st.session_state.virtual_history:
            st.markdown("**📌 已追加的虚拟历史脚本:**")
            for vf in st.session_state.virtual_history:
                st.write(f"📄 `{vf['name']}`")
            if st.button("🗑️ 清空虚拟历史"):
                st.session_state.virtual_history = []
                st.rerun()
        
        st.markdown("---")
        st.subheader("⚙️ 生成设置")
        script_mode = st.radio("生成模式", ["仿写现有格式生成新脚本", "根据历史序列向后发散续写（例如基于0,1,2,3,4续写5,6）"], horizontal=False, key="vs_script_mode")
        export_format = st.selectbox("导出格式", [".docx", ".md"], key="vs_export_format")
        
        st.markdown("---")
        st.subheader("📝 素材与提示词")
        prompt_input = st.text_area("输入新的核心观点、素材内容或具体要求...", height=150, placeholder="例如：今天我们来讲一下人工智能在医疗领域的最新应用，重点突出AI辅助诊断的高效性...", key="vs_prompt_input")
        
        if st.button("🚀 开始生成脚本", use_container_width=True, type="primary"):
            if not current_chan_config.get("api_key"):
                st.error("请先在左侧全局配置中填写 API Key！")
                st.stop()
                
            with st.status("正在启动脚本制作流程...", expanded=True) as status:
                st.write("🔍 分析历史脚本与素材...")
                
                history_texts = []
                for f in uploaded_files:
                    txt = parse_uploaded_file(f)
                    if txt: history_texts.append(f"【历史脚本：{f.name}】\n{txt}")
                for vf in st.session_state.virtual_history:
                    history_texts.append(f"【历史脚本（追加）：{vf['name']}】\n{vf['text']}")
                history_context = "\n\n".join(history_texts)
                
                system_prompt = "你是一个专业的金融/交易类视频脚本编导。请严格学习并模仿用户提供的历史脚本的文案风格、语气（如口语化、设问式）和排版格式。\n\n【重要排版指令】：请必须使用 Markdown 语法进行排版输出。为了作为提词器使用时的重音提示，请务必对文案中的核心观点、金句或转折词使用**加粗**（如 `**重点内容**`）或引用块（如 `> 核心金句`）进行高亮。"
                
                user_content = ""
                if history_context:
                    user_content += f"以下是你需要学习参考的历史脚本序列：\n\n{history_context}\n\n====================\n\n"
                
                if "仿写" in script_mode:
                    user_content += f"请根据以上提供的历史脚本风格进行学习。你的核心任务是**直接生成一篇全新的完整视频脚本（最终提词器念稿版本）**，绝对不要只输出分析总结，也不要向我提问索要素材。\n"
                    if prompt_input.strip():
                        user_content += f"\n[新素材与要求]：\n{prompt_input}\n\n请结合上述新素材生成脚本内容。"
                    else:
                        user_content += "\n由于用户没有提供新素材，请自行发挥你的专业金融编导水平，拟定一个符合当前市场热点的主题，直接撰写出这篇完整的视频脚本。"
                else:
                    user_content += f"请深度分析前面提供的历史脚本序列的故事线、知识递进逻辑和表达手法，自动推理出下一期的主题，并**直接【顺延生成】最新一期的完整视频脚本内容（最终提词器念稿版本）**。要求保持一贯的口语化、设问式风格。绝对不要向我提问索要素材，也绝对不要仅仅输出风格特征分析表！\n"
                    if prompt_input.strip():
                        user_content += f"\n在续写时，请必须结合以下新素材或要求：\n[新素材与要求]：\n{prompt_input}"
                    else:
                        user_content += "\n请注意：由于我没有提供新素材，请直接根据历史上下文的逻辑推演，自动拟定下一期主题并生成完整的视频文案！"
                
                st.write(f"🧠 AI 正在构思与生成 (模型: {current_chan_config.get('selected_model')})...")
                
                try:
                    client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                    response = call_chat_completion(
                        client,
                        current_chan_config.get("selected_model"),
                        [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_content}
                        ],
                        chan_config=current_chan_config
                    )
                    if not response.choices:
                        raise Exception(f"接口返回空数据，可能是该模型暂不支持或网络限流。详情: {response}")
                    script_content = response.choices[0].message.content
                    st.session_state['generated_script_preview'] = script_content
                    st.session_state['export_format_choice'] = export_format
                    
                    # 保存到历史归档
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    save_path = os.path.join(SCRIPT_OUTPUT_DIR, f"script_{timestamp}.md")
                    with open(save_path, "w", encoding="utf-8") as f:
                        f.write(script_content)
                        
                    status.update(label="脚本生成完成！", state="complete", expanded=False)
                    st.success("生成成功！请在右侧预览并下载。")
                except Exception as e:
                    status.update(label="任务异常终止", state="error")
                    st.error(f"脚本生成失败: {str(e)}")

    with col2:
        st.subheader("📺 生成结果预览")
        
        preview_text = st.session_state.get('generated_script_preview', '')
        if preview_text:
            styled_html = f"""
            <style>
                .script-preview-container {{
                    font-size: 16px;
                    line-height: 1.8;
                    color: #333;
                }}
                .script-preview-container strong {{
                    color: #d97706; /* 深橙色/主题色 */
                    font-weight: 900;
                    background-color: #fef3c7;
                    padding: 0 4px;
                    border-radius: 3px;
                }}
                .script-preview-container blockquote {{
                    border-left: 5px solid #07c160; /* 主题绿 */
                    padding: 12px 15px;
                    margin: 15px 0;
                    color: #4b5563;
                    background-color: #f9fafb;
                    border-radius: 0 8px 8px 0;
                    font-style: italic;
                }}
            </style>
            """
            st.markdown(styled_html, unsafe_allow_html=True)
            
            with st.container(border=True):
                st.markdown(f'<div class="script-preview-container">{markdown.markdown(preview_text, extensions=["extra", "nl2br"])}</div>', unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("💾 导出与下载")
            
            # Real download logic for generated scripts
            ext = st.session_state.get('export_format_choice', '.md')
            file_name = f"视频脚本_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
            
            if ext == ".docx":
                from io import BytesIO
                from docx import Document
                from docx.shared import RGBColor
                from docx.enum.text import WD_COLOR_INDEX
                import re
                
                doc = Document()
                for line in preview_text.split('\n'):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                        continue
                    
                    if line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                        continue
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                        continue
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                        continue
                    
                    is_quote = line.startswith('> ')
                    is_bullet = line.startswith('- ') or line.startswith('* ')
                    
                    if is_bullet:
                        p = doc.add_paragraph(style='List Bullet')
                        text = line[2:]
                    else:
                        p = doc.add_paragraph()
                        text = line[2:] if is_quote else line
                        
                    # 解析并映射加粗与颜色
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                            run.font.color.rgb = RGBColor(217, 119, 6) # 对应 CSS 的深橙色 #d97706
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW # 黄色背景高亮
                        elif part:
                            run = p.add_run(part)
                            if is_quote:
                                run.font.color.rgb = RGBColor(7, 193, 96) # 对应 CSS 的主题绿 #07c160
                                run.italic = True
                                
                bio = BytesIO()
                doc.save(bio)
                file_data = bio.getvalue()
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                file_data = preview_text.encode('utf-8')
                mime_type = "text/plain"
            
            st.download_button(
                label=f"⬇️ 下载 {ext} 文件",
                data=file_data,
                file_name=file_name,
                mime=mime_type,
                use_container_width=True
            )
            
            st.markdown("---")
            if st.button("➕ 将此篇加入参考库，继续生成下一期", use_container_width=True):
                next_index = len(st.session_state.virtual_history) + 1
                v_name = f"新生成_追加_{next_index}.md"
                st.session_state.virtual_history.append({"name": v_name, "text": preview_text})
                st.success(f"已成功加入左侧参考库：{v_name}")
                time.sleep(1)
                st.rerun()
        else:
            st.info("暂无生成内容，请先在左侧输入素材并点击开始生成。")
            
        st.write("---")
        st.subheader("📅 历史生成归档")
        hist_scripts = sorted(glob.glob(os.path.join(SCRIPT_OUTPUT_DIR, "*.md")), reverse=True)
        if not hist_scripts:
            st.info("暂无历史记录")
        else:
            d_groups = {}
            for h_sc in hist_scripts:
                try:
                    ds = os.path.basename(h_sc).split('_')[1]
                    dfmt = f"{ds[:4]}-{ds[4:6]}-{ds[6:8]}"
                except:
                    dfmt = "其他"
                if dfmt not in d_groups: d_groups[dfmt] = []
                d_groups[dfmt].append(h_sc)
            
            for d_key, d_scs in d_groups.items():
                with st.expander(f"📅 {d_key} ({len(d_scs)}份)"):
                    for single_sc in d_scs:
                        sc_name = os.path.basename(single_sc)
                        if st.button(f"📄 {sc_name}", key=f"hist_{single_sc}"):
                            with open(single_sc, "r", encoding="utf-8") as f:
                                st.session_state['generated_script_preview'] = f.read()
                            st.session_state['export_format_choice'] = '.md'
                            st.session_state['show_success_toast'] = "视频脚本生成完成！请在右侧预览"
                            st.rerun()

elif page_selection == "指标文档制作":
    st.title("📈 指标文档制作")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("⚙️ 1. 指标库管理")
        indicators = load_indicators()
        ind_options = ["✨ [新建空白指标]"] + list(indicators.keys())
        
        # Maintain selection state
        if "ind_selected" not in st.session_state: st.session_state["ind_selected"] = ind_options[0]
        sel_ind = st.selectbox("选择或新建指标", ind_options, index=ind_options.index(st.session_state["ind_selected"]) if st.session_state["ind_selected"] in ind_options else 0)
        
        if sel_ind != st.session_state["ind_selected"]:
            st.session_state["ind_selected"] = sel_ind
            st.rerun()
        
        if sel_ind == "✨ [新建空白指标]":
            def_name = ""
            def_code = ""
        else:
            def_name = sel_ind
            def_code = indicators[sel_ind]["code"]
            
        new_ind_name = st.text_input("指标名称", value=def_name, placeholder="例如：震荡顶底模型")
        new_ind_code = st.text_area("指标源码", value=def_code, height=200, placeholder="在此粘贴 Pine Script 或 Python 源码...")
        
        c_btn1, c_btn2 = st.columns(2)
        with c_btn1:
            if st.button("💾 保存/更新", use_container_width=True):
                if new_ind_name.strip() and new_ind_code.strip():
                    save_indicator(new_ind_name.strip(), new_ind_code.strip())
                    st.session_state["ind_selected"] = new_ind_name.strip()
                    st.success(f"指标 '{new_ind_name}' 已保存！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("名称和源码不能为空！")
        with c_btn2:
            if st.button("🗑️ 删除", use_container_width=True):
                if sel_ind != "✨ [新建空白指标]":
                    delete_indicator(sel_ind)
                    st.session_state["ind_selected"] = "✨ [新建空白指标]"
                    st.success(f"指标 '{sel_ind}' 已删除！")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.warning("新建状态不可删除")
                    
        st.markdown("---")
        st.subheader("⚡ 2. 生成分析文档")
        
        if not indicators:
            st.info("暂无已保存的指标，请先保存指标后再生成。")
        else:
            gen_ind = st.selectbox("选择目标指标", list(indicators.keys()), index=list(indicators.keys()).index(st.session_state["ind_selected"]) if st.session_state["ind_selected"] in indicators else 0)
            
            if st.button("🚀 生成标准合规版", type="primary", use_container_width=True):
                if not current_chan_config.get("api_key"): st.error("请先在左侧全局配置中填写 API Key！")
                else:
                    with st.status("正在生成标准合规分析文档...", expanded=True) as status:
                        st.write(f"🧠 AI 正在分析 {gen_ind} 的源码...")
                        user_content = f"【指标名称】：{gen_ind}\n【指标源码】：\n{indicators[gen_ind]['code']}"
                        try:
                            client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                            response = call_chat_completion(
                                client,
                                current_chan_config.get("selected_model"),
                                [
                                    {"role": "system", "content": PROMPT_INDICATOR_STANDARD},
                                    {"role": "user", "content": user_content}
                                ],
                                chan_config=current_chan_config
                            )
                            if not response.choices: raise Exception("接口返回空数据")
                            doc_content = response.choices[0].message.content
                            
                            ts = datetime.now().strftime("%Y%m%d%H%M%S")
                            md_filename = f"《{gen_ind}》---合规化_{ts}.md"
                            docx_filename = f"《{gen_ind}》---合规化_{ts}.docx"
                            md_path = os.path.join(INDICATOR_DOCS_DIR, md_filename)
                            docx_path = os.path.join(INDICATOR_DOCS_DIR, docx_filename)
                            
                            with open(md_path, "w", encoding="utf-8") as f: f.write(doc_content)
                            markdown_to_docx_file(doc_content, docx_path, indicator_name=gen_ind)
                            
                            st.session_state["ind_preview_title"] = f"{gen_ind} (标准合规版)"
                            st.session_state["ind_preview_content"] = doc_content
                            st.session_state["ind_preview_docx"] = docx_path
                            
                            status.update(label="标准文档生成完成！", state="complete", expanded=False)
                            st.rerun()
                        except Exception as e:
                            status.update(label="生成失败", state="error")
                            st.error(f"报错信息: {str(e)}")
                            
            if st.button("✨ 基于标准版一键转化【社群互动教学】版", use_container_width=True):
                if not current_chan_config.get("api_key"): st.error("请先填写 API Key！")
                elif "ind_preview_content" not in st.session_state or "合规化" not in st.session_state.get("ind_preview_title", ""):
                    st.warning("请先生成或从右侧历史预览一份【标准合规版】文档，再进行转化！")
                else:
                    with st.status("正在进行社群风格转换...", expanded=True) as status2:
                        try:
                            client = OpenAI(api_key=current_chan_config.get("api_key"), base_url=current_chan_config.get("base_url") if current_chan_config.get("base_url") else None)
                            response2 = call_chat_completion(
                                client,
                                current_chan_config.get("selected_model"),
                                [
                                    {"role": "system", "content": PROMPT_INDICATOR_COMMUNITY},
                                    {"role": "user", "content": f"请将以下标准技术文档转化为社群教学版本：\n\n{st.session_state['ind_preview_content']}"}
                                ],
                                chan_config=current_chan_config
                            )
                            if not response2.choices: raise Exception("接口返回空数据")
                            comm_doc = response2.choices[0].message.content
                            
                            ts = datetime.now().strftime("%Y%m%d%H%M%S")
                            md_filename = f"《{gen_ind}》--社群指标互动手册_{ts}.md"
                            docx_filename = f"《{gen_ind}》--社群指标互动手册_{ts}.docx"
                            md_path = os.path.join(INDICATOR_DOCS_DIR, md_filename)
                            docx_path = os.path.join(INDICATOR_DOCS_DIR, docx_filename)
                            
                            with open(md_path, "w", encoding="utf-8") as f: f.write(comm_doc)
                            markdown_to_docx_file(comm_doc, docx_path, indicator_name=gen_ind)
                            
                            st.session_state["ind_preview_title"] = f"{gen_ind} (社群互动教学版)"
                            st.session_state["ind_preview_content"] = comm_doc
                            st.session_state["ind_preview_docx"] = docx_path
                            
                            status2.update(label="社群文档转换完成！", state="complete", expanded=False)
                            st.rerun()
                        except Exception as e:
                            status2.update(label="转换失败", state="error")
                            st.error(f"报错信息: {str(e)}")

    with col2:
        st.subheader("👁️ 3. 预览与归档")
        
        preview_title = st.session_state.get("ind_preview_title", "预览区")
        preview_content = st.session_state.get("ind_preview_content", "")
        
        with st.container(border=True):
            st.markdown(f"#### {preview_title}")
            if preview_content:
                st.markdown(preview_content)
                docx_path = st.session_state.get("ind_preview_docx")
                if docx_path and os.path.exists(docx_path):
                    with open(docx_path, "rb") as f: docx_data = f.read()
                    st.download_button(
                        label="⬇️ 下载该 Docx 文档",
                        data=docx_data,
                        file_name=os.path.basename(docx_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        type="primary"
                    )
            else:
                st.info("右侧暂无预览，请在左侧操作生成或从历史归档中选择。")
                
        st.markdown("---")
        st.write("### 📅 历史报告归档")
        hist_docs = sorted(glob.glob(os.path.join(INDICATOR_DOCS_DIR, "*.md")), reverse=True)
        if not hist_docs:
            st.info("暂无生成的文档归档。")
        else:
            # 按日期分组
            grouped_files = {}
            for h_md in hist_docs:
                basename = os.path.basename(h_md)
                match = re.search(r'_(\d{8})\d{6}\.md$', basename)
                if match:
                    date_str = datetime.strptime(match.group(1), "%Y%m%d").strftime("%Y-%m-%d")
                else:
                    mtime = os.path.getmtime(h_md)
                    date_str = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
                
                if date_str not in grouped_files:
                    grouped_files[date_str] = []
                grouped_files[date_str].append(h_md)
            
            # 日期下拉菜单
            date_list = sorted(grouped_files.keys(), reverse=True)
            selected_date = st.selectbox("选择日期", date_list, index=0, key="hist_date_select")
            
            # 仅展示所选日期的文件
            if selected_date in grouped_files:
                for h_md in grouped_files[selected_date]:
                    basename = os.path.basename(h_md)
                    with st.expander(f"📄 {basename}"):
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("👁️ 在上方预览", key=f"prev_{h_md}", use_container_width=True):
                                with open(h_md, "r", encoding="utf-8") as f: content = f.read()
                                docx_p = h_md.replace(".md", ".docx")
                                st.session_state["ind_preview_title"] = basename.replace(".md", "")
                                st.session_state["ind_preview_content"] = content
                                st.session_state["ind_preview_docx"] = docx_p if os.path.exists(docx_p) else ""
                                st.rerun()
                        with c2:
                            docx_p = h_md.replace(".md", ".docx")
                            if os.path.exists(docx_p):
                                with open(docx_p, "rb") as f: db = f.read()
                                st.download_button(
                                    label="⬇️ 下载 Docx",
                                    data=db,
                                    file_name=os.path.basename(docx_p),
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_{docx_p}",
                                    use_container_width=True
                                )

